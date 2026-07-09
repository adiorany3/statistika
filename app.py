import io
import json
import re
import tempfile
import textwrap
import warnings
from datetime import datetime

import numpy as np
import pandas as pd
import streamlit as st
from scipy import stats

warnings.filterwarnings("ignore")

try:
    import plotly.express as px
except Exception:  # pragma: no cover
    px = None

try:
    import statsmodels.api as sm
    import statsmodels.formula.api as smf
    from statsmodels.stats.multicomp import pairwise_tukeyhsd
    from statsmodels.stats.diagnostic import het_breuschpagan
    from statsmodels.stats.outliers_influence import variance_inflation_factor
    from statsmodels.stats.multitest import multipletests
except Exception:  # pragma: no cover
    sm = None
    smf = None
    het_breuschpagan = None
    pairwise_tukeyhsd = None
    variance_inflation_factor = None
    multipletests = None

def patch_sklearn_check_array_for_factor_analyzer():
    """Patch kompatibilitas factor-analyzer dengan scikit-learn terbaru.

    factor-analyzer 0.5.x masih dapat memanggil check_array(force_all_finite=...),
    sementara scikit-learn baru memakai ensure_all_finite=.... Patch ini dipasang
    di sklearn.utils.validation, sklearn.utils, dan referensi internal factor_analyzer
    jika modulnya sudah terlanjur di-import.
    """
    try:
        import inspect
        import sklearn.utils as _sk_utils
        import sklearn.utils.validation as _sk_validation

        current = _sk_validation.check_array
        original = getattr(current, "_statpro_original_check_array", current)
        params = inspect.signature(original).parameters

        if "force_all_finite" not in params and "ensure_all_finite" in params:
            def _check_array_compat(*args, force_all_finite=None, ensure_all_finite=None, **kwargs):
                if ensure_all_finite is None and force_all_finite is not None:
                    ensure_all_finite = force_all_finite
                if ensure_all_finite is not None:
                    kwargs["ensure_all_finite"] = ensure_all_finite
                return original(*args, **kwargs)

            _check_array_compat._statpro_original_check_array = original
            _sk_validation.check_array = _check_array_compat
            _sk_utils.check_array = _check_array_compat

            # Jika factor_analyzer sudah meng-cache check_array saat import, ganti juga.
            try:
                import factor_analyzer.factor_analyzer as _fa_module
                if hasattr(_fa_module, "check_array"):
                    _fa_module.check_array = _check_array_compat
            except Exception:
                pass
            try:
                import factor_analyzer.utils as _fa_utils
                if hasattr(_fa_utils, "check_array"):
                    _fa_utils.check_array = _check_array_compat
            except Exception:
                pass
        return True
    except Exception:
        return False


try:
    from sklearn.preprocessing import StandardScaler
    from sklearn.decomposition import PCA, FactorAnalysis
    patch_sklearn_check_array_for_factor_analyzer()
except Exception:  # pragma: no cover
    StandardScaler = None
    PCA = None
    FactorAnalysis = None

try:
    import pyreadstat
except Exception:  # pragma: no cover
    pyreadstat = None

try:
    import scikit_posthocs as sp_posthoc
except Exception:  # pragma: no cover
    sp_posthoc = None

try:
    import pingouin as pg
except Exception:  # pragma: no cover
    pg = None

try:
    patch_sklearn_check_array_for_factor_analyzer()
    from factor_analyzer import FactorAnalyzer, calculate_bartlett_sphericity, calculate_kmo
    patch_sklearn_check_array_for_factor_analyzer()
except Exception:  # pragma: no cover
    FactorAnalyzer = None
    calculate_bartlett_sphericity = None
    calculate_kmo = None

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None


APP_NAME = "Statistik Pro+"
APP_SUBTITLE = "v5.0 Research Analytics Suite — alternatif SPSS berbasis Streamlit dengan UI terpandu, analisis lanjutan, dan interpretasi riset berbahasa Indonesia."

st.set_page_config(page_title=f"{APP_NAME} - Alternatif SPSS", page_icon="📊", layout="wide")


# -----------------------------------------------------------------------------
# Streamlit stability patch
# -----------------------------------------------------------------------------
def _install_auto_widget_keys():
    """Berikan key otomatis yang stabil untuk widget Streamlit tanpa key eksplisit.

    Catatan perbaikan v3.7:
    - st.sidebar/st.columns/st.tabs memakai DeltaGenerator; argumen pertama method adalah
      objek container, bukan label widget. Versi lama salah membaca ini sehingga beberapa
      widget sidebar mendapat key yang sama.
    - Wrapper dipasang ulang setiap rerun dengan mengembalikan fungsi asli via __wrapped__,
      sehingga counter key selalu bersih per-run.
    - Jika ada beberapa widget identik di baris yang sama, suffix counter ditambahkan secara
      stabil berdasarkan urutan render pada run tersebut.
    """
    try:
        import functools
        import hashlib
        import inspect
        from collections import defaultdict
        from streamlit.delta_generator import DeltaGenerator

        widget_names = [
            "button", "download_button", "checkbox", "toggle", "radio", "selectbox",
            "multiselect", "slider", "select_slider", "number_input", "text_input",
            "text_area", "file_uploader", "data_editor", "date_input", "time_input",
            "color_picker", "form_submit_button",
        ]

        # Dibuat baru setiap script run supaya key stabil antar-rerun.
        seen_this_run = defaultdict(int)
        internal_names = {"make_key", "wrapped", "wrap_function", "_install_auto_widget_keys"}

        def unwrap_auto(func):
            """Ambil fungsi Streamlit asli dari wrapper lama, bila ada."""
            original = getattr(func, "_statpro_auto_key_original", None)
            if original is not None:
                return original
            # functools.wraps menyimpan fungsi asli di __wrapped__.
            while getattr(func, "_statpro_auto_key_wrapped", False) and hasattr(func, "__wrapped__"):
                func = func.__wrapped__
            return func

        def find_user_caller():
            for frameinfo in inspect.stack(context=0):
                filename = frameinfo.filename or ""
                if not filename.endswith("app.py"):
                    continue
                if frameinfo.function in internal_names:
                    continue
                return frameinfo
            return None

        def get_widget_label(kind, args, kwargs):
            if "label" in kwargs:
                return str(kwargs.get("label", ""))[:160]
            # Untuk DeltaGenerator method: args[0] adalah container/sidebar/column/tab,
            # label widget biasanya args[1]. Untuk st.* top-level, label biasanya args[0].
            offset = 1 if kind.startswith("dg.") else 0
            if len(args) > offset:
                try:
                    return str(args[offset])[:160]
                except Exception:
                    return ""
            return ""

        def make_key(kind, args, kwargs):
            caller = find_user_caller()
            if caller is None:
                location = "unknown:0"
                function = "unknown"
            else:
                location = f"{caller.filename}:{caller.lineno}"
                function = caller.function

            label = get_widget_label(kind, args, kwargs)
            raw_base = f"{kind}|{location}|{function}|{label}"
            seen_this_run[raw_base] += 1
            occurrence = seen_this_run[raw_base]
            raw = raw_base if occurrence == 1 else f"{raw_base}|occurrence={occurrence}"
            return "auto_" + hashlib.md5(raw.encode("utf-8", errors="ignore")).hexdigest()[:16]

        def wrap_function(func, kind):
            original = unwrap_auto(func)

            @functools.wraps(original)
            def wrapped(*args, **kwargs):
                if kwargs.get("key") is None:
                    kwargs["key"] = make_key(kind, args, kwargs)
                return original(*args, **kwargs)

            wrapped._statpro_auto_key_wrapped = True
            wrapped._statpro_auto_key_original = original
            return wrapped

        for name in widget_names:
            if hasattr(st, name):
                setattr(st, name, wrap_function(getattr(st, name), f"st.{name}"))
            if hasattr(DeltaGenerator, name):
                setattr(DeltaGenerator, name, wrap_function(getattr(DeltaGenerator, name), f"dg.{name}"))
        return True
    except Exception:
        return False


_install_auto_widget_keys()


def _is_streamlit_control_exception(exc):
    """Jangan tangkap exception internal untuk st.rerun()/st.stop()."""
    name = exc.__class__.__name__.lower()
    return "rerun" in name or "stop" in name

st.markdown(
    """
    <style>
        footer {visibility: hidden;}
        .stDeployButton {display: none;}
        #MainMenu {visibility: hidden;}
        .block-container {padding-top: 1.3rem; padding-bottom: 5.2rem;}
        .statpro-footer-fixed {
            position: fixed;
            left: 0;
            right: 0;
            bottom: 0;
            z-index: 999999;
            padding: 0.62rem 1rem;
            text-align: center;
            color: #475569;
            font-size: 0.86rem;
            line-height: 1.25rem;
            background: rgba(255, 255, 255, 0.96);
            border-top: 1px solid #e5e7eb;
            box-shadow: 0 -4px 16px rgba(15, 23, 42, 0.06);
            backdrop-filter: blur(8px);
        }
        .statpro-footer-fixed strong {color:#0f172a;}
        .statpro-footer-spacer {height: 3.2rem;}
        .statpro-sidebar-footer {
            margin-top: 1rem;
            padding: .75rem .65rem;
            border-radius: .75rem;
            background: #f8fafc;
            border: 1px solid #e5e7eb;
            color: #64748b;
            font-size: .82rem;
            text-align: center;
        }
        .small-note {color: #6b7280; font-size: 0.92rem;}
        .stat-card {
            padding: 1rem; border-radius: 0.9rem; border: 1px solid #e5e7eb;
            background: linear-gradient(180deg, #ffffff 0%, #f8fafc 100%);
            min-height: 96px;
        }
        .guide-card {
            padding: 1rem; border-radius: 1rem; border: 1px solid #dbeafe;
            background: #eff6ff; margin-bottom: 0.7rem;
        }
        .soft-card {
            padding: 1rem; border-radius: 1rem; border: 1px solid #e5e7eb;
            background: #ffffff; margin-bottom: 0.7rem;
        }
        .ok-box {padding: .75rem; border-radius: .8rem; background:#ecfdf5; border:1px solid #bbf7d0;}
        .warn-box {padding: .75rem; border-radius: .8rem; background:#fffbeb; border:1px solid #fde68a;}
        .danger-box {padding: .75rem; border-radius: .8rem; background:#fef2f2; border:1px solid #fecaca;}
        .tiny {font-size: 0.86rem; color: #6b7280;}
        .step-pill {display:inline-block; padding: .25rem .55rem; border-radius:999px; background:#f3f4f6; margin:.15rem; font-size:.85rem;}
    </style>
    """,
    unsafe_allow_html=True,
)


APP_VERSION = "v5.2"
FOOTER_TEXT = "Developed by Galuh Adi Insani · Statistik Pro+ v5.2 · Research Analytics Suite"


def render_persistent_footer():
    """Footer tetap terlihat walau halaman panjang, rerun, atau menu aktif berubah."""
    st.markdown(
        f"""
        <div class="statpro-footer-fixed">
            <strong>Developed by Galuh Adi Insani</strong> · Statistik Pro+ {APP_VERSION} · Research Analytics Suite
        </div>
        """,
        unsafe_allow_html=True,
    )




# Render sejak awal agar footer tetap muncul meskipun menu tertentu mengalami error lokal.
render_persistent_footer()


def init_state():
    defaults = {
        "df": None,
        "raw_df": None,
        "file_name": None,
        "report_items": [],
        "metadata": None,
        "syntax_log": [],
        "split_by": "(tidak ada)",
        "active_alpha": 0.05,
        "last_action": None,
        "ui_mode": "Pemula",
        "detail_level": "Ringkas",
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


init_state()


def p_value_text(p):
    if pd.isna(p):
        return "NA"
    return "< 0.001" if p < 0.001 else f"{p:.4f}"


def decision_text(p, alpha):
    if pd.isna(p):
        return "Tidak dapat dihitung"
    return "Signifikan — tolak H₀" if p < alpha else "Tidak signifikan — gagal tolak H₀"


def numeric_cols(df):
    return df.select_dtypes(include=[np.number]).columns.tolist()


def categorical_cols(df):
    return [c for c in df.columns if c not in numeric_cols(df)]


def infer_measure(series):
    """Infer measurement level ala SPSS: nominal, ordinal, atau scale."""
    s = series.dropna()
    if s.empty:
        return "Nominal"
    if pd.api.types.is_numeric_dtype(series):
        unique_count = s.nunique(dropna=True)
        if unique_count <= 10 and np.allclose(pd.to_numeric(s, errors="coerce").dropna() % 1, 0):
            return "Ordinal"
        return "Scale"
    return "Nominal"


def build_metadata(df):
    rows = []
    for col in df.columns:
        rows.append(
            {
                "Name": col,
                "Label": col.replace("_", " ").title(),
                "Type": str(df[col].dtype),
                "Measure": infer_measure(df[col]),
                "Role": "Input",
                "Value Labels": "",
                "Missing Values": "",
                "Decimals": 2 if pd.api.types.is_numeric_dtype(df[col]) else 0,
            }
        )
    return pd.DataFrame(rows)


def sync_metadata(df):
    existing = st.session_state.get("metadata")
    fresh = build_metadata(df)
    if existing is None or not isinstance(existing, pd.DataFrame) or existing.empty:
        st.session_state.metadata = fresh
        return fresh
    old = existing.set_index("Name", drop=False)
    rows = []
    for _, row in fresh.iterrows():
        name = row["Name"]
        if name in old.index:
            merged = row.to_dict()
            for key in ["Label", "Measure", "Role", "Value Labels", "Missing Values", "Decimals"]:
                if key in old.columns:
                    merged[key] = old.loc[name, key]
            merged["Type"] = row["Type"]
            rows.append(merged)
        else:
            rows.append(row.to_dict())
    st.session_state.metadata = pd.DataFrame(rows)
    return st.session_state.metadata


def parse_missing_values(text):
    if pd.isna(text) or str(text).strip() == "":
        return []
    values = []
    for token in re.split(r"[,;|]", str(text)):
        token = token.strip()
        if token == "":
            continue
        try:
            values.append(float(token) if "." in token else int(token))
        except ValueError:
            values.append(token)
    return values


def parse_value_labels(text):
    """Format: 1=Laki-laki; 2=Perempuan"""
    mapping = {}
    if pd.isna(text) or str(text).strip() == "":
        return mapping
    for part in str(text).split(";"):
        if "=" in part:
            key, val = part.split("=", 1)
            key = key.strip()
            val = val.strip()
            try:
                key_obj = float(key) if "." in key else int(key)
            except ValueError:
                key_obj = key
            mapping[key_obj] = val
            mapping[str(key_obj)] = val
    return mapping


def apply_metadata_to_data(df, metadata):
    new_df = df.copy()
    if metadata is None or metadata.empty:
        return new_df
    for _, row in metadata.iterrows():
        col = row.get("Name")
        if col not in new_df.columns:
            continue
        missings = parse_missing_values(row.get("Missing Values", ""))
        if missings:
            new_df[col] = new_df[col].replace(missings, np.nan)
    return new_df


def display_with_value_labels(df, metadata):
    view = df.copy()
    if metadata is None or metadata.empty:
        return view
    for _, row in metadata.iterrows():
        col = row.get("Name")
        if col not in view.columns:
            continue
        labels = parse_value_labels(row.get("Value Labels", ""))
        if labels:
            view[col] = view[col].map(lambda x: labels.get(x, labels.get(str(x), x)))
    return view


def log_syntax(command):
    st.session_state.syntax_log.append({"time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"), "command": command})


def safe_col_name(name, existing_cols):
    name = re.sub(r"[^0-9a-zA-Z_]+", "_", str(name).strip()) or "new_variable"
    if name[0].isdigit():
        name = f"v_{name}"
    base = name
    i = 1
    while name in existing_cols:
        i += 1
        name = f"{base}_{i}"
    return name


def recode_series(series, rules_text, default="copy"):
    """Rules: old=new, low:high=new, *=value. Satu aturan per baris."""
    result = series.copy() if default == "copy" else pd.Series([np.nan] * len(series), index=series.index)
    for raw_line in str(rules_text).splitlines():
        line = raw_line.strip()
        if not line or "=" not in line:
            continue
        left, right = [x.strip() for x in line.split("=", 1)]
        try:
            right_value = float(right) if re.fullmatch(r"[-+]?\d*\.\d+|[-+]?\d+", right) else right
        except Exception:
            right_value = right
        if left == "*":
            result.loc[:] = right_value
            continue
        if ":" in left:
            lo, hi = [x.strip() for x in left.split(":", 1)]
            numeric = pd.to_numeric(series, errors="coerce")
            lo_val = -np.inf if lo.lower() in ["low", "min", ""] else float(lo)
            hi_val = np.inf if hi.lower() in ["high", "max", ""] else float(hi)
            result.loc[numeric.between(lo_val, hi_val, inclusive="both")] = right_value
        else:
            try:
                left_value = float(left) if re.fullmatch(r"[-+]?\d*\.\d+|[-+]?\d+", left) else left
            except Exception:
                left_value = left
            result.loc[series.astype(str) == str(left_value)] = right_value
            result.loc[series == left_value] = right_value
    return result


def auto_interpretation(title, table, alpha=None):
    alpha = alpha if alpha is not None else st.session_state.get("active_alpha", 0.05)
    try:
        df_table = table if isinstance(table, pd.DataFrame) else pd.DataFrame(table)
    except Exception:
        return ""
    lower_cols = {str(c).lower(): c for c in df_table.columns}
    p_col = None
    for key in ["p-value", "p", "prob(f)", "llr p-value"]:
        if key in lower_cols:
            p_col = lower_cols[key]
            break
    if p_col is not None:
        pvals = pd.to_numeric(df_table[p_col], errors="coerce").dropna()
        if len(pvals):
            p = float(pvals.iloc[0])
            decision = "signifikan" if p < alpha else "tidak signifikan"
            return f"Interpretasi otomatis: hasil {title} {decision} pada α = {alpha:.2f} (p = {p_value_text(p)})."
    if "Cronbach's Alpha" in df_table.columns:
        val = pd.to_numeric(df_table["Cronbach's Alpha"], errors="coerce").dropna()
        if len(val):
            a = float(val.iloc[0])
            if a >= 0.90:
                level = "sangat tinggi"
            elif a >= 0.80:
                level = "baik"
            elif a >= 0.70:
                level = "dapat diterima"
            elif a >= 0.60:
                level = "cukup untuk eksplorasi"
            else:
                level = "rendah"
            return f"Interpretasi otomatis: reliabilitas internal {level} (Cronbach's α = {a:.3f})."
    if "R²" in df_table.columns:
        val = pd.to_numeric(df_table["R²"], errors="coerce").dropna()
        if len(val):
            return f"Interpretasi otomatis: model menjelaskan sekitar {val.iloc[0]*100:.1f}% variasi pada variabel dependen."
    return ""


def _first_existing_column(df_table, candidates):
    lower_cols = {str(c).lower().strip(): c for c in df_table.columns}
    for cand in candidates:
        key = str(cand).lower().strip()
        if key in lower_cols:
            return lower_cols[key]
    # fallback contains match, useful for CI/eigen labels
    for cand in candidates:
        needle = str(cand).lower().strip()
        for lower, original in lower_cols.items():
            if needle in lower:
                return original
    return None


def _first_numeric_value(df_table, candidates):
    col = _first_existing_column(df_table, candidates)
    if col is None:
        return None
    values = pd.to_numeric(df_table[col], errors="coerce").dropna()
    if values.empty:
        return None
    return float(values.iloc[0])


def _format_number(value, digits=3):
    try:
        if value is None or pd.isna(value):
            return "NA"
        value = float(value)
        if abs(value) < 0.001 and value != 0:
            return f"{value:.2e}"
        return f"{value:.{digits}f}"
    except Exception:
        return str(value)


def _effect_size_label(value, kind="generic"):
    try:
        v = abs(float(value))
    except Exception:
        return "tidak dapat dikategorikan"
    kind = str(kind).lower()
    if kind in ["d", "cohen", "cohen's d", "dz"]:
        if v < 0.20:
            return "sangat kecil"
        if v < 0.50:
            return "kecil"
        if v < 0.80:
            return "sedang"
        return "besar"
    if kind in ["eta", "eta2", "η²", "omega", "ω²"]:
        if v < 0.01:
            return "sangat kecil"
        if v < 0.06:
            return "kecil"
        if v < 0.14:
            return "sedang"
        return "besar"
    if kind in ["r", "corr", "correlation"]:
        if v < 0.10:
            return "sangat lemah"
        if v < 0.30:
            return "lemah"
        if v < 0.50:
            return "sedang"
        if v < 0.70:
            return "kuat"
        return "sangat kuat"
    if kind in ["r2", "r²", "pseudo r²"]:
        return f"menjelaskan sekitar {v*100:.1f}% variasi"
    return "perlu dibaca bersama konteks riset"


def _p_strength(p, alpha):
    try:
        p = float(p)
    except Exception:
        return "Bukti statistik tidak dapat dinilai karena p-value tidak tersedia."
    if p < 0.001:
        return "Bukti statistik sangat kuat terhadap H₀."
    if p < alpha:
        return "Bukti statistik cukup untuk menolak H₀ pada taraf signifikansi yang dipilih."
    if p < 0.10 and alpha <= 0.05:
        return "Ada kecenderungan/pola, tetapi belum cukup kuat pada α = 0.05."
    return "Belum ada bukti statistik yang cukup untuk menolak H₀."


def _table_as_long_text(df_table, max_rows=3):
    try:
        small = df_table.head(max_rows).copy().fillna("")
        return "; ".join([", ".join([f"{c}={row[c]}" for c in small.columns[:5]]) for _, row in small.iterrows()])
    except Exception:
        return ""


def insight_for_output_item(item, alpha=None):
    """Ubah satu output statistik menjadi insight riset yang lebih bermakna."""
    alpha = alpha if alpha is not None else st.session_state.get("active_alpha", 0.05)
    title = str(item.get("title", "Output"))
    title_low = title.lower()
    note = str(item.get("note", "") or "")
    table = item.get("table")
    if table is None:
        if note:
            return [{"Aspek": "Catatan", "Insight": note}]
        return [{"Aspek": "Catatan", "Insight": "Output ini tidak memiliki tabel yang dapat diringkas."}]
    try:
        df_table = table.copy() if isinstance(table, pd.DataFrame) else pd.DataFrame(table)
    except Exception:
        return [{"Aspek": "Catatan", "Insight": "Tabel output tidak dapat dibaca untuk insight otomatis."}]

    insights = []
    p = _first_numeric_value(df_table, ["p-value", "p", "Prob(F)", "LLR p-value", "Bartlett p-value"])

    if p is not None:
        if "normal" in title_low or "jarque" in title_low:
            meaning = "Data/residual relatif memenuhi asumsi normalitas." if p >= alpha else "Ada indikasi penyimpangan dari normalitas; pertimbangkan transformasi, uji robust, atau uji nonparametrik."
        elif "levene" in title_low:
            meaning = "Asumsi homogenitas varians relatif terpenuhi." if p >= alpha else "Varians antar kelompok tidak homogen; gunakan Welch/opsi equal_var=False atau uji robust."
        elif "bartlett" in title_low or "kmo" in title_low:
            meaning = "Matriks korelasi layak dianalisis faktor bila KMO memadai dan Bartlett signifikan." if p < alpha else "Bartlett tidak signifikan; struktur korelasi mungkin belum cukup kuat untuk analisis faktor."
        elif "chi-square" in title_low or "crosstab" in title_low:
            meaning = "Ada hubungan/asosiasi antar variabel kategorik." if p < alpha else "Belum ada bukti hubungan yang cukup kuat antar variabel kategorik."
        elif "correlation" in title_low or "korelasi" in title_low:
            meaning = "Ada hubungan linear/monoton yang signifikan pada pasangan variabel terkait." if p < alpha else "Hubungan antar variabel belum signifikan pada taraf α yang dipilih."
        elif "regresi" in title_low or "regression" in title_low or "coefficients" in title_low:
            meaning = "Setidaknya satu parameter/prediktor menunjukkan kontribusi signifikan terhadap model." if p < alpha else "Kontribusi parameter/prediktor belum cukup kuat secara statistik."
        elif "anova" in title_low:
            meaning = "Terdapat perbedaan rerata antar kelompok/faktor; lanjutkan post-hoc atau simple effects untuk mengetahui sumber perbedaan." if p < alpha else "Belum ada bukti perbedaan rerata yang cukup kuat antar kelompok/faktor."
        elif "t-test" in title_low or "mann-whitney" in title_low or "wilcoxon" in title_low or "kruskal" in title_low or "friedman" in title_low:
            meaning = "Perbedaan yang diuji bermakna secara statistik." if p < alpha else "Perbedaan yang diuji belum bermakna secara statistik."
        else:
            meaning = "Hasil signifikan secara statistik." if p < alpha else "Hasil belum signifikan secara statistik."
        insights.append({"Aspek": "Makna statistik", "Insight": f"p = {p_value_text(p)} pada α = {alpha:.2f}. {meaning} {_p_strength(p, alpha)}"})

    # Effect sizes and model fit
    for col, kind, label in [
        ("Cohen's d", "d", "Cohen's d"),
        ("Cohen's dz", "d", "Cohen's dz"),
        ("η²", "eta", "η²"),
        ("ω²", "eta", "ω²"),
        ("R²", "r2", "R²"),
        ("Adjusted R²", "r2", "Adjusted R²"),
        ("Pseudo R²", "r2", "Pseudo R²"),
    ]:
        val = _first_numeric_value(df_table, [col])
        if val is not None:
            if kind == "r2":
                insights.append({"Aspek": "Kekuatan model", "Insight": f"{label} = {_format_number(val)}; model {_effect_size_label(val, kind)}. Nilai ini membantu menilai manfaat praktis model, bukan hanya signifikansinya."})
            else:
                insights.append({"Aspek": "Effect size", "Insight": f"{label} = {_format_number(val)}; besarnya efek tergolong {_effect_size_label(val, kind)}. Ini menunjukkan seberapa berarti temuan secara praktis/substantif."})
            break

    # Reliability
    alpha_val = _first_numeric_value(df_table, ["Cronbach's Alpha"])
    if alpha_val is not None:
        if alpha_val >= 0.90:
            level = "sangat tinggi; cek kemungkinan item terlalu repetitif"
        elif alpha_val >= 0.80:
            level = "baik"
        elif alpha_val >= 0.70:
            level = "dapat diterima untuk riset umum"
        elif alpha_val >= 0.60:
            level = "cukup untuk eksplorasi, tetapi perlu hati-hati"
        else:
            level = "rendah; skala perlu direvisi atau item perlu diperiksa"
        insights.append({"Aspek": "Makna instrumen", "Insight": f"Cronbach's α = {_format_number(alpha_val)} ({level}). Artinya konsistensi internal skala perlu dimaknai sebelum skor total dipakai dalam analisis lanjut."})

    # KMO/Bartlett
    kmo_val = _first_numeric_value(df_table, ["KMO"])
    if kmo_val is not None:
        if kmo_val >= 0.90:
            kmo_desc = "sangat baik"
        elif kmo_val >= 0.80:
            kmo_desc = "baik"
        elif kmo_val >= 0.70:
            kmo_desc = "cukup baik"
        elif kmo_val >= 0.60:
            kmo_desc = "cukup/minimal"
        elif kmo_val >= 0.50:
            kmo_desc = "lemah tetapi masih mungkin untuk eksplorasi"
        else:
            kmo_desc = "tidak memadai"
        insights.append({"Aspek": "Kelayakan faktor", "Insight": f"KMO = {_format_number(kmo_val)} ({kmo_desc}). EFA lebih layak bila KMO ≥ 0.60 dan Bartlett signifikan."})

    # Correlation matrix: strongest pair
    if ("korelasi" in title_low or "correlation" in title_low) and df_table.shape[0] >= 2:
        try:
            corr_df = df_table.copy()
            if corr_df.columns[0].lower() in ["index", "variabel", "variable"]:
                corr_df = corr_df.set_index(corr_df.columns[0])
            corr_num = corr_df.apply(pd.to_numeric, errors="coerce")
            pairs = []
            cols = corr_num.columns.tolist()
            for i, c1 in enumerate(cols):
                for j, c2 in enumerate(cols):
                    if j <= i:
                        continue
                    val = corr_num.loc[c1, c2] if c1 in corr_num.index else corr_num.iloc[i, j]
                    if pd.notna(val):
                        pairs.append((abs(float(val)), float(val), str(c1), str(c2)))
            if pairs:
                _, val, c1, c2 = sorted(pairs, reverse=True)[0]
                direction = "positif" if val > 0 else "negatif"
                insights.append({"Aspek": "Pola hubungan terkuat", "Insight": f"Pasangan terkuat adalah {c1} dan {c2} dengan r = {_format_number(val)} ({direction}, {_effect_size_label(val, 'r')}). Ini dapat menjadi kandidat hubungan utama untuk dibahas dalam riset."})
        except Exception:
            pass

    # Regression coefficients: significant predictors and direction
    if "coefficients" in title_low or "koefisien" in title_low:
        p_col = _first_existing_column(df_table, ["p-value", "P>|t|", "P>|z|", "p"])
        term_col = _first_existing_column(df_table, ["Variabel", "Variable", "Term", "index"])
        coef_col = _first_existing_column(df_table, ["Coef.", "Coefficient", "Koefisien", "coef"])
        if p_col is not None and term_col is not None:
            temp = df_table.copy()
            temp["__p"] = pd.to_numeric(temp[p_col], errors="coerce")
            sig = temp[(temp["__p"] < alpha) & (~temp[term_col].astype(str).str.lower().isin(["const", "intercept"]))]
            if not sig.empty:
                parts = []
                for _, row in sig.head(5).iterrows():
                    direction = ""
                    if coef_col is not None:
                        coef = pd.to_numeric(pd.Series([row[coef_col]]), errors="coerce").iloc[0]
                        if pd.notna(coef):
                            direction = "positif" if coef > 0 else "negatif"
                    parts.append(f"{row[term_col]}{f' ({direction})' if direction else ''}")
                insights.append({"Aspek": "Prediktor penting", "Insight": "Prediktor yang tampak berkontribusi signifikan: " + ", ".join(parts) + ". Fokuskan pembahasan pada arah dan makna teoritis prediktor tersebut."})
            else:
                insights.append({"Aspek": "Prediktor penting", "Insight": "Tidak ada prediktor non-konstanta yang signifikan pada α yang dipilih. Evaluasi ukuran sampel, kualitas pengukuran, multikolinearitas, atau teori model."})

    # EFA loadings
    if "factor loadings" in title_low or "efa" in title_low and "loadings" in title_low:
        try:
            var_col = _first_existing_column(df_table, ["Variable", "Variabel"])
            factor_cols = [c for c in df_table.columns if str(c).lower().startswith("factor")]
            if var_col is not None and factor_cols:
                assignments = []
                cross = []
                for _, row in df_table.iterrows():
                    values = pd.to_numeric(row[factor_cols], errors="coerce").abs()
                    if values.dropna().empty:
                        continue
                    best_factor = values.idxmax()
                    best_val = values.max()
                    if best_val >= 0.40:
                        assignments.append(f"{row[var_col]} → {best_factor} ({best_val:.2f})")
                    if (values >= 0.40).sum() > 1:
                        cross.append(str(row[var_col]))
                if assignments:
                    insights.append({"Aspek": "Struktur faktor", "Insight": "Item/variabel dengan loading kuat: " + "; ".join(assignments[:8]) + ". Gunakan pola ini untuk menamai faktor berdasarkan kesamaan konsep."})
                if cross:
                    insights.append({"Aspek": "Catatan struktur", "Insight": "Ada indikasi cross-loading pada: " + ", ".join(cross[:8]) + ". Item ini perlu ditinjau apakah ambigu secara konseptual."})
        except Exception:
            pass

    # Communalities
    if "communalities" in title_low:
        comm_col = _first_existing_column(df_table, ["Communality", "Komunalitas"])
        var_col = _first_existing_column(df_table, ["Variable", "Variabel"])
        if comm_col is not None and var_col is not None:
            temp = df_table.copy()
            temp["__comm"] = pd.to_numeric(temp[comm_col], errors="coerce")
            low = temp[temp["__comm"] < 0.30]
            if low.empty:
                insights.append({"Aspek": "Kualitas item", "Insight": "Communality mayoritas memadai. Variabel relatif terwakili oleh faktor yang diekstraksi."})
            else:
                insights.append({"Aspek": "Kualitas item", "Insight": "Variabel dengan communality rendah (<0.30): " + ", ".join(low[var_col].astype(str).head(8)) + ". Pertimbangkan revisi/penghapusan item setelah melihat teori."})

    # Descriptive overview
    if "deskriptif" in title_low or "descriptive" in title_low:
        mean_col = _first_existing_column(df_table, ["Mean"])
        sd_col = _first_existing_column(df_table, ["Std. Dev", "SD"])
        var_col = _first_existing_column(df_table, ["Variabel", "Variable"])
        miss_col = _first_existing_column(df_table, ["Missing", "Missing %"])
        if mean_col is not None and var_col is not None:
            insights.append({"Aspek": "Gambaran data", "Insight": "Statistik deskriptif menunjukkan posisi tengah dan sebaran tiap variabel. Variabel dengan rerata tinggi/rendah atau standar deviasi besar layak menjadi fokus awal pembahasan."})
        if miss_col is not None:
            try:
                miss = pd.to_numeric(df_table[miss_col], errors="coerce").fillna(0)
                if (miss > 0).any():
                    worst_idx = miss.idxmax()
                    insights.append({"Aspek": "Kualitas data", "Insight": f"Terdapat missing value, paling menonjol pada {df_table.loc[worst_idx, var_col] if var_col else 'salah satu variabel'}. Jelaskan strategi penanganan missing sebelum inferensi."})
            except Exception:
                pass

    if note and not any(note in x["Insight"] for x in insights):
        insights.append({"Aspek": "Catatan output", "Insight": note})
    if not insights:
        insights.append({"Aspek": "Ringkasan", "Insight": f"Output {title} perlu dibaca bersama desain penelitian, skala pengukuran, ukuran sampel, dan teori. Cuplikan tabel: {_table_as_long_text(df_table)}"})
    insights.append({"Aspek": "Saran pelaporan", "Insight": "Dalam laporan, jangan hanya menulis p-value. Sertakan arah temuan, ukuran efek/model fit, asumsi yang diuji, dan makna substantif terhadap pertanyaan penelitian."})
    return insights


def build_insight_table(report_items, alpha=None, selected_titles=None):
    rows = []
    alpha = alpha if alpha is not None else st.session_state.get("active_alpha", 0.05)
    for item in report_items:
        if selected_titles and item.get("title") not in selected_titles:
            continue
        item_insights = insight_for_output_item(item, alpha=alpha)
        for insight in item_insights:
            rows.append({"Output": item.get("title", "Output"), "Aspek": insight.get("Aspek", "Insight"), "Insight Riset": insight.get("Insight", "")})
    return pd.DataFrame(rows)


def build_research_synthesis(report_items, alpha=None):
    alpha = alpha if alpha is not None else st.session_state.get("active_alpha", 0.05)
    if not report_items:
        return pd.DataFrame([{"Bagian": "Belum ada output", "Sintesis": "Jalankan analisis terlebih dahulu agar aplikasi dapat menyusun insight riset."}])

    sig_outputs, nonsig_outputs, assumption_flags, measurement_flags = [], [], [], []
    for item in report_items:
        title = str(item.get("title", "Output"))
        table = item.get("table")
        if table is None:
            continue
        try:
            df_table = table if isinstance(table, pd.DataFrame) else pd.DataFrame(table)
        except Exception:
            continue
        p = _first_numeric_value(df_table, ["p-value", "p", "Prob(F)", "LLR p-value", "Bartlett p-value"])
        if p is not None:
            title_low = title.lower()
            if any(k in title_low for k in ["normal", "jarque", "levene", "breusch", "diagnostic"]):
                if ("levene" in title_low and p < alpha) or ("normal" in title_low and p < alpha) or ("breusch" in title_low and p < alpha) or ("diagnostic" in title_low and p < alpha):
                    assumption_flags.append(f"{title} (p={p_value_text(p)})")
            elif p < alpha:
                sig_outputs.append(f"{title} (p={p_value_text(p)})")
            else:
                nonsig_outputs.append(f"{title} (p={p_value_text(p)})")
        a = _first_numeric_value(df_table, ["Cronbach's Alpha"])
        if a is not None and a < 0.70:
            measurement_flags.append(f"Reliabilitas perlu perhatian: {title} (α={a:.3f})")
        kmo = _first_numeric_value(df_table, ["KMO"])
        if kmo is not None and kmo < 0.60:
            measurement_flags.append(f"KMO rendah pada {title} (KMO={kmo:.3f})")

    rows = []
    if sig_outputs:
        rows.append({"Bagian": "Temuan utama", "Sintesis": "Ada hasil yang mendukung adanya pola/perbedaan/hubungan: " + "; ".join(sig_outputs[:8]) + ". Temuan ini dapat menjadi inti pembahasan riset, tetapi tetap perlu dibaca bersama effect size dan teori."})
    else:
        rows.append({"Bagian": "Temuan utama", "Sintesis": "Belum terlihat hasil inferensial yang signifikan pada output tersimpan. Ini tidak otomatis berarti 'tidak ada pengaruh/perbedaan'; bisa jadi efek kecil, sampel kurang besar, instrumen kurang sensitif, atau model belum tepat."})
    if nonsig_outputs:
        rows.append({"Bagian": "Temuan non-signifikan", "Sintesis": "Output non-signifikan: " + "; ".join(nonsig_outputs[:8]) + ". Bahas sebagai keterbatasan bukti, bukan sebagai bukti mutlak bahwa fenomena tidak ada."})
    if assumption_flags:
        rows.append({"Bagian": "Asumsi & validitas analisis", "Sintesis": "Ada potensi masalah asumsi: " + "; ".join(assumption_flags[:8]) + ". Pertimbangkan transformasi, metode robust, nonparametrik, atau pelaporan caveat."})
    else:
        rows.append({"Bagian": "Asumsi & validitas analisis", "Sintesis": "Tidak ada bendera asumsi besar yang terdeteksi dari output tersimpan, atau uji asumsi belum dijalankan. Untuk laporan final, tetap dokumentasikan normalitas, homogenitas, outlier, dan multikolinearitas bila relevan."})
    if measurement_flags:
        rows.append({"Bagian": "Kualitas pengukuran", "Sintesis": "Catatan instrumen: " + "; ".join(measurement_flags[:8]) + ". Perbaiki kualitas item/skala sebelum menarik kesimpulan substantif yang kuat."})
    rows.append({"Bagian": "Narasi riset", "Sintesis": "Susun pembahasan dari pertanyaan riset → hasil deskriptif → uji asumsi → hasil inferensial/model → effect size → implikasi teoritis/praktis → keterbatasan."})
    rows.append({"Bagian": "Rekomendasi lanjutan", "Sintesis": "Jika hasil signifikan, lanjutkan post-hoc/diagnostik dan visualisasi. Jika tidak signifikan, cek power/ukuran sampel, kualitas variabel, outlier, reliabilitas, serta kesesuaian model dengan teori."})
    return pd.DataFrame(rows)


def df_to_markdown_safe(df: pd.DataFrame) -> str:
    """Render DataFrame ke Markdown tanpa wajib dependency optional `tabulate`."""
    try:
        return df.to_markdown(index=False)
    except ImportError:
        # Fallback manual agar tombol ekspor Markdown tetap jalan walau user belum install tabulate.
        safe_df = df.copy()
        safe_df.columns = [str(c) for c in safe_df.columns]
        safe_df = safe_df.astype(object).where(pd.notna(safe_df), "")

        def fmt(value):
            text = str(value)
            return text.replace("\n", "<br>").replace("|", "\\|")

        headers = [fmt(c) for c in safe_df.columns]
        rows = safe_df.values.tolist()
        lines = ["| " + " | ".join(headers) + " |"]
        lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
        for row in rows:
            lines.append("| " + " | ".join(fmt(v) for v in row) + " |")
        return "\n".join(lines)



def calculate_kmo_bartlett_safe(data: pd.DataFrame):
    """Hitung KMO dan Bartlett dengan fallback manual bila factor-analyzer bermasalah."""
    x = data.apply(pd.to_numeric, errors="coerce").dropna()
    if x.shape[0] < 3 or x.shape[1] < 2:
        return np.nan, np.nan, np.nan

    # Coba fungsi resmi factor-analyzer lebih dulu.
    if calculate_bartlett_sphericity is not None and calculate_kmo is not None:
        try:
            patch_sklearn_check_array_for_factor_analyzer()
            chi_square_value, bartlett_p = calculate_bartlett_sphericity(x)
            _, kmo_model = calculate_kmo(x)
            return float(kmo_model), float(chi_square_value), float(bartlett_p)
        except Exception:
            pass

    # Fallback manual: berbasis correlation matrix.
    arr = x.to_numpy(dtype=float)
    n, p = arr.shape
    corr = np.corrcoef(arr, rowvar=False)
    corr = np.nan_to_num(corr, nan=0.0, posinf=0.0, neginf=0.0)
    np.fill_diagonal(corr, 1.0)

    det_corr = float(np.linalg.det(corr))
    det_corr = max(det_corr, np.finfo(float).eps)
    bartlett_chi = -(n - 1 - ((2 * p + 5) / 6)) * np.log(det_corr)
    bartlett_df = p * (p - 1) / 2
    bartlett_p = stats.chi2.sf(bartlett_chi, bartlett_df)

    inv_corr = np.linalg.pinv(corr)
    denom = np.sqrt(np.outer(np.diag(inv_corr), np.diag(inv_corr)))
    partial = -inv_corr / denom
    partial = np.nan_to_num(partial, nan=0.0, posinf=0.0, neginf=0.0)
    np.fill_diagonal(corr, 0.0)
    np.fill_diagonal(partial, 0.0)
    corr_sq = corr ** 2
    partial_sq = partial ** 2
    kmo_model = corr_sq.sum() / (corr_sq.sum() + partial_sq.sum()) if (corr_sq.sum() + partial_sq.sum()) else np.nan
    return float(kmo_model), float(bartlett_chi), float(bartlett_p)


def varimax_rotation(loadings, gamma=1.0, q=50, tol=1e-6):
    """Rotasi varimax ringan untuk fallback EFA/PCA."""
    loadings = np.asarray(loadings, dtype=float)
    p, k = loadings.shape
    rotation_matrix = np.eye(k)
    previous = 0
    for _ in range(q):
        rotated = loadings @ rotation_matrix
        u, s, vh = np.linalg.svd(
            loadings.T @ (rotated ** 3 - (gamma / p) * rotated @ np.diag(np.diag(rotated.T @ rotated)))
        )
        rotation_matrix = u @ vh
        current = s.sum()
        if previous != 0 and current < previous * (1 + tol):
            break
        previous = current
    return loadings @ rotation_matrix


def efa_fallback_from_correlation(data: pd.DataFrame, columns, n_factors: int, rotation):
    """EFA fallback stabil tanpa factor-analyzer.

    Engine ini memakai Principal Axis Factoring (PAF) berbasis correlation matrix:
    1) standardisasi data,
    2) estimasi communality awal memakai squared multiple correlation,
    3) iterasi reduced correlation matrix,
    4) rotasi varimax opsional.

    Tujuannya agar fitur EFA tetap usable di komputer yang mengalami konflik
    factor-analyzer/scikit-learn, tanpa hanya jatuh ke PCA biasa.
    """
    if StandardScaler is None:
        raise RuntimeError("scikit-learn belum tersedia untuk EFA fallback.")

    x = data[columns].apply(pd.to_numeric, errors="coerce").dropna()
    if x.shape[0] < 5:
        raise RuntimeError("EFA sebaiknya memiliki minimal 5 baris lengkap; idealnya jauh lebih besar.")
    if x.shape[1] < 2:
        raise RuntimeError("EFA membutuhkan minimal 2 variabel numerik.")

    max_components = min(int(n_factors), x.shape[1] - 1, x.shape[0] - 1)
    if max_components < 1:
        raise RuntimeError("Jumlah faktor tidak valid untuk data yang dipilih.")

    z = StandardScaler().fit_transform(x)
    corr = np.corrcoef(z, rowvar=False)
    corr = np.nan_to_num(corr, nan=0.0, posinf=0.0, neginf=0.0)
    np.fill_diagonal(corr, 1.0)

    # Communality awal: squared multiple correlations (SMC).
    inv_corr = np.linalg.pinv(corr)
    smc = 1 - (1 / np.maximum(np.diag(inv_corr), np.finfo(float).eps))
    communalities = np.clip(np.nan_to_num(smc, nan=0.5, posinf=0.9, neginf=0.2), 0.05, 0.99)

    eigenvalues = None
    loadings_arr = None
    converged = False
    for iteration in range(1, 101):
        reduced_corr = corr.copy()
        np.fill_diagonal(reduced_corr, communalities)
        vals, vecs = np.linalg.eigh(reduced_corr)
        order = np.argsort(vals)[::-1]
        vals = vals[order]
        vecs = vecs[:, order]
        vals_pos = np.maximum(vals[:max_components], 0)
        candidate_loadings = vecs[:, :max_components] * np.sqrt(vals_pos)
        new_communalities = np.clip(np.sum(candidate_loadings ** 2, axis=1), 0.0, 0.999)
        if np.max(np.abs(new_communalities - communalities)) < 1e-5:
            loadings_arr = candidate_loadings
            eigenvalues = vals
            communalities = new_communalities
            converged = True
            break
        loadings_arr = candidate_loadings
        eigenvalues = vals
        communalities = new_communalities

    rotation_requested = "none" if rotation is None else str(rotation).lower()
    rotation_used = "none"
    if rotation_requested in ["varimax", "promax", "oblimin"] and max_components > 1:
        loadings_arr = varimax_rotation(loadings_arr)
        rotation_used = "varimax"
    elif rotation_requested not in ["none", "null"]:
        rotation_used = "none"

    factor_names = [f"Factor {i+1}" for i in range(max_components)]
    loadings = pd.DataFrame(loadings_arr, index=columns, columns=factor_names).reset_index(names="Variable").round(5)
    ss_loadings = np.sum(loadings_arr ** 2, axis=0)
    prop_var = ss_loadings / len(columns)
    variance = pd.DataFrame(
        [ss_loadings, prop_var, np.cumsum(prop_var)],
        index=["SS Loadings", "Proportion Var", "Cumulative Var"],
        columns=factor_names,
    ).reset_index(names="Metric").round(5)
    communalities_table = pd.DataFrame(
        {"Variable": columns, "Communality": np.sum(loadings_arr ** 2, axis=1)}
    ).round(5)

    eigen_table = pd.DataFrame({
        "Component": [f"Component {i+1}" for i in range(len(eigenvalues))],
        "Eigenvalue": eigenvalues,
    }).round(5)

    note = (
        "EFA berhasil dihitung menggunakan engine fallback stabil. "
        "Extraction: Principal Axis Factoring berbasis correlation matrix. "
        f"Rotasi diminta: {rotation_requested}; rotasi dipakai: {rotation_used}. "
        f"Konvergen: {'ya' if converged else 'tidak penuh, hasil iterasi terakhir dipakai'}."
    )
    return loadings, variance, communalities_table, eigen_table, note

def run_efa_analysis(data: pd.DataFrame, columns, n_factors: int, rotation, prefer_fallback=False):
    """Jalankan EFA robust: factor-analyzer jika bisa, fallback jika gagal."""
    x = data[columns].apply(pd.to_numeric, errors="coerce").dropna()
    if x.shape[0] < 5:
        raise RuntimeError("EFA sebaiknya memiliki minimal 5 baris lengkap; idealnya jauh lebih besar.")

    kmo_model, bartlett_chi, bartlett_p = calculate_kmo_bartlett_safe(x)
    kmo_table = pd.DataFrame([
        {"KMO": kmo_model, "Bartlett Chi-square": bartlett_chi, "Bartlett p-value": bartlett_p}
    ]).round(5)

    if not prefer_fallback and FactorAnalyzer is not None:
        try:
            patch_sklearn_check_array_for_factor_analyzer()
            fa = FactorAnalyzer(n_factors=n_factors, rotation=rotation)
            fa.fit(x)
            factor_names = [f"Factor {i+1}" for i in range(n_factors)]
            loadings = pd.DataFrame(fa.loadings_, index=columns, columns=factor_names).reset_index(names="Variable").round(5)
            variance = pd.DataFrame(
                fa.get_factor_variance(),
                index=["SS Loadings", "Proportion Var", "Cumulative Var"],
                columns=factor_names,
            ).reset_index(names="Metric").round(5)
            communalities = pd.DataFrame({"Variable": columns, "Communality": fa.get_communalities()}).round(5)
            eigen_table = pd.DataFrame({"Component": [f"Factor {i+1}" for i in range(n_factors)], "Eigenvalue/SS Loading": fa.get_factor_variance()[0]}).round(5)
            return kmo_table, loadings, variance, communalities, eigen_table, "EFA dihitung dengan factor-analyzer."
        except TypeError as exc:
            msg = str(exc)
            if "force_all_finite" not in msg and "ensure_all_finite" not in msg:
                raise
        except Exception as exc:
            # Jika penyebabnya dependency, fallback; selain itu fallback juga aman untuk menjaga UI tidak crash.
            fallback_reason = str(exc)

    loadings, variance, communalities, eigen_table, note = efa_fallback_from_correlation(x, columns, n_factors, rotation)
    return kmo_table, loadings, variance, communalities, eigen_table, note

def report_as_markdown(report_items):
    lines = [f"# Output Statistik Pro+", "", f"Dibuat: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ""]
    for i, item in enumerate(report_items, start=1):
        lines += [f"## {i}. {item['title']}", f"Waktu: {item.get('created_at', '')}", ""]
        if item.get("note"):
            lines += [item["note"], ""]
        if "table" in item:
            lines += [df_to_markdown_safe(item["table"]), ""]
    return "\n".join(lines)


def report_as_html(report_items):
    body = ["<html><head><meta charset='utf-8'><title>Output Statistik Pro+</title></head><body>"]
    body.append(f"<h1>Output Statistik Pro+</h1><p>Dibuat: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>")
    for i, item in enumerate(report_items, start=1):
        body.append(f"<h2>{i}. {item['title']}</h2><p><em>{item.get('created_at','')}</em></p>")
        if item.get("note"):
            body.append(f"<p>{item['note']}</p>")
        if "table" in item:
            body.append(item["table"].to_html(index=False, border=1))
    body.append("</body></html>")
    return "\n".join(body)


def report_as_docx(report_items):
    if Document is None:
        return None
    bio = io.BytesIO()
    doc = Document()
    doc.add_heading("Output Statistik Pro+", 0)
    doc.add_paragraph(f"Dibuat: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    for i, item in enumerate(report_items, start=1):
        doc.add_heading(f"{i}. {item['title']}", level=1)
        doc.add_paragraph(f"Waktu: {item.get('created_at', '')}")
        if item.get("note"):
            doc.add_paragraph(item["note"])
        if "table" in item:
            tdf = item["table"].fillna("")
            table = doc.add_table(rows=1, cols=len(tdf.columns))
            hdr = table.rows[0].cells
            for j, col in enumerate(tdf.columns):
                hdr[j].text = str(col)
            for _, row in tdf.iterrows():
                cells = table.add_row().cells
                for j, val in enumerate(row):
                    cells[j].text = str(val)
    doc.save(bio)
    return bio.getvalue()


def safe_numeric(series):
    return pd.to_numeric(series, errors="coerce")


def add_report(title, table=None, note=""):
    item = {"title": title, "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"), "note": note}
    if table is not None:
        item["table"] = table.copy() if isinstance(table, pd.DataFrame) else pd.DataFrame(table)
    st.session_state.report_items.append(item)


def show_table(title, table, note="", save=True):
    st.markdown(f"### {title}")
    st.dataframe(table, use_container_width=True)
    if not note:
        note = auto_interpretation(title, table)
    if note:
        st.caption(note)
    if save:
        add_report(title, table, note)


def get_excel_download(df, report_items):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        if df is not None:
            df.to_excel(writer, sheet_name="Data", index=False)
        if report_items:
            row = 0
            sheet_name = "Output"
            pd.DataFrame({"Generated": [datetime.now().strftime("%Y-%m-%d %H:%M:%S")]}).to_excel(
                writer, sheet_name=sheet_name, index=False, startrow=row
            )
            row += 3
            for item in report_items:
                pd.DataFrame([[item["title"]], [item.get("note", "")]]).to_excel(
                    writer, sheet_name=sheet_name, index=False, header=False, startrow=row
                )
                row += 3
                if "table" in item:
                    item["table"].to_excel(writer, sheet_name=sheet_name, index=False, startrow=row)
                    row += len(item["table"]) + 3
                else:
                    row += 1
    return output.getvalue()


def load_data(uploaded_file, csv_sep=",", decimal=".", sheet_name=None):
    suffix = uploaded_file.name.split(".")[-1].lower()
    raw = uploaded_file.getvalue()
    if suffix == "csv":
        return pd.read_csv(io.BytesIO(raw), sep=csv_sep, decimal=decimal)
    if suffix in ["xlsx", "xls"]:
        return pd.read_excel(io.BytesIO(raw), sheet_name=sheet_name)
    if suffix == "sav":
        if pyreadstat is None:
            raise RuntimeError("File .sav memerlukan dependency pyreadstat. Jalankan: pip install pyreadstat")
        with tempfile.NamedTemporaryFile(suffix=".sav", delete=True) as tmp:
            tmp.write(raw)
            tmp.flush()
            df, meta = pyreadstat.read_sav(tmp.name, apply_value_formats=True)
        return df
    raise ValueError("Format belum didukung. Gunakan CSV, XLSX/XLS, atau SAV.")


def sample_dataframe():
    rng = np.random.default_rng(42)
    n = 90
    group = np.repeat(["A", "B", "C"], n // 3)
    gender = rng.choice(["Pria", "Wanita"], size=n)
    anxiety = np.round(rng.normal(62, 10, n) + np.where(group == "C", 7, 0), 1)
    motivation = np.round(rng.normal(70, 8, n) - np.where(group == "A", 4, 0), 1)
    score = np.round(50 + 0.35 * motivation - 0.18 * anxiety + rng.normal(0, 6, n), 1)
    item_1 = np.clip(np.round(rng.normal(3.4, 0.9, n)), 1, 5)
    item_2 = np.clip(np.round(item_1 + rng.normal(0, 0.8, n)), 1, 5)
    item_3 = np.clip(np.round(item_1 + rng.normal(0, 0.9, n)), 1, 5)
    item_4 = np.clip(np.round(item_1 + rng.normal(0, 1.0, n)), 1, 5)
    passed = np.where(score >= np.nanmedian(score), 1, 0)
    return pd.DataFrame(
        {
            "kelompok": group,
            "gender": gender,
            "kecemasan": anxiety,
            "motivasi": motivation,
            "nilai_akhir": score,
            "lulus": passed,
            "item_1": item_1.astype(int),
            "item_2": item_2.astype(int),
            "item_3": item_3.astype(int),
            "item_4": item_4.astype(int),
        }
    )



def interpret_skewness_value(value):
    """Interpretasi ringkas kemencengan distribusi untuk user awam."""
    if pd.isna(value):
        return "Tidak cukup data"
    av = abs(float(value))
    direction = "ke kanan/positif" if value > 0 else "ke kiri/negatif" if value < 0 else "simetris"
    if av < 0.5:
        level = "hampir simetris"
    elif av < 1:
        level = "sedikit menceng"
    elif av < 2:
        level = "cukup menceng"
    else:
        level = "sangat menceng"
    if value == 0:
        return "Distribusi sangat simetris"
    return f"{level}; arah {direction}"


def interpret_kurtosis_value(value):
    """Interpretasi excess kurtosis. Normal biasanya mendekati 0."""
    if pd.isna(value):
        return "Tidak cukup data"
    av = abs(float(value))
    if av < 0.5:
        return "Keruncingan/ekor mirip normal"
    if value > 0:
        if value < 2:
            return "Lebih runcing/ekor agak berat; cek outlier ringan"
        return "Ekor berat/nilai ekstrem kuat; cek outlier dan robust/nonparametrik"
    if value > -2:
        return "Lebih datar dari normal; data relatif menyebar"
    return "Sangat datar/menyebar; cek skala dan pola kategori"


def distribution_shape_recommendation(skewness, kurtosis):
    """Saran praktis berdasarkan skewness dan kurtosis."""
    if pd.isna(skewness) or pd.isna(kurtosis):
        return "Tambah data atau pilih variabel dengan minimal data numerik yang cukup."
    a_skew, a_kurt = abs(float(skewness)), abs(float(kurtosis))
    if a_skew <= 1 and a_kurt <= 1:
        return "Distribusi relatif aman secara deskriptif; lanjut cek normalitas formal/Q-Q plot bila akan memakai uji parametrik."
    if a_skew <= 2 and a_kurt <= 2:
        return "Masih dapat diterima pada banyak riset, tetapi cek histogram, boxplot, dan Q-Q plot sebelum menyimpulkan."
    return "Distribusi jauh dari normal; cek outlier/input data, pertimbangkan transformasi, uji nonparametrik, bootstrap, atau metode robust."


def normality_action_recommendation(p, n=None, alpha=0.05, variable="variabel"):
    """Memberi tindakan praktis saat normalitas terpenuhi/tidak terpenuhi."""
    if pd.isna(p):
        return "Normalitas belum dapat dinilai. Tambah data numerik atau cek missing value."
    if p >= alpha:
        return "Normalitas tidak bermasalah secara formal. Lanjutkan analisis parametrik bila asumsi lain juga terpenuhi."
    small_n = n is not None and n < 30
    if small_n:
        return "Normalitas tidak terpenuhi dan N kecil. Cek outlier/Q-Q plot; pertimbangkan transformasi data atau gunakan uji nonparametrik."
    return "Normalitas tidak terpenuhi. Jika N besar, uji parametrik kadang tetap cukup robust, tetapi tetap cek grafik, outlier, effect size/CI, atau gunakan bootstrap/robust/nonparametrik."


def normality_solution_table():
    """Panduan tindakan ketika uji normalitas tidak terpenuhi."""
    return pd.DataFrame([
        {"Kondisi data/tujuan": "Ada salah input atau outlier ekstrem", "Yang sebaiknya dilakukan": "Cek data mentah, boxplot, z-score/IQR; koreksi jika salah input", "Catatan untuk user awam": "Jangan langsung hapus outlier. Pastikan dulu apakah itu kesalahan input atau data nyata."},
        {"Kondisi data/tujuan": "Data menceng kanan, misalnya pendapatan/waktu", "Yang sebaiknya dilakukan": "Coba transformasi log, sqrt, atau Box-Cox/Yeo-Johnson", "Catatan untuk user awam": "Transformasi mengubah skala, jadi interpretasi hasil harus disesuaikan."},
        {"Kondisi data/tujuan": "Membandingkan 2 kelompok independen", "Yang sebaiknya dilakukan": "Gunakan Welch t-test jika varians tidak homogen; gunakan Mann-Whitney jika distribusi sangat tidak normal/ordinal", "Catatan untuk user awam": "Mann-Whitney membandingkan peringkat/distribusi, bukan selalu rata-rata."},
        {"Kondisi data/tujuan": "Membandingkan pre-post/berpasangan", "Yang sebaiknya dilakukan": "Gunakan Wilcoxon Signed-Rank atau bootstrap CI untuk selisih", "Catatan untuk user awam": "Yang perlu dicek normalitasnya adalah selisih pre-post, bukan masing-masing kolom saja."},
        {"Kondisi data/tujuan": "Membandingkan 3+ kelompok", "Yang sebaiknya dilakukan": "Gunakan Welch ANOVA/Games-Howell atau Kruskal-Wallis + Dunn post-hoc", "Catatan untuk user awam": "Jika hasil signifikan, lanjut post-hoc untuk tahu kelompok mana yang berbeda."},
        {"Kondisi data/tujuan": "Korelasi antar variabel", "Yang sebaiknya dilakukan": "Gunakan Spearman/Kendall jika data tidak normal, ordinal, atau ada outlier", "Catatan untuk user awam": "Spearman cocok untuk hubungan yang meningkat/menurun tetapi tidak harus linear."},
        {"Kondisi data/tujuan": "Regresi", "Yang sebaiknya dilakukan": "Cek normalitas residual, bukan hanya Y; gunakan robust standard errors/transformasi/bootstrapping bila perlu", "Catatan untuk user awam": "Regresi lebih peduli pola residual daripada bentuk data mentah."},
        {"Kondisi data/tujuan": "Sampel besar", "Yang sebaiknya dilakukan": "Jangan hanya bergantung pada p normalitas; lihat Q-Q plot, skewness/kurtosis, residual, dan effect size", "Catatan untuk user awam": "Pada N besar, uji normalitas sangat sensitif dan mudah signifikan."},
        {"Kondisi data/tujuan": "Laporan riset", "Yang sebaiknya dilakukan": "Laporkan hasil uji, grafik, keputusan metode, effect size, dan alasan memilih alternatif", "Catatan untuk user awam": "Tulis keputusan metodologis secara transparan agar pembaca paham alasan analisis."},
    ])


def skewness_kurtosis_reference_table():
    return pd.DataFrame([
        {"Indikator": "Skewness", "Nilai": "-0.5 s.d. +0.5", "Makna": "Distribusi hampir simetris", "Tindakan": "Biasanya aman; tetap cek grafik bila analisis inferensial."},
        {"Indikator": "Skewness", "Nilai": "+0.5 s.d. +1 atau -1 s.d. -0.5", "Makna": "Sedikit menceng", "Tindakan": "Cek histogram/Q-Q plot; lanjut hati-hati."},
        {"Indikator": "Skewness", "Nilai": "+1 s.d. +2 atau -2 s.d. -1", "Makna": "Cukup menceng", "Tindakan": "Cek outlier dan pertimbangkan transformasi/alternatif nonparametrik."},
        {"Indikator": "Skewness", "Nilai": "> +2 atau < -2", "Makna": "Sangat menceng", "Tindakan": "Jangan langsung pakai uji parametrik; evaluasi outlier, transformasi, robust/nonparametrik."},
        {"Indikator": "Kurtosis", "Nilai": "Mendekati 0", "Makna": "Keruncingan/ekor mirip normal", "Tindakan": "Cukup baik secara deskriptif."},
        {"Indikator": "Kurtosis", "Nilai": "> 0", "Makna": "Ekor lebih berat/lebih banyak nilai ekstrem", "Tindakan": "Cek outlier, boxplot, dan sensitivitas hasil."},
        {"Indikator": "Kurtosis", "Nilai": "< 0", "Makna": "Distribusi lebih datar/menyebar", "Tindakan": "Cek apakah variabel berskala kategori/ordinal atau rentang nilai terbatas."},
        {"Indikator": "Patokan praktis", "Nilai": "-2 s.d. +2", "Makna": "Sering masih diterima pada banyak riset terapan", "Tindakan": "Tetap kombinasikan dengan Shapiro/Q-Q plot dan konteks penelitian."},
    ])

def descriptive_table(df, cols):
    rows = []
    for col in cols:
        s = safe_numeric(df[col]).dropna()
        skew = stats.skew(s, nan_policy="omit") if len(s) > 2 else np.nan
        kurt = stats.kurtosis(s, nan_policy="omit") if len(s) > 3 else np.nan
        rows.append(
            {
                "Variabel": col,
                "N": len(s),
                "Missing": int(df[col].isna().sum()),
                "Mean": s.mean() if len(s) else np.nan,
                "Median": s.median() if len(s) else np.nan,
                "Std. Dev": s.std(ddof=1) if len(s) > 1 else np.nan,
                "Min": s.min() if len(s) else np.nan,
                "Max": s.max() if len(s) else np.nan,
                "Skewness": skew,
                "Makna Skewness": interpret_skewness_value(skew),
                "Kurtosis": kurt,
                "Makna Kurtosis": interpret_kurtosis_value(kurt),
                "Saran Distribusi": distribution_shape_recommendation(skew, kurt),
            }
        )
    out = pd.DataFrame(rows)
    num = out.select_dtypes(include=[np.number]).columns
    out[num] = out[num].round(4)
    return out


def frequency_table(df, col):
    vc = df[col].value_counts(dropna=False)
    pct = df[col].value_counts(dropna=False, normalize=True) * 100
    return pd.DataFrame({"Kategori": vc.index.astype(str), "Frekuensi": vc.values, "Persen": pct.values.round(2)})


def normality_table(df, cols, alpha=None):
    if alpha is None:
        alpha = st.session_state.get("active_alpha", 0.05)
    rows = []
    for col in cols:
        s = safe_numeric(df[col]).dropna()
        if len(s) < 3:
            stat, p = np.nan, np.nan
            test = "Shapiro-Wilk perlu N ≥ 3"
        elif len(s) <= 5000:
            stat, p = stats.shapiro(s)
            test = "Shapiro-Wilk"
        else:
            stat, p = stats.normaltest(s)
            test = "D'Agostino K²"
        rows.append({
            "Variabel": col,
            "N": len(s),
            "Uji": test,
            "Statistic": stat,
            "p-value": p,
            "Keputusan": "Normalitas OK" if (not pd.isna(p) and p >= alpha) else "Tidak normal/perlu cek" if not pd.isna(p) else "Belum dapat dinilai",
            "Apa yang sebaiknya dilakukan": normality_action_recommendation(p, len(s), alpha, col),
        })
    out = pd.DataFrame(rows)
    num = out.select_dtypes(include=[np.number]).columns
    out[num] = out[num].round(5)
    return out


def one_sample_ttest(s, mu, alpha):
    s = s.dropna()
    t_stat, p = stats.ttest_1samp(s, mu)
    df_val = len(s) - 1
    mean_diff = s.mean() - mu
    se = stats.sem(s)
    ci = stats.t.interval(1 - alpha, df_val, loc=s.mean(), scale=se)
    d = mean_diff / s.std(ddof=1) if s.std(ddof=1) != 0 else np.nan
    return pd.DataFrame(
        [
            {
                "N": len(s),
                "Mean": s.mean(),
                "μ₀": mu,
                "Mean Difference": mean_diff,
                "t": t_stat,
                "df": df_val,
                "p-value": p,
                f"CI Lower ({(1-alpha)*100:.0f}%)": ci[0],
                f"CI Upper ({(1-alpha)*100:.0f}%)": ci[1],
                "Cohen's d": d,
                "Keputusan": decision_text(p, alpha),
            }
        ]
    ).round(5)


def unequal_n_label(ns):
    """Ringkasan sederhana untuk kondisi jumlah data per grup yang tidak sama."""
    ns = [int(n) for n in ns if pd.notna(n)]
    if not ns or min(ns) <= 0:
        return "Tidak dapat dinilai"
    if len(set(ns)) == 1:
        return "Seimbang"
    ratio = max(ns) / max(1, min(ns))
    if ratio < 1.5:
        return "Unequal ringan"
    if ratio < 3:
        return "Unequal sedang"
    return "Unequal kuat"


def unequal_n_recommendation(ns, analysis="t-test", levene_p=None):
    """Saran praktis saat ukuran sampel/observasi antar grup tidak sama."""
    ns = [int(n) for n in ns if pd.notna(n)]
    if not ns or min(ns) <= 0:
        return "Data belum cukup untuk menilai kesetaraan jumlah observasi."
    if len(set(ns)) == 1:
        return "Jumlah observasi antar grup seimbang. Lanjutkan dengan tetap memeriksa asumsi normalitas dan homogenitas varians."

    ratio = max(ns) / max(1, min(ns))
    homogeneity_bad = levene_p is not None and pd.notna(levene_p) and levene_p < 0.05

    if analysis.lower().startswith("t"):
        if homogeneity_bad or ratio >= 1.5:
            return "Jumlah data antar grup tidak sama. Gunakan Welch t-test/equal_var=False sebagai pilihan utama; Mann-Whitney dapat dipakai bila data sangat tidak normal atau ordinal."
        return "Jumlah data antar grup tidak sama, tetapi perbedaannya ringan. Student t-test masih mungkin jika varians homogen; Welch t-test tetap lebih aman."

    if "anova" in analysis.lower():
        if homogeneity_bad or ratio >= 1.5:
            return "Jumlah data antar grup tidak sama. ANOVA biasa masih bisa menghitung, tetapi Welch ANOVA lebih aman bila varians tidak homogen; gunakan Games-Howell atau Kruskal-Wallis + Dunn untuk alternatif/post-hoc."
        return "Jumlah data antar grup tidak sama secara ringan. ANOVA biasa masih dapat digunakan jika varians homogen dan distribusi wajar; tetap laporkan ukuran grup."

    return "Jumlah observasi tidak sama. Pilih metode yang toleran terhadap unequal N, cek asumsi, dan laporkan N setiap grup."


def group_size_summary(groups, original_lengths=None):
    """Buat tabel N valid, missing, dan status unequal untuk daftar grup."""
    rows = []
    ns = []
    for idx, (label, values) in enumerate(groups.items()):
        ser = pd.Series(values)
        n_valid = int(ser.dropna().shape[0])
        original_n = int(original_lengths.get(label, len(ser))) if isinstance(original_lengths, dict) else int(len(ser))
        ns.append(n_valid)
        rows.append({
            "Grup": str(label),
            "N Valid": n_valid,
            "Missing": max(0, original_n - n_valid),
            "Mean": safe_numeric(ser).dropna().mean() if n_valid else np.nan,
            "SD": safe_numeric(ser).dropna().std(ddof=1) if n_valid > 1 else np.nan,
        })
    status = unequal_n_label(ns)
    ratio = max(ns) / max(1, min(ns)) if ns and min(ns) > 0 else np.nan
    out = pd.DataFrame(rows)
    out["Status N"] = status
    out["Rasio Max/Min N"] = round(ratio, 3) if pd.notna(ratio) else np.nan
    return out.round(5)


def safe_levene(*samples):
    """Levene test defensif untuk homogenitas varians."""
    try:
        cleaned = [pd.Series(s).dropna().astype(float) for s in samples]
        if len(cleaned) < 2 or any(len(x) < 2 for x in cleaned):
            return np.nan
        if any(x.var(ddof=1) == 0 for x in cleaned):
            return np.nan
        return float(stats.levene(*cleaned, center="median").pvalue)
    except Exception:
        return np.nan


def welch_anova_from_groups(groups, alpha=0.05):
    """Welch ANOVA manual untuk grup dengan ukuran/varians tidak sama."""
    cleaned = {str(k): pd.Series(v).dropna().astype(float) for k, v in groups.items()}
    cleaned = {k: v for k, v in cleaned.items() if len(v) >= 2 and v.var(ddof=1) > 0}
    k = len(cleaned)
    if k < 2:
        return pd.DataFrame({"Pesan": ["Welch ANOVA membutuhkan minimal 2 grup dengan N ≥ 2 dan varians > 0."]})
    n = np.array([len(v) for v in cleaned.values()], dtype=float)
    means = np.array([v.mean() for v in cleaned.values()], dtype=float)
    variances = np.array([v.var(ddof=1) for v in cleaned.values()], dtype=float)
    w = n / variances
    w_sum = w.sum()
    mean_w = (w * means).sum() / w_sum
    df1 = k - 1
    numerator = (w * (means - mean_w) ** 2).sum() / df1
    correction_terms = ((1 - (w / w_sum)) ** 2) / (n - 1)
    correction = 1 + (2 * (k - 2) / (k**2 - 1)) * correction_terms.sum() if k > 2 else 1
    f_value = numerator / correction
    df2 = (k**2 - 1) / (3 * correction_terms.sum()) if correction_terms.sum() > 0 else np.nan
    p_value = stats.f.sf(f_value, df1, df2) if pd.notna(df2) else np.nan
    return pd.DataFrame([{
        "Metode": "Welch ANOVA",
        "F": f_value,
        "df1": df1,
        "df2": df2,
        "p-value": p_value,
        "Keputusan": decision_text(p_value, alpha) if pd.notna(p_value) else "Tidak dapat dihitung",
        "Catatan": "Direkomendasikan saat N/varians antar grup tidak sama.",
    }]).round(5)


def games_howell_table(long_df, alpha=0.05):
    """Post-hoc Games-Howell jika pingouin tersedia; fallback pesan ramah jika tidak."""
    if pg is None:
        return pd.DataFrame({"Pesan": ["Games-Howell membutuhkan package pingouin. Jalankan pip install -r requirements.txt."]})
    try:
        work = long_df[["nilai", "grup"]].dropna().copy()
        if work["grup"].nunique() < 2:
            return pd.DataFrame({"Pesan": ["Minimal 2 grup diperlukan."]})
        res = pg.pairwise_gameshowell(dv="nilai", between="grup", data=work)
        if "pval" in res.columns:
            res["Keputusan"] = np.where(res["pval"] < alpha, "Signifikan", "Tidak signifikan")
        return res.round(5)
    except Exception as exc:
        return pd.DataFrame({"Pesan": [f"Games-Howell gagal dihitung: {exc}"]})


def independent_ttest(s1, s2, label1, label2, alpha, equal_var=True):
    s1, s2 = s1.dropna(), s2.dropna()
    if len(s1) < 2 or len(s2) < 2:
        return pd.DataFrame({"Pesan": ["Setiap grup membutuhkan minimal 2 data valid untuk independent t-test."]})
    levene_p = safe_levene(s1, s2)
    t_stat, p = stats.ttest_ind(s1, s2, equal_var=equal_var)
    if equal_var:
        df_val = len(s1) + len(s2) - 2
        pooled_sd = np.sqrt(((len(s1) - 1) * s1.var(ddof=1) + (len(s2) - 1) * s2.var(ddof=1)) / df_val)
        se_diff = pooled_sd * np.sqrt(1 / len(s1) + 1 / len(s2))
        method = "Student t-test"
    else:
        v1, v2 = s1.var(ddof=1), s2.var(ddof=1)
        se_diff = np.sqrt(v1 / len(s1) + v2 / len(s2))
        denom = ((v1 / len(s1)) ** 2 / (len(s1) - 1) + (v2 / len(s2)) ** 2 / (len(s2) - 1))
        df_val = (v1 / len(s1) + v2 / len(s2)) ** 2 / denom if denom != 0 else np.nan
        pooled_sd = np.sqrt((s1.var(ddof=1) + s2.var(ddof=1)) / 2)
        method = "Welch t-test"
    diff = s1.mean() - s2.mean()
    ci = stats.t.interval(1 - alpha, df_val, loc=diff, scale=se_diff) if pd.notna(df_val) else (np.nan, np.nan)
    d = diff / pooled_sd if pooled_sd != 0 else np.nan
    ns = [len(s1), len(s2)]
    return pd.DataFrame(
        [
            {
                "Metode": method,
                "Grup 1": label1,
                "N1": len(s1),
                "Mean1": s1.mean(),
                "Grup 2": label2,
                "N2": len(s2),
                "Mean2": s2.mean(),
                "Status N": unequal_n_label(ns),
                "Levene p": levene_p,
                "Mean Difference": diff,
                "t": t_stat,
                "df": df_val,
                "p-value": p,
                f"CI Lower ({(1-alpha)*100:.0f}%)": ci[0],
                f"CI Upper ({(1-alpha)*100:.0f}%)": ci[1],
                "Cohen's d": d,
                "Keputusan": decision_text(p, alpha),
                "Saran jika N tidak sama": unequal_n_recommendation(ns, "t-test", levene_p),
            }
        ]
    ).round(5)

def paired_ttest(s1, s2, label1, label2, alpha):
    pair = pd.concat([s1, s2], axis=1).dropna()
    a, b = pair.iloc[:, 0], pair.iloc[:, 1]
    t_stat, p = stats.ttest_rel(a, b)
    diff = a - b
    df_val = len(diff) - 1
    se = stats.sem(diff)
    ci = stats.t.interval(1 - alpha, df_val, loc=diff.mean(), scale=se)
    d = diff.mean() / diff.std(ddof=1) if diff.std(ddof=1) != 0 else np.nan
    return pd.DataFrame(
        [
            {
                "Pasangan": f"{label1} - {label2}",
                "N": len(diff),
                "Mean Difference": diff.mean(),
                "t": t_stat,
                "df": df_val,
                "p-value": p,
                f"CI Lower ({(1-alpha)*100:.0f}%)": ci[0],
                f"CI Upper ({(1-alpha)*100:.0f}%)": ci[1],
                "Cohen's dz": d,
                "Keputusan": decision_text(p, alpha),
            }
        ]
    ).round(5)


def anova_wide(df, cols, alpha):
    data = [safe_numeric(df[c]).dropna() for c in cols]
    f, p = stats.f_oneway(*data)
    k = len(cols)
    n_total = sum(len(g) for g in data)
    grand = np.concatenate([g.to_numpy() for g in data]).mean()
    ss_between = sum(len(g) * (g.mean() - grand) ** 2 for g in data)
    ss_within = sum(((g - g.mean()) ** 2).sum() for g in data)
    ss_total = ss_between + ss_within
    df_between = k - 1
    df_within = n_total - k
    ms_between = ss_between / df_between
    ms_within = ss_within / df_within
    eta_sq = ss_between / ss_total if ss_total > 0 else np.nan
    omega_sq = (ss_between - df_between * ms_within) / (ss_total + ms_within) if ss_total > 0 else np.nan
    ns = [len(g) for g in data]
    levene_p = safe_levene(*data)
    table = pd.DataFrame(
        [
            {"Sumber": "Between Groups", "SS": ss_between, "df": df_between, "MS": ms_between, "F": f, "p-value": p},
            {"Sumber": "Within Groups", "SS": ss_within, "df": df_within, "MS": ms_within, "F": np.nan, "p-value": np.nan},
            {"Sumber": "Total", "SS": ss_total, "df": n_total - 1, "MS": np.nan, "F": np.nan, "p-value": np.nan},
        ]
    ).round(5)
    effects = pd.DataFrame([{
        "η²": eta_sq,
        "ω²": omega_sq,
        "Status N": unequal_n_label(ns),
        "Levene p": levene_p,
        "Keputusan": decision_text(p, alpha),
        "Saran jika N tidak sama": unequal_n_recommendation(ns, "anova", levene_p),
    }]).round(5)
    long = pd.concat([pd.DataFrame({"nilai": safe_numeric(df[c]).dropna(), "grup": c}) for c in cols], ignore_index=True)
    return table, effects, long


def anova_long(df, dv, group_col, alpha):
    work = df[[dv, group_col]].copy()
    work[dv] = safe_numeric(work[dv])
    work = work.dropna()
    groups = [g[dv] for _, g in work.groupby(group_col)]
    f, p = stats.f_oneway(*groups)
    labels = work[group_col].astype(str).unique().tolist()
    table, effects, long = anova_wide(
        pd.DataFrame({str(label): work.loc[work[group_col].astype(str) == str(label), dv].reset_index(drop=True) for label in labels}),
        [str(label) for label in labels],
        alpha,
    )
    return table, effects, work.rename(columns={dv: "nilai", group_col: "grup"})


def tukey_table(long_df, alpha):
    if pairwise_tukeyhsd is None:
        return pd.DataFrame({"Pesan": ["statsmodels belum tersedia untuk Tukey HSD."]})
    tk = pairwise_tukeyhsd(endog=long_df["nilai"].astype(float), groups=long_df["grup"].astype(str), alpha=alpha)
    return pd.DataFrame(tk.summary().data[1:], columns=tk.summary().data[0])



# -----------------------------------------------------------------------------
# Research interpretation helpers: Reliability, PCA, and EFA
# -----------------------------------------------------------------------------
def interpret_cronbach_reliability(alpha_value, item_table=None, n_items=None, n_cases=None):
    """Bangun interpretasi riset untuk Cronbach's Alpha dan item-total statistics."""
    rows = []
    try:
        a = float(alpha_value)
    except Exception:
        a = np.nan

    if pd.isna(a):
        rows.append({
            "Aspek": "Status reliabilitas",
            "Makna untuk riset": "Cronbach's Alpha tidak dapat dihitung. Biasanya ini terjadi karena item terlalu sedikit, varians total nol, atau data lengkap terlalu sedikit.",
            "Saran tindakan": "Pastikan minimal 2 item numerik, item memiliki variasi jawaban, dan missing value tidak terlalu banyak.",
        })
        return pd.DataFrame(rows)

    if a >= 0.95:
        level = "sangat tinggi, tetapi perlu dicek"
        meaning = "Item sangat konsisten, namun nilai terlalu tinggi dapat menandakan beberapa item terlalu mirip atau repetitif."
        action = "Pertahankan skala jika sesuai teori, tetapi cek apakah ada item yang redundan agar kuesioner lebih ringkas."
    elif a >= 0.90:
        level = "sangat baik"
        meaning = "Konsistensi internal sangat kuat. Skala umumnya layak dipakai untuk membentuk skor komposit."
        action = "Gunakan skor total/rata-rata jika item memang satu konstruk; tetap cek item-total correlation."
    elif a >= 0.80:
        level = "baik"
        meaning = "Konsistensi internal baik. Item relatif bergerak searah dalam mengukur konstruk yang sama."
        action = "Skala dapat digunakan; periksa item dengan korelasi item-total rendah untuk penyempurnaan."
    elif a >= 0.70:
        level = "dapat diterima"
        meaning = "Konsistensi internal cukup memadai untuk banyak riset sosial/pendidikan/manajemen."
        action = "Skala dapat digunakan dengan catatan; cek item yang melemahkan alpha."
    elif a >= 0.60:
        level = "cukup untuk eksplorasi"
        meaning = "Reliabilitas masih lemah untuk kesimpulan kuat, tetapi kadang dapat diterima pada studi awal/eksploratori."
        action = "Revisi item, tambah jumlah item yang relevan, cek item negatif yang belum di-reverse coding."
    else:
        level = "rendah"
        meaning = "Item belum konsisten mengukur konstruk yang sama. Skor total berisiko tidak stabil."
        action = "Jangan langsung memakai skor total; cek salah input, reverse coding, item tidak relevan, atau pisahkan dimensi dengan EFA/PCA."

    rows.append({
        "Aspek": "Status reliabilitas",
        "Makna untuk riset": f"Cronbach's Alpha = {_format_number(a)}; reliabilitas tergolong {level}. {meaning}",
        "Saran tindakan": action,
    })

    if n_items is not None or n_cases is not None:
        item_txt = f"{n_items} item" if n_items is not None else "jumlah item tidak diketahui"
        case_txt = f"{n_cases} kasus lengkap" if n_cases is not None else "jumlah kasus lengkap tidak diketahui"
        rows.append({
            "Aspek": "Kecukupan data",
            "Makna untuk riset": f"Perhitungan didasarkan pada {item_txt} dan {case_txt}.",
            "Saran tindakan": "Semakin banyak item dan responden yang relevan, semakin stabil estimasi reliabilitas. Untuk kuesioner formal, cek juga validitas isi dan struktur faktor.",
        })

    if item_table is not None and isinstance(item_table, pd.DataFrame) and not item_table.empty:
        tbl = item_table.copy()
        item_col = _first_existing_column(tbl, ["Item", "Variable", "Variabel"])
        corr_col = _first_existing_column(tbl, ["Corrected Item-Total Corr", "Item-Total", "Corrected Item Total Correlation"])
        del_col = _first_existing_column(tbl, ["Alpha if Deleted", "Cronbach's Alpha if Item Deleted"])

        if corr_col is not None:
            tbl["__corr"] = pd.to_numeric(tbl[corr_col], errors="coerce")
            low = tbl[tbl["__corr"] < 0.30]
            negative = tbl[tbl["__corr"] < 0]
            if len(negative):
                items = ", ".join(negative[item_col].astype(str).head(8)) if item_col else f"{len(negative)} item"
                rows.append({
                    "Aspek": "Item bermasalah",
                    "Makna untuk riset": f"Ada item dengan korelasi item-total negatif: {items}. Ini sering menunjukkan item berlawanan arah, belum di-reverse coding, atau tidak mengukur konstruk yang sama.",
                    "Saran tindakan": "Periksa redaksi item, lakukan reverse coding untuk item negatif, atau keluarkan item jika tidak sesuai teori.",
                })
            elif len(low):
                items = ", ".join(low[item_col].astype(str).head(8)) if item_col else f"{len(low)} item"
                rows.append({
                    "Aspek": "Item perlu ditinjau",
                    "Makna untuk riset": f"Ada item dengan corrected item-total correlation < 0.30: {items}. Item ini kurang sejalan dengan skor total skala.",
                    "Saran tindakan": "Tinjau ulang item tersebut. Jika secara teori lemah dan alpha meningkat saat dihapus, pertimbangkan revisi atau penghapusan.",
                })
            else:
                rows.append({
                    "Aspek": "Konsistensi item",
                    "Makna untuk riset": "Tidak ada item dengan corrected item-total correlation di bawah 0.30 yang terdeteksi.",
                    "Saran tindakan": "Item relatif konsisten; keputusan akhir tetap perlu mengikuti teori/konstruk penelitian.",
                })

        if del_col is not None:
            tbl["__alpha_deleted"] = pd.to_numeric(tbl[del_col], errors="coerce")
            improves = tbl[tbl["__alpha_deleted"] > (a + 0.02)]
            if len(improves):
                items = ", ".join(improves[item_col].astype(str).head(8)) if item_col else f"{len(improves)} item"
                rows.append({
                    "Aspek": "Alpha if item deleted",
                    "Makna untuk riset": f"Reliabilitas tampak meningkat jika item berikut dihapus: {items}.",
                    "Saran tindakan": "Jangan hapus otomatis hanya karena angka meningkat. Cocokkan dengan teori, validitas isi, dan arah item sebelum memutuskan.",
                })

    rows.append({
        "Aspek": "Makna substantif",
        "Makna untuk riset": "Reliabilitas menunjukkan konsistensi alat ukur, bukan bukti bahwa konstruk pasti valid. Skala yang reliabel masih perlu didukung validitas teori, isi, dan/atau struktur faktor.",
        "Saran tindakan": "Jika item berasal dari kuesioner, lanjutkan dengan validitas item, EFA/PCA, atau konfirmasi berdasarkan teori sebelum memakai skor komposit sebagai variabel riset utama.",
    })
    return pd.DataFrame(rows)


def interpret_pca_results(explained, loadings, selected_cols=None, n_rows=None):
    """Bangun interpretasi riset untuk PCA: variance, loading, cross-loading, dan rekomendasi."""
    rows = []
    if explained is None or loadings is None or not isinstance(explained, pd.DataFrame) or not isinstance(loadings, pd.DataFrame):
        return pd.DataFrame([{
            "Aspek": "Status PCA",
            "Makna untuk riset": "Output PCA belum tersedia atau tidak dapat dibaca.",
            "Saran tindakan": "Jalankan PCA dengan minimal 2 variabel numerik dan data lengkap yang memadai.",
        }])

    exp = explained.copy()
    load = loadings.copy()
    var_col = _first_existing_column(exp, ["Explained Variance %", "Variance %", "% Variance"])
    cum_col = _first_existing_column(exp, ["Cumulative %", "Cumulative Variance %"])
    eig_col = _first_existing_column(exp, ["Eigenvalue", "Eigenvalue/SS Loading"])
    comp_col = _first_existing_column(exp, ["Component", "Factor"])

    if cum_col is not None:
        vals = pd.to_numeric(exp[cum_col], errors="coerce").dropna()
        if len(vals):
            cum = float(vals.iloc[-1])
            if cum >= 70:
                level = "sangat kuat"
                action = "Komponen yang dipilih sudah merangkum sebagian besar informasi variabel."
            elif cum >= 60:
                level = "baik"
                action = "Cukup baik untuk banyak riset sosial; tetap cek loading tiap variabel."
            elif cum >= 50:
                level = "cukup/eksploratori"
                action = "Masih dapat dipakai untuk eksplorasi, tetapi pertimbangkan menambah komponen atau meninjau variabel."
            else:
                level = "rendah"
                action = "Komponen belum cukup merangkum data; pertimbangkan jumlah komponen lebih banyak, buang variabel lemah, atau gunakan pendekatan lain."
            rows.append({
                "Aspek": "Varians yang dijelaskan",
                "Makna untuk riset": f"Komponen yang dipilih menjelaskan total sekitar {_format_number(cum, 2)}% variasi data; kategori {level}.",
                "Saran tindakan": action,
            })

    if eig_col is not None:
        eig = pd.to_numeric(exp[eig_col], errors="coerce")
        n_gt1 = int((eig > 1).sum())
        if n_gt1:
            rows.append({
                "Aspek": "Jumlah komponen potensial",
                "Makna untuk riset": f"Berdasarkan aturan eigenvalue > 1, ada sekitar {n_gt1} komponen yang layak dipertimbangkan.",
                "Saran tindakan": "Gunakan ini sebagai petunjuk awal saja. Cocokkan dengan scree plot, teori konstruk, dan interpretasi loading.",
            })
        else:
            rows.append({
                "Aspek": "Jumlah komponen potensial",
                "Makna untuk riset": "Tidak ada eigenvalue > 1 pada komponen yang ditampilkan, atau nilai eigen tidak cukup kuat.",
                "Saran tindakan": "Cek kembali jumlah komponen, kualitas variabel, korelasi antar variabel, dan ukuran sampel.",
            })

    variable_col = _first_existing_column(load, ["Variable", "Variabel", "Item"])
    comp_cols = [c for c in load.columns if str(c).upper().startswith("PC") or str(c).lower().startswith("component")]
    if not comp_cols:
        comp_cols = [c for c in load.columns if c != variable_col and pd.api.types.is_numeric_dtype(pd.to_numeric(load[c], errors="coerce"))]

    if variable_col is not None and comp_cols:
        assignments = []
        weak = []
        cross = []
        for _, row in load.iterrows():
            vals = pd.to_numeric(row[comp_cols], errors="coerce").abs().dropna()
            if vals.empty:
                continue
            best_comp = vals.idxmax()
            best_val = float(vals.max())
            strong_count = int((vals >= 0.40).sum())
            var_name = str(row[variable_col])
            assignments.append((var_name, str(best_comp), best_val))
            if best_val < 0.40:
                weak.append(var_name)
            if strong_count >= 2:
                cross.append(var_name)

        if assignments:
            top = sorted(assignments, key=lambda x: x[2], reverse=True)[:8]
            top_txt = "; ".join([f"{v} → {pc} ({val:.2f})" for v, pc, val in top])
            rows.append({
                "Aspek": "Loading utama",
                "Makna untuk riset": f"Variabel dengan kontribusi/loading terbesar: {top_txt}.",
                "Saran tindakan": "Namai komponen berdasarkan variabel-variabel dengan loading tertinggi dan sesuai teori. Loading ≥0.40 biasanya mulai layak dibaca; ≥0.60 kuat.",
            })
        if weak:
            rows.append({
                "Aspek": "Variabel lemah",
                "Makna untuk riset": "Variabel dengan loading tertinggi < 0.40: " + ", ".join(weak[:10]) + ". Variabel ini kurang terwakili oleh komponen yang dipilih.",
                "Saran tindakan": "Pertimbangkan revisi, penghapusan, atau tambah komponen bila variabel tersebut penting secara teori.",
            })
        if cross:
            rows.append({
                "Aspek": "Cross-loading",
                "Makna untuk riset": "Variabel yang loading ≥0.40 pada lebih dari satu komponen: " + ", ".join(cross[:10]) + ". Ini menunjukkan variabel mungkin mengukur lebih dari satu dimensi.",
                "Saran tindakan": "Cek redaksi/definisi variabel. Untuk kuesioner, pertimbangkan EFA dengan rotasi atau revisi item.",
            })

    if selected_cols is not None or n_rows is not None:
        rows.append({
            "Aspek": "Konteks data",
            "Makna untuk riset": f"PCA dijalankan pada {len(selected_cols) if selected_cols is not None else 'beberapa'} variabel dan {n_rows if n_rows is not None else 'sejumlah'} kasus lengkap.",
            "Saran tindakan": "Untuk hasil yang stabil, jumlah responden sebaiknya memadai dan variabel memiliki korelasi yang bermakna. PCA merangkum variasi data; ia bukan bukti final validitas konstruk.",
        })

    rows.append({
        "Aspek": "Kesimpulan praktis",
        "Makna untuk riset": "PCA membantu mereduksi banyak variabel menjadi beberapa komponen utama sehingga analisis dan interpretasi lebih ringkas.",
        "Saran tindakan": "Gunakan skor komponen untuk analisis lanjutan bila tujuannya reduksi data. Jika tujuannya menemukan konstruk laten kuesioner, bandingkan dengan EFA dan teori.",
    })
    return pd.DataFrame(rows)


def interpret_efa_results(kmo_table, loadings, variance=None, communalities=None, eigen_table=None, alpha=None):
    """Bangun interpretasi riset untuk EFA/PAF agar user awam tahu makna output."""
    alpha = alpha if alpha is not None else st.session_state.get("active_alpha", 0.05)
    rows = []
    try:
        kmo_df = kmo_table.copy() if isinstance(kmo_table, pd.DataFrame) else pd.DataFrame(kmo_table)
    except Exception:
        kmo_df = pd.DataFrame()

    if not kmo_df.empty:
        kmo_val = _first_numeric_value(kmo_df, ["KMO", "KMO Overall"])
        bart_p = _first_numeric_value(kmo_df, ["Bartlett p-value", "p-value", "p"])
        if kmo_val is not None:
            if kmo_val >= 0.80:
                desc = "baik/sangat memadai"
                action = "Analisis faktor relatif layak dilanjutkan."
            elif kmo_val >= 0.70:
                desc = "cukup baik"
                action = "Analisis faktor dapat dilanjutkan, sambil cek communality dan loading."
            elif kmo_val >= 0.60:
                desc = "cukup/minimal"
                action = "Masih bisa dipakai untuk eksplorasi, tetapi interpretasi harus hati-hati."
            elif kmo_val >= 0.50:
                desc = "lemah"
                action = "Pertimbangkan buang item yang lemah, tambah sampel, atau revisi konstruk."
            else:
                desc = "tidak memadai"
                action = "EFA sebaiknya tidak dijadikan dasar kesimpulan kuat; cek korelasi antar item dan jumlah sampel."
            rows.append({
                "Aspek": "Kelayakan EFA - KMO",
                "Makna untuk riset": f"KMO = {_format_number(kmo_val)}; kelayakan sampling {desc}.",
                "Saran tindakan": action,
            })
        if bart_p is not None:
            if bart_p < alpha:
                meaning = "Bartlett signifikan; matriks korelasi tidak identik, sehingga ada dasar korelasi untuk analisis faktor."
                action = "Lanjutkan membaca factor loading, communality, dan variance explained."
            else:
                meaning = "Bartlett tidak signifikan; korelasi antar item mungkin belum cukup kuat untuk membentuk faktor."
                action = "Cek apakah item memang satu domain, tambah sampel, atau pertimbangkan tidak memakai EFA."
            rows.append({
                "Aspek": "Kelayakan EFA - Bartlett",
                "Makna untuk riset": f"p Bartlett = {p_value_text(bart_p)}. {meaning}",
                "Saran tindakan": action,
            })

    try:
        load = loadings.copy() if isinstance(loadings, pd.DataFrame) else pd.DataFrame(loadings)
    except Exception:
        load = pd.DataFrame()
    if not load.empty:
        variable_col = _first_existing_column(load, ["Variable", "Variabel", "Item"])
        factor_cols = [c for c in load.columns if str(c).lower().startswith("factor")]
        if variable_col is not None and factor_cols:
            weak, cross, assignments = [], [], []
            for _, row in load.iterrows():
                vals = pd.to_numeric(row[factor_cols], errors="coerce").abs().dropna()
                if vals.empty:
                    continue
                best_factor = vals.idxmax()
                best_val = float(vals.max())
                strong_count = int((vals >= 0.40).sum())
                item = str(row[variable_col])
                assignments.append((item, str(best_factor), best_val))
                if best_val < 0.40:
                    weak.append(item)
                if strong_count >= 2:
                    cross.append(item)
            if assignments:
                top_txt = "; ".join([f"{v} → {f} ({val:.2f})" for v, f, val in sorted(assignments, key=lambda x: x[2], reverse=True)[:8]])
                rows.append({
                    "Aspek": "Struktur faktor",
                    "Makna untuk riset": f"Item/variabel paling kuat memuat faktor: {top_txt}.",
                    "Saran tindakan": "Beri nama faktor berdasarkan kumpulan item dengan loading tertinggi dan dasar teori, bukan hanya angka terbesar."
                })
            if weak:
                rows.append({
                    "Aspek": "Item lemah",
                    "Makna untuk riset": "Item dengan loading tertinggi < 0.40: " + ", ".join(weak[:10]) + ". Item ini belum jelas masuk ke faktor mana.",
                    "Saran tindakan": "Pertimbangkan revisi atau penghapusan item jika tidak penting secara teori."
                })
            if cross:
                rows.append({
                    "Aspek": "Cross-loading",
                    "Makna untuk riset": "Item yang memuat lebih dari satu faktor: " + ", ".join(cross[:10]) + ". Item ini dapat menimbulkan ambiguitas konstruk.",
                    "Saran tindakan": "Cek redaksi item; pertimbangkan rotasi lain atau keputusan teoritis."
                })

    try:
        comm = communalities.copy() if isinstance(communalities, pd.DataFrame) else pd.DataFrame(communalities)
    except Exception:
        comm = pd.DataFrame()
    if not comm.empty:
        comm_col = _first_existing_column(comm, ["Communality", "Extraction", "h2"])
        item_col = _first_existing_column(comm, ["Variable", "Variabel", "Item"])
        if comm_col is not None:
            comm["__h2"] = pd.to_numeric(comm[comm_col], errors="coerce")
            low = comm[comm["__h2"] < 0.30]
            if len(low):
                items = ", ".join(low[item_col].astype(str).head(10)) if item_col else f"{len(low)} item"
                rows.append({
                    "Aspek": "Communality rendah",
                    "Makna untuk riset": f"Ada item dengan communality < 0.30: {items}. Faktor belum menjelaskan item tersebut dengan baik.",
                    "Saran tindakan": "Tinjau item ini. Jika tidak penting secara teori, pertimbangkan dikeluarkan dan jalankan ulang EFA."
                })

    try:
        var_df = variance.copy() if isinstance(variance, pd.DataFrame) else pd.DataFrame(variance)
    except Exception:
        var_df = pd.DataFrame()
    if not var_df.empty:
        cum_col = _first_existing_column(var_df, ["Cumulative %", "Cumulative Variance %"])
        if cum_col is not None:
            vals = pd.to_numeric(var_df[cum_col], errors="coerce").dropna()
            if len(vals):
                cum = float(vals.iloc[-1])
                rows.append({
                    "Aspek": "Varians faktor",
                    "Makna untuk riset": f"Faktor yang dipilih menjelaskan sekitar {_format_number(cum, 2)}% variasi bersama item.",
                    "Saran tindakan": "Jika persentase rendah, cek item lemah, jumlah faktor, atau kesesuaian konstruk dengan teori."
                })

    if not rows:
        rows.append({
            "Aspek": "Status EFA",
            "Makna untuk riset": "Output EFA belum cukup untuk ditafsirkan otomatis.",
            "Saran tindakan": "Pastikan KMO/Bartlett, loading, communalities, dan variance table tersedia."
        })
    return pd.DataFrame(rows)

def cronbach_alpha(df_items):
    data = df_items.apply(pd.to_numeric, errors="coerce").dropna()
    k = data.shape[1]
    if k < 2 or data.shape[0] < 2:
        return np.nan, pd.DataFrame()
    item_vars = data.var(axis=0, ddof=1)
    total_var = data.sum(axis=1).var(ddof=1)
    alpha = (k / (k - 1)) * (1 - item_vars.sum() / total_var) if total_var != 0 else np.nan
    rows = []
    for col in data.columns:
        total_minus = data.drop(columns=[col]).sum(axis=1)
        corr = data[col].corr(total_minus)
        alpha_if_deleted, _ = cronbach_alpha(data.drop(columns=[col])) if k > 2 else (np.nan, None)
        rows.append({"Item": col, "Corrected Item-Total Corr": corr, "Alpha if Deleted": alpha_if_deleted})
    return alpha, pd.DataFrame(rows).round(5)


def vif_table(X):
    if variance_inflation_factor is None:
        return pd.DataFrame()
    Xc = sm.add_constant(X, has_constant="add")
    rows = []
    for i, col in enumerate(Xc.columns):
        if col == "const":
            continue
        rows.append({"Variabel": col, "VIF": variance_inflation_factor(Xc.values, i)})
    return pd.DataFrame(rows).round(4)


def make_design_matrix(df, predictors):
    X = df[predictors].copy()
    for c in predictors:
        if not pd.api.types.is_numeric_dtype(X[c]):
            X[c] = X[c].astype("category")
    X = pd.get_dummies(X, drop_first=True, dtype=float)
    return X



# -----------------------------------------------------------------------------
# Data compatibility assistant for beginner users
# -----------------------------------------------------------------------------
def _sev_rank(severity):
    order = {"Kritis": 0, "Tinggi": 1, "Sedang": 2, "Ringan": 3, "Info": 4, "OK": 5}
    return order.get(str(severity), 9)


def _issue(severity, area, problem, columns, why, fix, easy_action, status=None):
    return {
        "Status": status or ("✅ Siap" if severity == "OK" else "⚠️ Perlu dicek"),
        "Prioritas": severity,
        "Area": area,
        "Masalah/Tanda": problem,
        "Kolom Terdampak": ", ".join(map(str, columns)) if isinstance(columns, (list, tuple, set)) else str(columns),
        "Mengapa Penting": why,
        "Sebaiknya Diubah/Ditambahkan/Diganti": fix,
        "Langkah Mudah untuk User Awam": easy_action,
    }


def _is_numeric_like(series, threshold=0.80):
    s = series.dropna()
    if s.empty or pd.api.types.is_numeric_dtype(series):
        return False, 0.0
    cleaned = s.astype(str).str.strip()
    cleaned = cleaned.str.replace(".", "", regex=False).str.replace(",", ".", regex=False)
    cleaned = cleaned.str.replace(r"[^0-9\-\.]+", "", regex=True)
    converted = pd.to_numeric(cleaned.replace("", np.nan), errors="coerce")
    ratio = float(converted.notna().mean()) if len(converted) else 0.0
    return ratio >= threshold, ratio


def _is_date_like(series, threshold=0.80):
    s = series.dropna()
    if s.empty or pd.api.types.is_datetime64_any_dtype(series):
        return False, 0.0
    if pd.api.types.is_numeric_dtype(series):
        return False, 0.0
    parsed = pd.to_datetime(s.astype(str), errors="coerce", dayfirst=True)
    ratio = float(parsed.notna().mean()) if len(parsed) else 0.0
    return ratio >= threshold, ratio


def _column_profile(df, metadata=None):
    rows = []
    meta_map = {}
    if metadata is not None and isinstance(metadata, pd.DataFrame) and "Name" in metadata.columns:
        try:
            meta_map = metadata.set_index("Name").to_dict("index")
        except Exception:
            meta_map = {}
    for col in df.columns:
        s = df[col]
        miss = int(s.isna().sum())
        nonmiss = int(s.notna().sum())
        unique = int(s.nunique(dropna=True))
        numeric_like, numeric_ratio = _is_numeric_like(s)
        date_like, date_ratio = _is_date_like(s)
        measure = meta_map.get(col, {}).get("Measure", infer_measure(s))
        role = meta_map.get(col, {}).get("Role", "Input")
        rows.append({
            "Kolom": col,
            "Tipe terbaca": str(s.dtype),
            "Measure": measure,
            "Role": role,
            "Non-missing": nonmiss,
            "Missing": miss,
            "% Missing": round((miss / max(len(df), 1)) * 100, 2),
            "Nilai unik": unique,
            "Contoh nilai": ", ".join(map(str, s.dropna().astype(str).head(3).tolist())) if nonmiss else "-",
            "Terlihat numerik?": "Ya" if numeric_like else "Tidak",
            "Terlihat tanggal?": "Ya" if date_like else "Tidak",
            "Catatan singkat": _quick_column_note(s, measure, numeric_like, date_like, unique, len(df)),
        })
    return pd.DataFrame(rows)


def _quick_column_note(series, measure, numeric_like, date_like, unique, n_rows):
    if series.isna().all():
        return "Kosong; perlu diisi atau dihapus."
    if numeric_like:
        return "Seharusnya angka; ubah menjadi numeric sebelum analisis."
    if date_like:
        return "Sepertinya tanggal; ubah menjadi date/datetime bila dipakai sebagai waktu."
    if unique <= 1:
        return "Variasi terlalu rendah; biasanya tidak berguna untuk uji statistik."
    if unique == n_rows and n_rows > 10:
        return "Kemungkinan ID/kode unik; jangan dipakai sebagai grup/prediktor utama tanpa alasan."
    if pd.api.types.is_numeric_dtype(series) and str(measure).lower() in ["nominal", "ordinal"] and unique > 15:
        return "Tipe metadata mungkin perlu diubah ke Scale bila ini skor/ukuran kontinu."
    return "Cukup aman."


def analyze_data_compatibility(df, metadata=None):
    issues = []
    n_rows, n_cols = df.shape
    num = numeric_cols(df)
    cat = categorical_cols(df)

    if n_rows == 0 or n_cols == 0:
        issues.append(_issue(
            "Kritis", "Struktur data", "Dataset kosong", "Semua data",
            "Analisis tidak dapat dilakukan tanpa baris dan kolom.",
            "Tambahkan data mentah dengan minimal satu variabel dan beberapa observasi.",
            "Upload ulang file yang benar atau gunakan Data Contoh untuk mencoba alur aplikasi.",
        ))
        return pd.DataFrame(issues)

    if n_rows < 5:
        issues.append(_issue(
            "Tinggi", "Ukuran sampel", f"Jumlah baris sangat kecil: {n_rows}", "Semua analisis inferensial",
            "Sebagian besar uji statistik membutuhkan sampel memadai agar hasil stabil.",
            "Tambahkan responden/observasi. Untuk latihan boleh lanjut, tetapi jangan jadikan hasil sebagai kesimpulan riset final.",
            "Isi minimal 20-30 baris untuk analisis sederhana; EFA/regresi biasanya perlu lebih banyak.",
        ))
    elif n_rows < 30:
        issues.append(_issue(
            "Sedang", "Ukuran sampel", f"Jumlah baris masih terbatas: {n_rows}", "Semua analisis inferensial",
            "Sampel kecil membuat p-value dan estimasi efek kurang stabil.",
            "Laporkan keterbatasan ukuran sampel dan prioritaskan effect size serta uji nonparametrik bila asumsi tidak terpenuhi.",
            "Jika memungkinkan, tambahkan data sampai minimal 30 baris untuk analisis dasar.",
        ))

    duplicated = df.columns[df.columns.duplicated()].tolist()
    if duplicated:
        issues.append(_issue(
            "Kritis", "Nama kolom", "Ada nama kolom ganda", duplicated,
            "Aplikasi dan rumus analisis dapat salah memilih kolom jika namanya sama.",
            "Ganti nama kolom agar unik, misalnya pre_test dan post_test.",
            "Buka Transform → Rename/Drop Variables, lalu ubah nama kolom yang duplikat.",
        ))

    unnamed = [c for c in df.columns if str(c).strip() == "" or str(c).lower().startswith("unnamed")]
    if unnamed:
        issues.append(_issue(
            "Sedang", "Nama kolom", "Ada kolom tanpa nama jelas", unnamed,
            "Pengguna awam sulit memilih variabel jika nama kolom tidak bermakna.",
            "Ganti nama menjadi deskriptif, misalnya jenis_kelamin, usia, skor_total.",
            "Buka Transform → Rename/Drop Variables, lalu beri nama yang pendek tanpa spasi.",
        ))

    if not num:
        maybe_num = []
        for col in df.columns:
            is_like, ratio = _is_numeric_like(df[col])
            if is_like:
                maybe_num.append(col)
        issues.append(_issue(
            "Tinggi", "Tipe data", "Tidak ada kolom numerik yang terbaca", maybe_num or "-",
            "T-test, ANOVA, korelasi, regresi, reliabilitas, PCA, dan EFA membutuhkan angka.",
            "Ubah kolom skor/nilai menjadi numeric. Bersihkan simbol %, Rp, spasi, atau koma desimal yang salah.",
            "Cek kolom yang terlihat seperti angka, lalu gunakan Transform → Compute Variable atau bersihkan file CSV/Excel.",
        ))
    else:
        issues.append(_issue(
            "OK", "Tipe data", f"Ada {len(num)} kolom numerik", num[:8],
            "Analisis berbasis skor dapat dilakukan pada kolom numerik.",
            "Pastikan kolom numerik benar-benar skor/ukuran, bukan kode kategori seperti 1=Laki-laki.",
            "Gunakan Variable View untuk menandai Scale/Nominal/Ordinal.",
        ))

    numeric_like_cols = []
    date_like_cols = []
    for col in df.columns:
        is_num_like, ratio = _is_numeric_like(df[col])
        if is_num_like:
            numeric_like_cols.append(f"{col} ({ratio*100:.0f}% terlihat angka)")
        is_date_like, dratio = _is_date_like(df[col])
        if is_date_like:
            date_like_cols.append(f"{col} ({dratio*100:.0f}% terlihat tanggal)")
    if numeric_like_cols:
        issues.append(_issue(
            "Tinggi", "Konversi angka", "Kolom terlihat numerik tetapi terbaca teks", numeric_like_cols,
            "Kolom seperti ini tidak muncul di daftar variabel numerik sehingga analisis bisa tidak tersedia.",
            "Konversi menjadi angka; samakan tanda desimal, hapus satuan/simbol mata uang, dan pastikan missing value kosong/NA.",
            "Saat upload CSV, coba ubah 'Tanda desimal' menjadi koma atau titik. Jika masih gagal, bersihkan di Excel lalu upload ulang.",
        ))
    if date_like_cols:
        issues.append(_issue(
            "Info", "Tanggal/waktu", "Kolom terlihat seperti tanggal", date_like_cols,
            "Tanggal tidak boleh diperlakukan sebagai angka biasa kecuali sudah diubah menjadi durasi/umur/periode.",
            "Ubah tanggal menjadi variabel yang bermakna, misalnya usia_hari, bulan, semester, atau before/after.",
            "Untuk analisis sederhana, buat kolom baru di Excel: umur, lama_kerja, bulan_pengamatan, atau kategori periode.",
        ))

    placeholder_tokens = ["-", "--", "?", "NA", "N/A", "null", "None", "missing", "tidak ada"]
    placeholder_cols = []
    for col in df.columns:
        if df[col].dtype == object:
            vals = df[col].astype(str).str.strip()
            count = int(vals.isin(placeholder_tokens).sum())
            if count:
                placeholder_cols.append(f"{col} ({count} nilai)")
    if placeholder_cols:
        issues.append(_issue(
            "Sedang", "Missing value", "Ada kode missing yang masih terbaca sebagai teks", placeholder_cols,
            "Kode seperti '-' atau 'NA' bisa dianggap kategori valid sehingga frekuensi/crosstab bias.",
            "Ubah kode tersebut menjadi kosong/NaN atau daftarkan sebagai Missing Values di Variable View.",
            "Buka Variable View → isi Missing Values, misalnya `-, NA, 99`, atau bersihkan file sebelum upload.",
        ))

    missing_high = []
    missing_any = []
    for col in df.columns:
        pct = df[col].isna().mean() * 100
        if pct >= 40:
            missing_high.append(f"{col} ({pct:.1f}%)")
        elif pct > 0:
            missing_any.append(f"{col} ({pct:.1f}%)")
    if missing_high:
        issues.append(_issue(
            "Tinggi", "Missing value", "Missing value sangat tinggi", missing_high,
            "Kolom dengan banyak data kosong dapat melemahkan analisis dan membuat hasil tidak representatif.",
            "Pertimbangkan hapus kolom, imputasi, atau kumpulkan data ulang tergantung konteks riset.",
            "Untuk pemula: jika missing >40%, jangan jadikan variabel utama kecuali ada alasan metodologis kuat.",
        ))
    if missing_any:
        issues.append(_issue(
            "Ringan", "Missing value", "Ada beberapa data kosong", missing_any[:10],
            "Sebagian analisis akan otomatis membuang baris kosong sehingga N bisa berbeda antar output.",
            "Tentukan strategi: listwise deletion, pairwise deletion, imputasi mean/median, atau kategori 'Tidak menjawab'.",
            "Laporkan jumlah missing di bagian metode/hasil, lalu cek N pada setiap output.",
        ))

    constant_cols = [c for c in df.columns if df[c].nunique(dropna=True) <= 1]
    if constant_cols:
        issues.append(_issue(
            "Tinggi", "Variasi data", "Ada kolom tanpa variasi", constant_cols,
            "Kolom konstan tidak bisa menjelaskan perbedaan/hubungan karena nilainya sama semua.",
            "Hapus dari analisis atau cek apakah data salah input.",
            "Gunakan Transform → Drop Variables bila memang kolom tidak diperlukan.",
        ))

    id_like = [c for c in df.columns if df[c].nunique(dropna=True) == len(df) and len(df) >= 10]
    if id_like:
        issues.append(_issue(
            "Info", "ID/kode unik", "Kolom tampak seperti ID unik", id_like[:8],
            "ID unik biasanya tidak cocok sebagai grup, prediktor kategorik, atau item reliabilitas.",
            "Tandai sebagai Role=None atau gunakan hanya untuk pelacakan responden.",
            "Di Variable View, ubah Role menjadi None/ID secara konseptual dan jangan pilih kolom ini untuk uji statistik.",
        ))

    high_card_cat = []
    for c in cat:
        u = df[c].nunique(dropna=True)
        if u > 20 and u < len(df):
            high_card_cat.append(f"{c} ({u} kategori)")
    if high_card_cat:
        issues.append(_issue(
            "Sedang", "Kategori", "Kategori terlalu banyak untuk variabel grup", high_card_cat,
            "ANOVA/t-test/crosstab sulit dimaknai jika grup terlalu banyak atau terlalu kecil.",
            "Gabungkan kategori yang serupa atau buat kategori baru yang lebih ringkas.",
            "Gunakan Transform → Recode Variable untuk mengelompokkan kategori.",
        ))

    sparse_groups = []
    for c in cat:
        vc = df[c].dropna().astype(str).value_counts()
        if len(vc) >= 2 and (vc < 2).any():
            sparse_groups.append(f"{c} (ada kategori n<2)")
    if sparse_groups:
        issues.append(_issue(
            "Sedang", "Ukuran grup", "Ada kategori dengan anggota sangat sedikit", sparse_groups[:8],
            "Uji beda butuh data memadai di setiap grup; grup sangat kecil membuat hasil tidak stabil.",
            "Gabungkan kategori kecil, hapus kategori outlier, atau tambah data untuk kategori tersebut.",
            "Cek Deskriptif → Frekuensi Kategori sebelum memakai kolom sebagai grup.",
        ))

    outlier_cols = []
    for c in num:
        s = safe_numeric(df[c]).dropna()
        if len(s) >= 5 and s.std(ddof=1) > 0:
            z = np.abs(stats.zscore(s, nan_policy="omit"))
            count = int(np.sum(z > 3))
            if count:
                outlier_cols.append(f"{c} ({count} outlier |z|>3)")
    if outlier_cols:
        issues.append(_issue(
            "Ringan", "Outlier", "Ada nilai ekstrem pada variabel numerik", outlier_cols[:10],
            "Outlier dapat memengaruhi mean, korelasi, regresi, t-test, dan ANOVA.",
            "Periksa apakah outlier adalah salah input, fenomena nyata, atau perlu transformasi/winsorizing.",
            "Buka Visualisasi → boxplot/histogram untuk melihat bentuk distribusi sebelum memutuskan.",
        ))

    if metadata is not None and isinstance(metadata, pd.DataFrame) and not metadata.empty:
        mismatch = []
        for _, row in metadata.iterrows():
            c = row.get("Name")
            if c not in df.columns:
                continue
            measure = str(row.get("Measure", "")).lower()
            if measure == "scale" and not pd.api.types.is_numeric_dtype(df[c]):
                mismatch.append(f"{c} (Scale tetapi bukan numeric)")
            if measure in ["nominal", "ordinal"] and pd.api.types.is_numeric_dtype(df[c]) and df[c].nunique(dropna=True) > 20:
                mismatch.append(f"{c} ({row.get('Measure')} tetapi unik >20)")
        if mismatch:
            issues.append(_issue(
                "Sedang", "Variable View", "Measure level mungkin tidak cocok", mismatch[:10],
                "Measure level memandu pemilihan uji statistik. Salah label bisa membuat pengguna memilih analisis yang keliru.",
                "Perbaiki Measure: Scale untuk skor/usia/nilai kontinu; Nominal untuk kategori; Ordinal untuk urutan/Likert item.",
                "Buka Data → Variable View, lalu ubah Measure sesuai arti variabel.",
            ))

    if not cat:
        issues.append(_issue(
            "Info", "Variabel grup", "Tidak ada kolom kategorik/teks", "-",
            "Uji beda antar kelompok butuh variabel grup, misalnya gender/kelas/perlakuan.",
            "Jika desain riset membandingkan kelompok, tambahkan kolom grup atau recode skor menjadi kategori yang bermakna.",
            "Contoh kolom grup: kelompok, kelas, jenis_kelamin, perlakuan, kategori_usia.",
        ))

    if not issues:
        issues.append(_issue(
            "OK", "Kesiapan data", "Tidak ditemukan masalah besar", "Semua data",
            "Struktur data cukup aman untuk eksplorasi awal.",
            "Tetap cek asumsi sesuai analisis yang dipilih.",
            "Lanjut ke Deskriptif/Uji Statistik, lalu baca tab Insight Riset.",
        ))

    result = pd.DataFrame(issues)
    result["_rank"] = result["Prioritas"].map(_sev_rank)
    return result.sort_values(["_rank", "Area"]).drop(columns="_rank").reset_index(drop=True)


def compatibility_score(issues_df):
    if issues_df is None or issues_df.empty:
        return 100
    penalties = {"Kritis": 35, "Tinggi": 20, "Sedang": 10, "Ringan": 4, "Info": 0, "OK": 0}
    score = 100
    for sev in issues_df.get("Prioritas", []):
        score -= penalties.get(str(sev), 0)
    return int(max(0, min(100, score)))


def beginner_data_recipe(df, issues_df):
    rows = []
    score = compatibility_score(issues_df)
    if score < 60:
        rows.append({"Urutan": 1, "Langkah": "Rapikan struktur file", "Apa yang dilakukan": "Pastikan baris pertama adalah nama kolom, tidak ada kolom kosong/duplikat, dan setiap baris adalah satu responden/observasi."})
        rows.append({"Urutan": 2, "Langkah": "Bersihkan tipe data", "Apa yang dilakukan": "Kolom skor/nilai harus terbaca numeric. Hapus simbol Rp, %, spasi, atau tanda koma/titik yang salah."})
        rows.append({"Urutan": 3, "Langkah": "Tangani missing value", "Apa yang dilakukan": "Ubah kode '-', 99, atau NA menjadi missing value yang konsisten, lalu tentukan apakah dihapus atau diimputasi."})
    else:
        rows.append({"Urutan": 1, "Langkah": "Cek Variable View", "Apa yang dilakukan": "Pastikan Measure sudah benar: Scale untuk skor/nilai/usia, Nominal untuk kategori, Ordinal untuk Likert/tingkatan."})
        rows.append({"Urutan": 2, "Langkah": "Jalankan deskriptif", "Apa yang dilakukan": "Lihat mean, standar deviasi, frekuensi, dan missing value sebelum uji hipotesis."})
    rows.append({"Urutan": len(rows)+1, "Langkah": "Pilih uji sesuai desain", "Apa yang dilakukan": "Hubungan dua skor → korelasi; beda dua grup → t-test; beda >2 grup → ANOVA/Kruskal; prediksi → regresi; kuesioner → reliabilitas/EFA."})
    rows.append({"Urutan": len(rows)+1, "Langkah": "Baca insight riset", "Apa yang dilakukan": "Setelah output dibuat, buka tab Insight Riset untuk menafsirkan p-value, effect size, asumsi, dan makna substantif."})
    return pd.DataFrame(rows)


def suggest_compatible_analyses(df):
    num = numeric_cols(df)
    cat = categorical_cols(df)
    rows = []

    def add(name, readiness, needed, suggestion, beginner_note):
        rows.append({"Analisis": name, "Kesiapan": readiness, "Yang Dibutuhkan": needed, "Saran Variabel": suggestion, "Catatan untuk User Awam": beginner_note})

    if num:
        add("Statistik deskriptif", "✅ Siap", "Minimal 1 kolom numerik", ", ".join(num[:6]), "Mulai dari sini untuk memahami nilai tengah, sebaran, dan missing.")
    else:
        add("Statistik deskriptif numerik", "❌ Belum siap", "Minimal 1 kolom numerik", "Ubah kolom skor/nilai menjadi angka", "Cek apakah angka masih terbaca sebagai teks.")

    if len(num) >= 2:
        add("Korelasi Pearson/Spearman", "✅ Siap", "Minimal 2 kolom numerik", f"{num[0]} + {num[1]}", "Gunakan untuk melihat hubungan antar skor/nilai.")
    else:
        add("Korelasi", "❌ Belum siap", "Minimal 2 kolom numerik", "Tambahkan/ubah satu lagi variabel numeric", "Contoh: motivasi dan prestasi.")

    two_level = []
    multi_level = []
    for c in cat + num:
        u = df[c].dropna().nunique()
        if u == 2:
            two_level.append(c)
        if 3 <= u <= 10:
            multi_level.append(c)
    if num and two_level:
        add("Independent T-Test", "✅ Siap", "1 skor numeric + 1 grup dua kategori", f"DV: {num[0]}, Grup: {two_level[0]}", "Cocok untuk membandingkan rata-rata dua kelompok.")
    else:
        add("Independent T-Test", "⚠️ Perlu disiapkan", "1 skor numeric + 1 grup dua kategori", "Tambahkan kolom grup dengan tepat 2 kategori", "Contoh grup: kontrol vs eksperimen, laki-laki vs perempuan.")

    if len(num) >= 2:
        add("Paired T-Test", "✅ Siap secara format wide", "2 kolom numeric berpasangan", f"Before: {num[0]}, After: {num[1]}", "Pastikan dua kolom berasal dari orang/objek yang sama sebelum-sesudah.")
    else:
        add("Paired T-Test", "⚠️ Perlu disiapkan", "2 kolom numeric berpasangan", "Tambahkan kolom pre dan post", "Contoh: pretest dan posttest.")

    if num and multi_level:
        add("One-Way ANOVA", "✅ Siap", "1 skor numeric + 1 grup dengan ≥3 kategori", f"DV: {num[0]}, Grup: {multi_level[0]}", "Cocok untuk membandingkan rata-rata lebih dari dua kelompok.")
    elif len(num) >= 3:
        add("One-Way ANOVA format wide", "✅ Siap secara format wide", "≥3 kolom numeric yang mewakili grup", ", ".join(num[:3]), "Setiap kolom dianggap grup terpisah.")
    else:
        add("One-Way ANOVA", "⚠️ Perlu disiapkan", "1 skor numeric + grup ≥3 kategori, atau ≥3 kolom grup numeric", "Tambahkan kolom grup atau kolom skor per grup", "Contoh grup: kelas A/B/C.")

    if len(cat) >= 2:
        add("Crosstab & Chi-Square", "✅ Siap", "2 kolom kategorik", f"{cat[0]} × {cat[1]}", "Gunakan untuk hubungan antar kategori.")
    else:
        add("Crosstab & Chi-Square", "⚠️ Perlu disiapkan", "2 kolom kategorik", "Tambahkan minimal 2 variabel kategori", "Contoh: gender dan pilihan_produk.")

    if len(num) >= 2:
        add("Regresi Linear", "✅ Siap format dasar", "1 target numeric + minimal 1 prediktor", f"Y: {num[0]}, X: {', '.join(num[1:4])}", "Gunakan untuk memprediksi nilai Y dari X.")
    else:
        add("Regresi Linear", "❌ Belum siap", "Minimal 2 kolom numeric atau prediktor kategorik yang jelas", "Tambahkan target dan prediktor", "Contoh: prestasi diprediksi oleh motivasi dan jam_belajar.")

    if len(num) >= 3:
        add("Reliabilitas Cronbach Alpha", "✅ Siap secara teknis", "Minimal 3 item numeric satu konstruk", ", ".join(num[:6]), "Pastikan item memang mengukur konstruk yang sama.")
    else:
        add("Reliabilitas Cronbach Alpha", "⚠️ Perlu disiapkan", "Minimal 3 item numeric", "Tambahkan item kuesioner", "Contoh: item_motivasi_1 sampai item_motivasi_5.")

    if len(num) >= 3 and len(df) >= max(30, len(num) * 5):
        add("EFA / Analisis Faktor", "✅ Cukup siap", "≥3 item numeric dan sampel cukup", ", ".join(num[:8]), "Idealnya jumlah responden ≥5 kali jumlah item, lebih baik ≥100.")
    elif len(num) >= 3:
        add("EFA / Analisis Faktor", "⚠️ Bisa dicoba, sampel mungkin kurang", "≥3 item numeric dan sampel cukup", ", ".join(num[:8]), "Gunakan sebagai eksplorasi; hati-hati menafsirkan jika N kecil.")
    else:
        add("EFA / Analisis Faktor", "❌ Belum siap", "Minimal 3 item numeric", "Tambahkan item kuesioner numeric", "EFA tidak cocok untuk hanya 1-2 variabel.")

    return pd.DataFrame(rows)


def suggest_analysis_table(df):
    """Return quick-start analysis suggestions with stable column names.

    v4.0/v4.1 memanggil helper ini dari halaman Mulai Cepat, tetapi
    implementasi yang ikut terpaket adalah `suggest_compatible_analyses`.
    Fungsi ini menjadi adaptor resmi: output-nya lebih ringkas untuk user
    awam dan aman walau dataset kosong/kolom campuran.
    """
    try:
        base = suggest_compatible_analyses(df)
    except Exception as exc:
        return pd.DataFrame([{
            "Uji/Analisis": "Belum dapat memberi rekomendasi",
            "Status": "⚠️ Perlu cek data",
            "Data yang Dibutuhkan": "Dataset yang valid",
            "Saran Variabel": "Periksa format file dan tipe kolom",
            "Catatan Singkat": f"Rekomendasi gagal dibuat: {exc}",
            "Menu yang Digunakan": "🧰 Kompatibilitas Data",
        }])

    if base is None or len(base) == 0:
        return pd.DataFrame([{
            "Uji/Analisis": "Deskriptif awal",
            "Status": "ℹ️ Mulai dari data",
            "Data yang Dibutuhkan": "Minimal 1 kolom data",
            "Saran Variabel": "Upload atau input data terlebih dahulu",
            "Catatan Singkat": "Setelah data tersedia, aplikasi akan menyarankan uji yang cocok.",
            "Menu yang Digunakan": "📥 Input Data",
        }])

    rename_map = {
        "Analisis": "Uji/Analisis",
        "Kesiapan": "Status",
        "Yang Dibutuhkan": "Data yang Dibutuhkan",
        "Catatan untuk User Awam": "Catatan Singkat",
    }
    out = base.rename(columns=rename_map).copy()

    # Lengkapi kolom yang mungkin belum ada agar UI/export tidak error.
    for col in ["Uji/Analisis", "Status", "Data yang Dibutuhkan", "Saran Variabel", "Catatan Singkat"]:
        if col not in out.columns:
            out[col] = "-"

    def _menu_for_analysis(name):
        name = str(name).lower()
        if any(k in name for k in ["deskriptif", "frekuensi"]):
            return "📋 Deskriptif"
        if any(k in name for k in ["korelasi", "t-test", "anova", "chi-square", "mann", "wilcoxon", "kruskal"]):
            return "🧪 Uji Statistik / 🧙 Smart Assistant"
        if "regresi" in name:
            return "📈 Regresi"
        if any(k in name for k in ["reliabilitas", "cronbach", "efa", "faktor", "pca"]):
            return "🧭 Reliabilitas & Faktor"
        return "🧙 Smart Assistant"

    out["Menu yang Digunakan"] = out["Uji/Analisis"].apply(_menu_for_analysis)

    # Urutkan supaya rekomendasi yang siap tampil dulu, lalu yang perlu disiapkan.
    def _rank(status):
        s = str(status).lower()
        if "✅" in s or "siap" in s:
            return 0
        if "⚠" in s or "perlu" in s or "bisa" in s:
            return 1
        if "❌" in s or "belum" in s:
            return 2
        return 3

    out["_rank"] = out["Status"].apply(_rank)
    out = out.sort_values(["_rank", "Uji/Analisis"], kind="stable").drop(columns="_rank")
    return out[["Uji/Analisis", "Status", "Data yang Dibutuhkan", "Saran Variabel", "Catatan Singkat", "Menu yang Digunakan"]].reset_index(drop=True)


def analysis_specific_guidance(df, analysis_name):
    num = numeric_cols(df)
    cat = categorical_cols(df)
    rows = []
    def row(check, status, fix):
        rows.append({"Checklist": check, "Status": status, "Apa yang perlu dilakukan": fix})

    if analysis_name == "Independent T-Test":
        two_level = [c for c in df.columns if df[c].dropna().nunique() == 2]
        row("Ada variabel numeric sebagai nilai/DV", "✅ Ada" if num else "❌ Belum ada", "Ubah skor/nilai menjadi numeric atau tambahkan kolom skor.")
        row("Ada variabel grup tepat 2 kategori", "✅ Ada: " + ", ".join(two_level[:4]) if two_level else "❌ Belum ada", "Tambahkan/recode grup menjadi dua kategori, misalnya kontrol vs eksperimen.")
        row("Setiap grup minimal 2 data", "ℹ️ Cek frekuensi", "Buka Deskriptif → Frekuensi Kategori untuk memastikan setiap grup punya cukup data.")
    elif analysis_name == "ANOVA":
        multi = [c for c in df.columns if 3 <= df[c].dropna().nunique() <= 10]
        row("Ada variabel numeric sebagai DV", "✅ Ada" if num else "❌ Belum ada", "Ubah skor/nilai menjadi numeric.")
        row("Ada grup ≥3 kategori", "✅ Ada: " + ", ".join(multi[:4]) if multi else "⚠️ Belum jelas", "Tambahkan/recode kolom grup menjadi 3+ kategori yang bermakna.")
        row("Ukuran grup memadai", "ℹ️ Cek frekuensi", "Gabungkan kategori dengan n sangat kecil atau tambah data.")
    elif analysis_name == "Korelasi":
        row("Minimal 2 variabel numeric", "✅ Siap" if len(num) >= 2 else "❌ Belum siap", "Tambahkan/konversi variabel numeric.")
        row("Hubungan kira-kira linear/monoton", "ℹ️ Perlu grafik", "Buka Visualisasi → Scatter plot sebelum menafsirkan korelasi.")
        row("Tidak didominasi outlier", "ℹ️ Perlu cek", "Gunakan boxplot/scatter; jika banyak outlier, pertimbangkan Spearman.")
    elif analysis_name == "Regresi Linear":
        row("Ada target numeric", "✅ Ada" if num else "❌ Belum ada", "Pilih kolom hasil/target yang numeric.")
        row("Ada prediktor", "✅ Ada" if len(df.columns) >= 2 else "❌ Belum ada", "Tambahkan variabel prediktor teoritis.")
        row("Multikolinearitas terkendali", "ℹ️ Perlu VIF", "Setelah regresi, cek VIF; VIF tinggi berarti prediktor saling tumpang tindih.")
    elif analysis_name == "Reliabilitas / Cronbach Alpha":
        row("Minimal 3 item numeric", "✅ Siap" if len(num) >= 3 else "❌ Belum siap", "Tambahkan minimal 3 item skala yang mengukur konstruk sama.")
        row("Arah item sama", "ℹ️ Perlu cek", "Reverse coding item negatif sebelum menghitung alpha.")
        row("Item satu konstruk", "ℹ️ Perlu teori", "Jangan campur item dari dimensi berbeda dalam satu alpha kecuali memang satu skala.")
    elif analysis_name == "EFA":
        row("Minimal 3 item numeric", "✅ Ada" if len(num) >= 3 else "❌ Belum ada", "Tambahkan item numeric yang akan difaktorkan.")
        row("Sampel memadai", "✅ Cukup" if len(df) >= max(30, len(num)*5) else "⚠️ Terbatas", "Ideal: N ≥ 5×jumlah item, lebih baik ≥100.")
        row("Korelasi antar item cukup", "ℹ️ Cek KMO/Bartlett", "Jika KMO <0.60, revisi/hapus item yang kurang berkorelasi.")
    else:
        row("Pilih analisis", "ℹ️ Info", "Gunakan tabel saran analisis untuk melihat uji yang cocok dengan dataset.")
    return pd.DataFrame(rows)



# -----------------------------------------------------------------------------
# Smart Statistical Assistant v3.9
# -----------------------------------------------------------------------------
def sanitize_column_name(name):
    """Buat nama kolom ramah formula/syntax: huruf kecil, underscore, unik ditangani di helper lain."""
    value = str(name).strip()
    value = re.sub(r"\s+", "_", value)
    value = re.sub(r"[^0-9A-Za-z_]+", "", value)
    value = re.sub(r"_+", "_", value).strip("_")
    if not value:
        value = "var"
    if re.match(r"^\d", value):
        value = "v_" + value
    return value.lower()


def make_unique_columns(columns):
    seen = {}
    result = []
    for col in columns:
        base = sanitize_column_name(col)
        seen[base] = seen.get(base, 0) + 1
        result.append(base if seen[base] == 1 else f"{base}_{seen[base]}")
    return result


def detect_repair_actions(df):
    actions = []
    if any(str(c).strip() != sanitize_column_name(c) for c in df.columns) or len(set(map(str, df.columns))) != len(df.columns):
        actions.append({
            "Kode": "clean_names",
            "Masalah": "Nama kolom belum ramah analisis atau ada potensi duplikat",
            "Dampak": "Rumus/regresi/syntax lebih mudah error bila nama kolom berisi spasi, simbol, atau duplikat.",
            "Tindakan": "Rapikan nama kolom menjadi huruf kecil + underscore dan pastikan unik.",
        })
    num_like = []
    for col in df.columns:
        ok, ratio = _is_numeric_like(df[col])
        if ok:
            num_like.append(f"{col} ({ratio*100:.0f}% terlihat angka)")
    if num_like:
        actions.append({
            "Kode": "convert_numeric_like",
            "Masalah": "Kolom angka masih terbaca sebagai teks",
            "Dampak": "Kolom tidak muncul untuk t-test, ANOVA, korelasi, regresi, reliabilitas, dan EFA.",
            "Tindakan": "Konversi kolom yang terlihat numerik menjadi numeric.",
        })
    placeholder_tokens = ["-", "--", "?", "NA", "N/A", "na", "n/a", "null", "None", "missing", "tidak ada", ""]
    placeholder_cols = []
    for col in df.columns:
        if df[col].dtype == object:
            count = int(df[col].astype(str).str.strip().isin(placeholder_tokens).sum())
            if count:
                placeholder_cols.append(f"{col} ({count})")
    if placeholder_cols:
        actions.append({
            "Kode": "replace_missing_tokens",
            "Masalah": "Ada kode missing yang masih terbaca sebagai nilai biasa",
            "Dampak": "Frekuensi kategori bisa bias dan konversi numerik bisa gagal.",
            "Tindakan": "Ubah kode missing umum seperti '-', '?', 'NA', 'null' menjadi NaN.",
        })
    constant_cols = [c for c in df.columns if df[c].nunique(dropna=True) <= 1]
    if constant_cols:
        actions.append({
            "Kode": "drop_constant",
            "Masalah": "Ada kolom tanpa variasi",
            "Dampak": "Kolom konstan tidak berguna untuk korelasi, regresi, EFA, atau uji beda.",
            "Tindakan": "Hapus kolom konstan bila bukan metadata penting.",
        })
    all_missing = [c for c in df.columns if df[c].isna().all()]
    if all_missing:
        actions.append({
            "Kode": "drop_empty",
            "Masalah": "Ada kolom kosong total",
            "Dampak": "Membingungkan user dan bisa mengganggu pilihan variabel.",
            "Tindakan": "Hapus kolom yang seluruh nilainya kosong.",
        })
    text_cols_with_spaces = []
    for col in df.columns:
        if df[col].dtype == object:
            before = df[col].astype(str)
            after = before.str.strip().str.replace(r"\s+", " ", regex=True)
            if int((before != after).sum()) > 0:
                text_cols_with_spaces.append(col)
    if text_cols_with_spaces:
        actions.append({
            "Kode": "trim_text_spaces",
            "Masalah": "Ada kategori/teks dengan spasi berlebih",
            "Dampak": "Kategori seperti 'A' dan ' A ' bisa dianggap berbeda.",
            "Tindakan": "Rapikan spasi pada kolom teks/kategori.",
        })
    if df.isna().any().any():
        actions.append({
            "Kode": "impute_basic",
            "Masalah": "Ada missing value",
            "Dampak": "Beberapa analisis membuang baris kosong sehingga sampel efektif mengecil.",
            "Tindakan": "Isi missing sederhana: numerik=median, kategori=modus.",
        })
    outlier_cols = []
    for col in df.select_dtypes(include=[np.number]).columns:
        s = df[col].dropna()
        if len(s) >= 8:
            q1, q3 = s.quantile([0.25, 0.75])
            iqr = q3 - q1
            if iqr > 0:
                count = int(((s < q1 - 3*iqr) | (s > q3 + 3*iqr)).sum())
                if count:
                    outlier_cols.append(f"{col} ({count})")
    if outlier_cols:
        actions.append({
            "Kode": "winsorize_extreme_outliers",
            "Masalah": "Ada outlier numerik ekstrem",
            "Dampak": "Mean, korelasi, regresi, dan ANOVA bisa sangat dipengaruhi nilai ekstrem.",
            "Tindakan": "Winsorize outlier ekstrem ke batas 3×IQR.",
        })
    return pd.DataFrame(actions)


def apply_repair_action(df, action_code):
    repaired = df.copy()
    note = ""
    if action_code == "clean_names":
        old_cols = list(repaired.columns)
        repaired.columns = make_unique_columns(old_cols)
        note = "Nama kolom dirapikan menjadi format aman: huruf kecil, underscore, dan unik."
    elif action_code == "replace_missing_tokens":
        tokens = ["-", "--", "?", "NA", "N/A", "na", "n/a", "null", "None", "missing", "tidak ada", ""]
        repaired = repaired.replace(tokens, np.nan)
        note = "Kode missing umum diganti menjadi NaN."
    elif action_code == "convert_numeric_like":
        converted_cols = []
        for col in repaired.columns:
            ok, _ = _is_numeric_like(repaired[col])
            if ok:
                cleaned = repaired[col].astype(str).str.strip()
                # Jika ada koma sebagai desimal dan titik sebagai ribuan, pola ini umumnya berhasil untuk data Indonesia.
                cleaned = cleaned.str.replace(".", "", regex=False).str.replace(",", ".", regex=False)
                cleaned = cleaned.str.replace(r"[^0-9\-\.]+", "", regex=True)
                repaired[col] = pd.to_numeric(cleaned.replace("", np.nan), errors="coerce")
                converted_cols.append(col)
        note = "Kolom terlihat numerik dikonversi: " + (", ".join(map(str, converted_cols)) if converted_cols else "tidak ada kolom yang perlu dikonversi")
    elif action_code == "drop_constant":
        cols = [c for c in repaired.columns if repaired[c].nunique(dropna=True) <= 1]
        repaired = repaired.drop(columns=cols)
        note = "Kolom tanpa variasi dihapus: " + (", ".join(map(str, cols)) if cols else "tidak ada")
    elif action_code == "drop_empty":
        cols = [c for c in repaired.columns if repaired[c].isna().all()]
        repaired = repaired.drop(columns=cols)
        note = "Kolom kosong total dihapus: " + (", ".join(map(str, cols)) if cols else "tidak ada")
    elif action_code == "trim_text_spaces":
        cols = []
        for col in repaired.columns:
            if repaired[col].dtype == object:
                repaired[col] = repaired[col].astype(str).str.strip().str.replace(r"\s+", " ", regex=True).replace({"nan": np.nan, "None": np.nan})
                cols.append(col)
        note = "Spasi berlebih pada teks/kategori dirapikan: " + (", ".join(map(str, cols)) if cols else "tidak ada")
    elif action_code == "impute_basic":
        changes = []
        for col in repaired.columns:
            miss = int(repaired[col].isna().sum())
            if miss == 0:
                continue
            if pd.api.types.is_numeric_dtype(repaired[col]):
                fill = repaired[col].median()
                repaired[col] = repaired[col].fillna(fill)
                changes.append(f"{col}=median")
            else:
                mode = repaired[col].mode(dropna=True)
                if not mode.empty:
                    repaired[col] = repaired[col].fillna(mode.iloc[0])
                    changes.append(f"{col}=modus")
        note = "Missing value diisi sederhana (gunakan hanya bila sesuai desain riset): " + (", ".join(changes) if changes else "tidak ada")
    elif action_code == "winsorize_extreme_outliers":
        changes = []
        for col in repaired.select_dtypes(include=[np.number]).columns:
            s = repaired[col].dropna()
            if len(s) < 8:
                continue
            q1, q3 = s.quantile([0.25, 0.75])
            iqr = q3 - q1
            if iqr <= 0:
                continue
            lo, hi = q1 - 3*iqr, q3 + 3*iqr
            before = repaired[col].copy()
            repaired[col] = repaired[col].clip(lo, hi)
            changed = int((before != repaired[col]).sum())
            if changed:
                changes.append(f"{col} ({changed})")
        note = "Outlier ekstrem di-winsorize ke batas 3×IQR: " + (", ".join(changes) if changes else "tidak ada")
    else:
        note = "Tidak ada tindakan yang diterapkan."
    return repaired, note


def classify_variable(df, col):
    if not col or col not in df.columns:
        return "Tidak dipilih"
    s = df[col].dropna()
    if s.empty:
        return "Kosong"
    if pd.api.types.is_numeric_dtype(df[col]):
        if s.nunique() <= 2:
            return "Numeric biner/kode kategori"
        if s.nunique() <= 10 and np.allclose(pd.to_numeric(s, errors="coerce").dropna() % 1, 0):
            return "Ordinal/numeric diskrit"
        return "Scale/numeric kontinu"
    unique = s.nunique()
    if unique == 2:
        return "Kategori 2 kelompok"
    if unique <= 10:
        return "Kategori multi-kelompok"
    return "Teks/kategori banyak"


def smart_test_recommendation(df, objective, y=None, x=None, group=None, paired_y=None):
    alpha_note = "Selalu baca bersama effect size, confidence interval, dan asumsi data."
    rows = []
    def add(priority, test, condition, alternative, why, next_step):
        rows.append({
            "Prioritas": priority,
            "Uji/Analisis Disarankan": test,
            "Kondisi Data": condition,
            "Alternatif Jika Asumsi Lemah": alternative,
            "Mengapa Cocok": why,
            "Langkah Berikutnya": next_step,
        })

    y_type = classify_variable(df, y)
    x_type = classify_variable(df, x)
    g_type = classify_variable(df, group)
    num = numeric_cols(df)
    cat = categorical_cols(df)

    if objective == "Membandingkan rata-rata antar kelompok":
        if y in num and group in df.columns:
            k = df[group].dropna().nunique()
            if k == 2:
                add("Utama", "Independent Samples T-Test", f"Y={y_type}; Grup={g_type}", "Mann-Whitney U", "Satu skor numerik dibandingkan pada dua kelompok independen.", "Cek normalitas per grup, homogenitas varians, lalu jalankan t-test.")
            elif k >= 3:
                add("Utama", "One-Way ANOVA", f"Y={y_type}; Grup={k} kategori", "Kruskal-Wallis + Dunn post-hoc", "Satu skor numerik dibandingkan pada tiga kelompok atau lebih.", "Cek Levene, normalitas residual, lalu lanjut post-hoc bila signifikan.")
            else:
                add("Perlu data", "Belum bisa ditentukan", "Variabel grup kurang dari 2 kategori", "-", "Uji beda butuh minimal dua kelompok.", "Tambahkan/recode grup agar memiliki 2+ kategori.")
        elif len(num) >= 2 and not group:
            add("Alternatif", "Paired T-Test / Repeated comparison format wide", "Ada ≥2 kolom numerik", "Wilcoxon Signed-Rank", "Jika dua kolom adalah before-after orang yang sama, gunakan paired t-test.", "Pilih kolom before dan after di Uji Statistik.")
        else:
            add("Perlu data", "Uji beda belum siap", "Butuh Y numerik dan grup kategori", "-", "Format data belum cukup untuk membandingkan rata-rata.", "Tambahkan kolom skor dan kolom kelompok.")
    elif objective == "Melihat hubungan antar variabel":
        if y in num and x in num:
            add("Utama", "Korelasi Pearson", f"X={x_type}; Y={y_type}", "Spearman/Kendall", "Dua variabel numerik dapat diuji hubungan linear/monoton.", "Buat scatter plot, cek outlier, lalu jalankan korelasi.")
        elif x in df.columns and y in df.columns:
            add("Alternatif", "Crosstab/Chi-Square atau Korelasi Spearman", f"X={x_type}; Y={y_type}", "Fisher Exact untuk tabel kecil", "Jenis variabel tidak sama-sama kontinu sehingga perlu uji kategori/ordinal.", "Pastikan measurement level benar di Variable View.")
        else:
            add("Perlu data", "Korelasi belum siap", "Butuh minimal dua variabel", "-", "Hubungan butuh pasangan variabel yang jelas.", "Pilih dua kolom skor atau dua kategori yang ingin dikaitkan.")
    elif objective == "Memprediksi variabel hasil":
        if y in num and (x in df.columns or group in df.columns):
            add("Utama", "Regresi Linear", f"Y={y_type}; Prediktor={x_type if x else g_type}", "Regresi robust/transformasi jika asumsi lemah", "Target numerik dapat diprediksi dari satu atau beberapa prediktor.", "Jalankan regresi, cek R², koefisien, VIF, residual, dan outlier.")
        elif y in df.columns and df[y].dropna().nunique() == 2:
            add("Utama", "Regresi Logistik", f"Y={y_type}; kategori biner", "Chi-square/crosstab untuk eksplorasi", "Target biner cocok dianalisis dengan odds ratio.", "Pastikan kategori target dikodekan 0/1 atau dua label konsisten.")
        else:
            add("Perlu data", "Regresi belum siap", "Butuh target Y dan prediktor", "-", "Prediksi membutuhkan variabel hasil yang jelas.", "Pilih target numerik/biner dan beberapa prediktor teoritis.")
    elif objective == "Menguji kuesioner/skala":
        n_num = len(num)
        if n_num >= 3:
            add("Utama", "Reliabilitas Cronbach's Alpha", f"Ada {n_num} item numerik", "McDonald's Omega / item-total correlation", "Item numerik dapat diuji konsistensi internalnya.", "Pastikan item satu konstruk dan reverse coding item negatif.")
            add("Lanjutan", "EFA / Analisis Faktor", f"Ada {n_num} item; N={len(df)}", "PCA eksploratori jika sampel kecil", "EFA membantu menemukan dimensi/faktor laten.", "Cek KMO ≥0.60, Bartlett signifikan, loading, cross-loading, dan communality.")
        else:
            add("Perlu data", "Reliabilitas/EFA belum siap", "Butuh minimal 3 item numerik", "-", "Skala/kuesioner butuh beberapa item yang mengukur konstruk sama.", "Tambahkan/konversi item Likert menjadi numeric.")
    elif objective == "Menganalisis data kategori":
        if len(cat) >= 2 or (x in df.columns and y in df.columns):
            add("Utama", "Crosstab & Chi-Square", f"X={x_type}; Y={y_type}", "Fisher Exact bila expected count kecil", "Dua variabel kategori dapat diuji asosiasinya.", "Cek persentase baris/kolom dan Cramer's V.")
        else:
            add("Perlu data", "Chi-Square belum siap", "Butuh dua variabel kategori", "-", "Data kategori butuh dua kolom kategori yang jelas.", "Tambahkan variabel kategori seperti gender, kelas, status, pilihan.")
    else:
        add("Info", "Eksplorasi Deskriptif", "Tujuan belum spesifik", "-", "Deskriptif adalah langkah awal semua analisis.", alpha_note)

    return pd.DataFrame(rows)


def compute_quick_effect_sizes(df, y=None, x=None, group=None):
    rows = []
    if y in numeric_cols(df) and group in df.columns and df[group].dropna().nunique() == 2:
        tmp = df[[y, group]].dropna()
        levels = tmp[group].dropna().unique().tolist()
        if len(levels) == 2:
            a = tmp[tmp[group] == levels[0]][y].astype(float)
            b = tmp[tmp[group] == levels[1]][y].astype(float)
            if len(a) >= 2 and len(b) >= 2:
                pooled = np.sqrt(((len(a)-1)*a.var(ddof=1) + (len(b)-1)*b.var(ddof=1)) / max(len(a)+len(b)-2, 1))
                d = (a.mean() - b.mean()) / pooled if pooled else np.nan
                correction = 1 - (3 / (4*(len(a)+len(b))-9)) if (len(a)+len(b)) > 2 else 1
                rows.append({"Ukuran Efek": "Cohen's d", "Nilai": d, "Interpretasi": effect_size_plain_label(abs(d), "d")})
                rows.append({"Ukuran Efek": "Hedges' g", "Nilai": d * correction, "Interpretasi": effect_size_plain_label(abs(d * correction), "d")})
    if y in numeric_cols(df) and x in numeric_cols(df):
        tmp = df[[x, y]].dropna()
        if len(tmp) >= 3:
            r, p = stats.pearsonr(tmp[x], tmp[y])
            rows.append({"Ukuran Efek": "Pearson r", "Nilai": r, "Interpretasi": effect_size_plain_label(abs(r), "r")})
            rows.append({"Ukuran Efek": "r²", "Nilai": r*r, "Interpretasi": f"Sekitar {r*r*100:.1f}% variasi bersama secara linear"})
    if y in numeric_cols(df) and group in df.columns and df[group].dropna().nunique() >= 3:
        tmp = df[[y, group]].dropna()
        groups = [g[y].astype(float).values for _, g in tmp.groupby(group)]
        if len(groups) >= 3 and all(len(g) >= 2 for g in groups):
            all_vals = np.concatenate(groups)
            grand = all_vals.mean()
            ss_between = sum(len(g) * (g.mean() - grand) ** 2 for g in groups)
            ss_total = sum((v - grand) ** 2 for v in all_vals)
            eta = ss_between / ss_total if ss_total else np.nan
            rows.append({"Ukuran Efek": "Eta squared (η²)", "Nilai": eta, "Interpretasi": effect_size_plain_label(eta, "eta")})
    return pd.DataFrame(rows)


def effect_size_plain_label(value, kind):
    try:
        v = abs(float(value))
    except Exception:
        return "Tidak dapat dinilai"
    if kind == "d":
        if v < 0.2: return "sangat kecil"
        if v < 0.5: return "kecil"
        if v < 0.8: return "sedang"
        return "besar"
    if kind == "r":
        if v < 0.1: return "sangat kecil"
        if v < 0.3: return "kecil"
        if v < 0.5: return "sedang"
        return "besar"
    if kind == "eta":
        if v < 0.01: return "sangat kecil"
        if v < 0.06: return "kecil"
        if v < 0.14: return "sedang"
        return "besar"
    return "Perlu konteks"


def normal_power_sample_size(test_type, alpha, power, effect_size, groups=2, predictors=1, r=None, moe=None, proportion=0.5):
    """Kalkulator sample size ringan. Memakai statsmodels bila tersedia, fallback ke pendekatan normal."""
    z_alpha = stats.norm.ppf(1 - alpha / 2)
    z_power = stats.norm.ppf(power)
    effect_size = float(effect_size) if effect_size else np.nan
    try:
        from statsmodels.stats.power import TTestPower, TTestIndPower, FTestAnovaPower
        if test_type == "Independent t-test" and effect_size > 0:
            n = TTestIndPower().solve_power(effect_size=effect_size, alpha=alpha, power=power, ratio=1.0, alternative="two-sided")
            return int(np.ceil(n)), "per kelompok", "Dihitung dengan statsmodels TTestIndPower."
        if test_type in ["One-sample t-test", "Paired t-test"] and effect_size > 0:
            n = TTestPower().solve_power(effect_size=effect_size, alpha=alpha, power=power, alternative="two-sided")
            return int(np.ceil(n)), "total", "Dihitung dengan statsmodels TTestPower."
        if test_type == "One-way ANOVA" and effect_size > 0:
            n = FTestAnovaPower().solve_power(effect_size=effect_size, k_groups=max(2, int(groups)), alpha=alpha, power=power)
            return int(np.ceil(n)), "total", "Effect size ANOVA memakai Cohen's f."
    except Exception:
        pass

    if test_type == "Independent t-test" and effect_size > 0:
        n = 2 * ((z_alpha + z_power) / effect_size) ** 2
        return int(np.ceil(n)), "per kelompok (aproksimasi)", "Fallback normal approximation."
    if test_type in ["One-sample t-test", "Paired t-test"] and effect_size > 0:
        n = ((z_alpha + z_power) / effect_size) ** 2
        return int(np.ceil(n)), "total (aproksimasi)", "Fallback normal approximation."
    if test_type == "Correlation" and r and abs(float(r)) > 0 and abs(float(r)) < 1:
        fisher_z = np.arctanh(abs(float(r)))
        n = ((z_alpha + z_power) / fisher_z) ** 2 + 3
        return int(np.ceil(n)), "pasangan data", "Aproksimasi Fisher z untuk korelasi."
    if test_type == "Survey proportion / margin of error" and moe and float(moe) > 0:
        n = (z_alpha ** 2) * float(proportion) * (1 - float(proportion)) / (float(moe) ** 2)
        return int(np.ceil(n)), "responden", "Rumus proporsi sederhana tanpa finite population correction."
    if test_type == "Multiple regression":
        m = max(1, int(predictors))
        n_model = 50 + 8*m
        n_predictor = 104 + m
        return int(max(n_model, n_predictor)), "responden (aturan praktis)", "Aturan praktis Green: N ≥ 50+8m dan 104+m."
    return None, "", "Masukkan effect size/parameter yang valid."


def parse_number_list(text_value):
    vals = re.split(r"[\s,;]+", str(text_value).strip())
    out = []
    for v in vals:
        if not v:
            continue
        try:
            out.append(float(v.replace(",", ".")))
        except Exception:
            pass
    return pd.Series(out, dtype="float64")


def descriptive_calculator_table(values):
    s = pd.Series(values).dropna().astype(float)
    if s.empty:
        return pd.DataFrame([{"Statistik": "Error", "Nilai": "Tidak ada angka valid"}])
    q1, q3 = s.quantile(0.25), s.quantile(0.75)
    mode_vals = s.mode().tolist()
    return pd.DataFrame([
        {"Statistik": "N", "Nilai": len(s)},
        {"Statistik": "Mean", "Nilai": s.mean()},
        {"Statistik": "Median", "Nilai": s.median()},
        {"Statistik": "Modus", "Nilai": ", ".join(f"{x:g}" for x in mode_vals[:5]) if mode_vals else "-"},
        {"Statistik": "Std. Deviasi", "Nilai": s.std(ddof=1) if len(s) > 1 else np.nan},
        {"Statistik": "Varians", "Nilai": s.var(ddof=1) if len(s) > 1 else np.nan},
        {"Statistik": "Minimum", "Nilai": s.min()},
        {"Statistik": "Q1", "Nilai": q1},
        {"Statistik": "Q3", "Nilai": q3},
        {"Statistik": "IQR", "Nilai": q3 - q1},
        {"Statistik": "Maximum", "Nilai": s.max()},
        {"Statistik": "Standard Error", "Nilai": stats.sem(s) if len(s) > 1 else np.nan},
        {"Statistik": "CI 95% Mean", "Nilai": f"{stats.t.interval(0.95, len(s)-1, loc=s.mean(), scale=stats.sem(s))[0]:.4f} s.d. {stats.t.interval(0.95, len(s)-1, loc=s.mean(), scale=stats.sem(s))[1]:.4f}" if len(s) > 1 and s.std(ddof=1) > 0 else "-"},
    ])


def distribution_calculator(dist_name, mode, value, df1=None, df2=None, n=None, p=None, lam=None, alpha=0.05):
    if dist_name == "Normal/Z":
        if mode == "P(X ≤ nilai)":
            return stats.norm.cdf(value)
        if mode == "P(X ≥ nilai)":
            return 1 - stats.norm.cdf(value)
        return stats.norm.ppf(1 - alpha/2)
    if dist_name == "t":
        d = int(df1 or 1)
        if mode == "P(X ≤ nilai)":
            return stats.t.cdf(value, d)
        if mode == "P(X ≥ nilai)":
            return 1 - stats.t.cdf(value, d)
        return stats.t.ppf(1 - alpha/2, d)
    if dist_name == "Chi-square":
        d = int(df1 or 1)
        if mode == "P(X ≤ nilai)":
            return stats.chi2.cdf(value, d)
        if mode == "P(X ≥ nilai)":
            return 1 - stats.chi2.cdf(value, d)
        return stats.chi2.ppf(1 - alpha, d)
    if dist_name == "F":
        d1, d2 = int(df1 or 1), int(df2 or 1)
        if mode == "P(X ≤ nilai)":
            return stats.f.cdf(value, d1, d2)
        if mode == "P(X ≥ nilai)":
            return 1 - stats.f.cdf(value, d1, d2)
        return stats.f.ppf(1 - alpha, d1, d2)
    if dist_name == "Binomial":
        nn, pp = int(n or 1), float(p or 0.5)
        if mode == "P(X ≤ nilai)":
            return stats.binom.cdf(int(value), nn, pp)
        if mode == "P(X ≥ nilai)":
            return 1 - stats.binom.cdf(int(value)-1, nn, pp)
        return stats.binom.ppf(1 - alpha, nn, pp)
    if dist_name == "Poisson":
        ll = float(lam or 1)
        if mode == "P(X ≤ nilai)":
            return stats.poisson.cdf(int(value), ll)
        if mode == "P(X ≥ nilai)":
            return 1 - stats.poisson.cdf(int(value)-1, ll)
        return stats.poisson.ppf(1 - alpha, ll)
    return np.nan


def build_report_template(kind, research_title, hypothesis, selected_output_titles):
    title = research_title.strip() or "[Judul Penelitian]"
    hyp = hypothesis.strip() or "[Hipotesis penelitian]"
    output_text = ", ".join(selected_output_titles) if selected_output_titles else "[output analisis yang dipilih]"
    if kind == "BAB 4 Skripsi/Tesis":
        return f"""### Template Narasi BAB 4

Penelitian berjudul **{title}** bertujuan untuk menguji: **{hyp}**. Analisis dilakukan menggunakan output: **{output_text}**.

Berdasarkan hasil analisis, peneliti perlu menyajikan tiga hal utama: (1) gambaran deskriptif data, (2) hasil uji asumsi atau kelayakan analisis, dan (3) hasil uji hipotesis. Nilai signifikansi tidak boleh dibaca sendirian; hasil perlu dilengkapi dengan arah hubungan/perbedaan, ukuran efek, interval kepercayaan, serta relevansinya terhadap teori.

Jika hasil signifikan, narasi pembahasan dapat diarahkan pada dukungan empiris terhadap hipotesis. Jika hasil tidak signifikan, pembahasan perlu menekankan bahwa bukti statistik pada data ini belum cukup kuat, bukan berarti hubungan/perbedaan pasti tidak ada. Pertimbangkan ukuran sampel, kualitas instrumen, variasi data, dan kesesuaian model.

**Kalimat siap pakai:**
> Berdasarkan hasil analisis, diperoleh temuan bahwa [isi temuan utama]. Temuan ini menunjukkan bahwa [makna substantif]. Dengan demikian, hasil penelitian ini [mendukung/tidak cukup mendukung] hipotesis yang diajukan, dengan tetap memperhatikan ukuran efek, asumsi statistik, dan keterbatasan data.
"""
    if kind == "APA Style":
        return f"""### Template APA Style

Study: **{title}**  
Hypothesis: **{hyp}**  
Selected output: **{output_text}**

Use this structure:

> A statistical analysis was conducted to examine {hyp}. Report the descriptive statistics first, followed by assumption checks and the main inferential test. Include the test statistic, degrees of freedom when available, *p*-value, confidence interval, and effect size.

Examples:

- Independent t-test: *t*(df) = value, *p* = value, 95% CI [LL, UL], Cohen's *d* = value.
- ANOVA: *F*(df1, df2) = value, *p* = value, η² = value.
- Correlation: *r*(df) = value, *p* = value, 95% CI [LL, UL].
- Regression: β = value, *t* = value, *p* = value, R² = value.
"""
    return f"""### Template Ringkasan Manajerial

**Riset:** {title}  
**Pertanyaan/Hipotesis:** {hyp}  
**Output yang digunakan:** {output_text}

**Inti temuan:**  
Tuliskan 1–3 temuan paling penting dari output statistik.

**Makna praktis:**  
Jelaskan apa arti temuan tersebut untuk keputusan, kebijakan, pembelajaran, layanan, atau pengembangan program.

**Risiko interpretasi:**  
Sebutkan keterbatasan seperti ukuran sampel, missing value, asumsi yang tidak terpenuhi, atau effect size kecil.

**Rekomendasi:**  
Berikan langkah berikutnya: tambah data, perbaiki instrumen, lakukan post-hoc, uji model alternatif, atau validasi pada sampel lain.
"""



# -----------------------------------------------------------------------------
# Guided UI & comprehensive guidance v4.0
# -----------------------------------------------------------------------------
def ui_mode_is_beginner():
    return st.session_state.get("ui_mode", "Pemula") == "Pemula"


def detail_is_full():
    return st.session_state.get("detail_level", "Ringkas") == "Lengkap"


def card_html(title, body, icon="", tone="soft"):
    klass = "guide-card" if tone == "guide" else "soft-card"
    return f"""<div class='{klass}'><b>{icon} {title}</b><br><span class='tiny'>{body}</span></div>"""


def status_color_label(score):
    if score >= 85:
        return "Sangat siap", "ok-box"
    if score >= 70:
        return "Cukup siap", "ok-box"
    if score >= 50:
        return "Perlu perbaikan ringan", "warn-box"
    return "Perlu dibersihkan dulu", "danger-box"


def variable_role_suggestions(df):
    rows = []
    n = len(df)
    for col in df.columns:
        s = df[col]
        nonmiss = int(s.notna().sum())
        miss_pct = float(s.isna().mean() * 100)
        unique = int(s.nunique(dropna=True))
        var_type = classify_variable(df, col)
        role = []
        avoid = []
        if pd.api.types.is_numeric_dtype(s):
            if unique <= 1:
                role.append("Tidak disarankan untuk analisis")
                avoid.append("Kolom tidak punya variasi")
            elif unique == n and str(col).lower() in ["id", "kode", "no", "nomor", "responden", "nama"]:
                role.append("ID/identitas")
                avoid.append("Jangan dipakai sebagai variabel statistik")
            elif unique <= 10 and np.allclose(pd.to_numeric(s.dropna(), errors="coerce") % 1, 0):
                role.append("Ordinal/kode kategori")
                role.append("Grup jika maknanya kategori")
                avoid.append("Jangan dianggap skala jika sebenarnya kode")
            else:
                role.append("Variabel hasil/skor")
                role.append("Prediktor numerik")
                role.append("Item skala/kuesioner bila berasal dari Likert")
        else:
            if unique <= 1:
                role.append("Tidak disarankan")
                avoid.append("Tidak ada variasi")
            elif unique == 2:
                role.append("Grup 2 kategori")
                role.append("Target logistik biner")
            elif unique <= 10:
                role.append("Grup/kategori")
                role.append("Faktor ANOVA/Chi-square")
            else:
                role.append("Label/teks bebas")
                avoid.append("Recode/ringkas dulu bila ingin dianalisis")
        if miss_pct > 30:
            avoid.append("Missing value tinggi")
        rows.append({
            "Kolom": col,
            "Tipe terbaca": var_type,
            "Isi valid": f"{nonmiss}/{n}",
            "Missing %": round(miss_pct, 2),
            "Kategori/nilai unik": unique,
            "Peran yang cocok": "; ".join(role),
            "Perlu hati-hati": "; ".join(avoid) if avoid else "-",
        })
    return pd.DataFrame(rows)


def _issues_to_records(issues):
    """Normalize compatibility issues to list[dict].

    Beberapa menu mengirim issues sebagai DataFrame, sementara helper lain
    bisa mengirim list of dict atau list of strings. Fungsi ini membuat
    semua bentuk tersebut aman dibaca agar UI tidak jatuh hanya karena
    format diagnosis berbeda.
    """
    if issues is None:
        return []
    if isinstance(issues, pd.DataFrame):
        if issues.empty:
            return []
        return issues.fillna("").to_dict("records")
    if isinstance(issues, dict):
        return [issues]
    records = []
    if isinstance(issues, (list, tuple, set)):
        for item in issues:
            if isinstance(item, dict):
                records.append(item)
            else:
                records.append({
                    "Prioritas": "Info",
                    "Masalah": str(item),
                    "Area": "Catatan",
                    "Solusi": "Tinjau kembali diagnosis kompatibilitas data.",
                })
        return records
    return [{
        "Prioritas": "Info",
        "Masalah": str(issues),
        "Area": "Catatan",
        "Solusi": "Tinjau kembali diagnosis kompatibilitas data.",
    }]


def build_next_best_actions(df, issues, score):
    num = numeric_cols(df)
    cat = categorical_cols(df)
    issue_records = _issues_to_records(issues)
    actions = []
    critical = [i for i in issue_records if str(i.get("Prioritas", "")).lower() == "kritis"]
    high = [i for i in issue_records if str(i.get("Prioritas", "")).lower() == "tinggi"]
    if critical:
        actions.append({"Urutan": 1, "Langkah": "Perbaiki masalah kritis", "Menu": "🧰 Kompatibilitas Data", "Kenapa": critical[0].get("Masalah", "Ada masalah kritis pada data")})
    elif high:
        actions.append({"Urutan": 1, "Langkah": "Bersihkan format data", "Menu": "🧙 Smart Assistant → Data Repair Assistant", "Kenapa": high[0].get("Masalah", "Ada masalah data prioritas tinggi")})
    elif score < 70:
        actions.append({"Urutan": 1, "Langkah": "Tinjau kualitas data", "Menu": "🧰 Kompatibilitas Data", "Kenapa": "Skor kesiapan data masih perlu ditingkatkan sebelum analisis utama"})
    else:
        actions.append({"Urutan": 1, "Langkah": "Lihat profil dan pola awal", "Menu": "📋 Deskriptif", "Kenapa": "Data cukup siap; mulai dari ringkasan dan frekuensi"})
    if len(num) >= 2:
        actions.append({"Urutan": 2, "Langkah": "Cek hubungan dan grafik", "Menu": "🎨 Visualisasi + 🧪 Uji Statistik", "Kenapa": "Ada minimal dua variabel numerik untuk korelasi/scatter"})
    elif num and cat:
        actions.append({"Urutan": 2, "Langkah": "Bandingkan skor antar kelompok", "Menu": "🧙 Smart Assistant → Wizard Uji Otomatis", "Kenapa": "Ada skor numerik dan variabel grup/kategori"})
    else:
        actions.append({"Urutan": 2, "Langkah": "Atur tipe variabel", "Menu": "🗂️ Data → Variable View", "Kenapa": "Variabel numerik/kategori belum cukup jelas"})
    actions.append({"Urutan": 3, "Langkah": "Jalankan analisis utama", "Menu": "🧪 Uji Statistik / 📈 Regresi / 🧭 Reliabilitas & Faktor", "Kenapa": "Pilih sesuai pertanyaan riset"})
    actions.append({"Urutan": 4, "Langkah": "Maknai dan ekspor laporan", "Menu": "🧠 Insight Riset + 📤 Output & Ekspor", "Kenapa": "Hasil perlu diterjemahkan menjadi kesimpulan riset"})
    return pd.DataFrame(actions)


def analysis_decision_matrix(df):
    num = numeric_cols(df)
    cat = categorical_cols(df)
    rows = [
        {"Pertanyaan user awam": "Apakah satu nilai rata-rata berbeda dari target tertentu?", "Data yang dibutuhkan": "1 variabel numerik", "Uji utama": "One-Sample T-Test", "Alternatif aman": "Wilcoxon one-sample/cek CI", "Status dataset": "Siap" if len(num) >= 1 else "Belum siap"},
        {"Pertanyaan user awam": "Apakah dua kelompok punya rata-rata berbeda?", "Data yang dibutuhkan": "1 variabel skor + 1 grup dua kategori", "Uji utama": "Independent T-Test", "Alternatif aman": "Mann-Whitney U", "Status dataset": "Siap" if num and any(df[c].dropna().nunique()==2 for c in df.columns) else "Perlu grup 2 kategori"},
        {"Pertanyaan user awam": "Apakah nilai sebelum-sesudah berubah?", "Data yang dibutuhkan": "2 kolom numerik berpasangan", "Uji utama": "Paired T-Test", "Alternatif aman": "Wilcoxon Signed-Rank", "Status dataset": "Siap" if len(num) >= 2 else "Perlu 2 kolom numerik"},
        {"Pertanyaan user awam": "Apakah 3+ kelompok berbeda?", "Data yang dibutuhkan": "Skor numerik + grup 3+ kategori", "Uji utama": "One-Way ANOVA", "Alternatif aman": "Kruskal-Wallis + Dunn", "Status dataset": "Siap" if num and any(3 <= df[c].dropna().nunique() <= 20 for c in df.columns) else "Perlu grup 3+ kategori"},
        {"Pertanyaan user awam": "Apakah dua variabel berhubungan?", "Data yang dibutuhkan": "2 variabel numerik/ordinal", "Uji utama": "Pearson/Spearman", "Alternatif aman": "Kendall/visual scatter", "Status dataset": "Siap" if len(num) >= 2 else "Perlu 2 variabel numerik"},
        {"Pertanyaan user awam": "Faktor apa yang memprediksi hasil?", "Data yang dibutuhkan": "Y numerik + X prediktor", "Uji utama": "Regresi Linear", "Alternatif aman": "Transformasi/robust/cek VIF", "Status dataset": "Siap" if len(num) >= 2 else "Perlu target dan prediktor"},
        {"Pertanyaan user awam": "Apakah dua kategori saling berkaitan?", "Data yang dibutuhkan": "2 variabel kategori", "Uji utama": "Chi-Square", "Alternatif aman": "Fisher Exact", "Status dataset": "Siap" if len(cat) >= 2 else "Perlu 2 kategori"},
        {"Pertanyaan user awam": "Apakah item kuesioner konsisten?", "Data yang dibutuhkan": "≥3 item numeric satu konstruk", "Uji utama": "Cronbach's Alpha", "Alternatif aman": "Item-total, Omega", "Status dataset": "Siap" if len(num) >= 3 else "Perlu ≥3 item"},
        {"Pertanyaan user awam": "Dimensi/faktor apa yang muncul dari item?", "Data yang dibutuhkan": "≥3 item numeric + sampel cukup", "Uji utama": "EFA/PAF", "Alternatif aman": "PCA eksploratori", "Status dataset": "Siap/eksploratif" if len(num) >= 3 else "Perlu ≥3 item"},
    ]
    return pd.DataFrame(rows)


def comprehensive_ui_feature_map():
    return pd.DataFrame([
        {"Area": "Data", "Fitur": "Import CSV/Excel/SPSS, Data View, Variable View", "Untuk pemula": "Mulai dari Upload/Data Contoh", "Untuk ahli": "Atur value labels, measure, missing values"},
        {"Area": "Kualitas data", "Fitur": "Compatibility Checker, Quality Score, Data Repair", "Untuk pemula": "Ikuti tombol saran perbaikan", "Untuk ahli": "Audit missing/outlier/tipe variabel"},
        {"Area": "Transformasi", "Fitur": "Compute, recode, reverse coding, z-score, filter, split", "Untuk pemula": "Pakai wizard/repair dulu", "Untuk ahli": "Gunakan Transform untuk workflow SPSS-like"},
        {"Area": "Analisis dasar", "Fitur": "Deskriptif, frekuensi, normalitas, crosstab", "Untuk pemula": "Baca ringkasan dan interpretasi", "Untuk ahli": "Cek asumsi dan distribusi"},
        {"Area": "Uji hipotesis", "Fitur": "T-test, ANOVA, nonparametrik, chi-square, post-hoc", "Untuk pemula": "Pilih lewat Wizard Uji Otomatis", "Untuk ahli": "Pilih langsung di menu Uji Statistik"},
        {"Area": "Model", "Fitur": "Regresi linear/logistik, VIF, residual diagnostics", "Untuk pemula": "Gunakan saat ingin memprediksi Y", "Untuk ahli": "Cek model fit, asumsi, multikolinearitas"},
        {"Area": "Instrumen", "Fitur": "Reliabilitas, PCA, EFA/PAF, KMO, Bartlett", "Untuk pemula": "Gunakan untuk item kuesioner", "Untuk ahli": "Tinjau loading, communality, variance"},
        {"Area": "Pelaporan", "Fitur": "Insight riset, template BAB 4/APA, ekspor Excel/Word/HTML/MD", "Untuk pemula": "Ambil narasi awal lalu sesuaikan teori", "Untuk ahli": "Gunakan output viewer dan syntax log"},
    ])


def research_design_planner_table(design, objective, sample_context):
    base = [
        {"Tahap": "1. Rumusan masalah", "Yang perlu diisi": "Apa fenomena/variabel utama?", "Contoh keputusan": objective or "Bandingkan/hubungkan/prediksi variabel"},
        {"Tahap": "2. Desain data", "Yang perlu diisi": "Cross-sectional, eksperimen, pre-post, longitudinal", "Contoh keputusan": design},
        {"Tahap": "3. Variabel", "Yang perlu diisi": "Tentukan Y, X, grup, kontrol, item skala", "Contoh keputusan": "Gunakan Variable View agar tipe/measure jelas"},
        {"Tahap": "4. Ukuran sampel", "Yang perlu diisi": "Target power, alpha, effect size", "Contoh keputusan": "Gunakan Sample Size & Power"},
        {"Tahap": "5. Kualitas data", "Yang perlu diisi": "Missing, outlier, coding, normalitas, homogenitas", "Contoh keputusan": "Jalankan Kompatibilitas Data sebelum uji"},
        {"Tahap": "6. Analisis utama", "Yang perlu diisi": "Uji/model sesuai pertanyaan", "Contoh keputusan": "Wizard akan menyarankan uji utama + alternatif"},
        {"Tahap": "7. Pelaporan", "Yang perlu diisi": "p-value, effect size, CI, asumsi, makna substantif", "Contoh keputusan": "Gunakan Insight Riset + Output Export"},
    ]
    if sample_context:
        base.append({"Tahap": "Catatan konteks", "Yang perlu diisi": "Konteks sampel/populasi", "Contoh keputusan": sample_context})
    return pd.DataFrame(base)


def assumption_playbook_table():
    return pd.DataFrame([
        {"Asumsi/masalah": "Normalitas", "Kapan dicek": "T-test/ANOVA terutama N kecil; regresi cek residual", "Cara cek": "Shapiro-Wilk/D'Agostino, skewness, kurtosis, histogram, Q-Q plot", "Jika bermasalah": "Cek outlier/input, transformasi, pilih uji nonparametrik, bootstrap/robust, atau laporkan keterbatasan"},
        {"Asumsi/masalah": "Skewness & kurtosis ekstrem", "Kapan dicek": "Saat deskriptif dan sebelum uji parametrik", "Cara cek": "Skewness jauh dari 0; kurtosis jauh dari 0; lihat histogram/boxplot", "Jika bermasalah": "Gunakan median/IQR, cek outlier, transformasi, Spearman/Mann-Whitney/Wilcoxon/Kruskal sesuai tujuan"},
        {"Asumsi/masalah": "Homogenitas varians", "Kapan dicek": "T-test independen/ANOVA", "Cara cek": "Levene test", "Jika bermasalah": "Welch t-test/ANOVA atau nonparametrik"},
        {"Asumsi/masalah": "Outlier", "Kapan dicek": "Semua analisis numerik", "Cara cek": "Boxplot, z-score, IQR, Cook's distance", "Jika bermasalah": "Verifikasi input, winsorize dengan alasan, analisis sensitif"},
        {"Asumsi/masalah": "Linearitas", "Kapan dicek": "Korelasi/regresi", "Cara cek": "Scatter plot, residual plot", "Jika bermasalah": "Transformasi, model non-linear, Spearman"},
        {"Asumsi/masalah": "Multikolinearitas", "Kapan dicek": "Regresi berganda", "Cara cek": "VIF/Tolerance", "Jika bermasalah": "Hapus/gabung prediktor, PCA, pilih variabel berdasar teori"},
        {"Asumsi/masalah": "Independensi", "Kapan dicek": "Semua uji inferensial", "Cara cek": "Desain sampling; Durbin-Watson untuk residual berurutan", "Jika bermasalah": "Gunakan paired/repeated/mixed model sesuai desain"},
        {"Asumsi/masalah": "Expected count kecil", "Kapan dicek": "Chi-square", "Cara cek": "Expected frequency tabel kontingensi", "Jika bermasalah": "Fisher Exact, gabungkan kategori bermakna"},
        {"Asumsi/masalah": "Kelayakan EFA", "Kapan dicek": "Analisis faktor", "Cara cek": "KMO, Bartlett, communality", "Jika bermasalah": "Hapus item lemah, tambah sampel, revisi konstruk"},
    ])


def effect_size_reference_table():
    return pd.DataFrame([
        {"Analisis": "T-test", "Ukuran efek": "Cohen's d / Hedges g", "Kecil": "≈0.20", "Sedang": "≈0.50", "Besar": "≈0.80", "Makna awam": "Besar perbedaan antar rata-rata"},
        {"Analisis": "ANOVA", "Ukuran efek": "η² / partial η² / ω²", "Kecil": "≈0.01", "Sedang": "≈0.06", "Besar": "≈0.14", "Makna awam": "Proporsi variasi Y yang dijelaskan kelompok"},
        {"Analisis": "Korelasi", "Ukuran efek": "r / r²", "Kecil": "≈0.10", "Sedang": "≈0.30", "Besar": "≈0.50", "Makna awam": "Kekuatan hubungan antar variabel"},
        {"Analisis": "Regresi", "Ukuran efek": "R² / adjusted R² / β", "Kecil": "kontekstual", "Sedang": "kontekstual", "Besar": "kontekstual", "Makna awam": "Seberapa baik model menjelaskan/memprediksi Y"},
        {"Analisis": "Chi-square", "Ukuran efek": "Cramer's V", "Kecil": "≈0.10", "Sedang": "≈0.30", "Besar": "≈0.50", "Makna awam": "Kekuatan asosiasi antar kategori"},
        {"Analisis": "Mann-Whitney", "Ukuran efek": "Rank-biserial / r", "Kecil": "≈0.10", "Sedang": "≈0.30", "Besar": "≈0.50", "Makna awam": "Seberapa jelas perbedaan peringkat antar kelompok"},
        {"Analisis": "Reliabilitas", "Ukuran efek": "Cronbach's alpha", "Rendah": "<0.60", "Cukup": "0.60–0.70", "Baik": ">=0.70", "Makna awam": "Konsistensi item dalam satu skala"},
    ])


def glossary_table():
    return pd.DataFrame([
        {"Istilah": "p-value", "Makna sederhana": "Seberapa kuat data menentang H0; bukan ukuran besar efek", "Jangan disalahartikan sebagai": "Probabilitas hipotesis benar/salah"},
        {"Istilah": "Alpha (α)", "Makna sederhana": "Batas keputusan, sering 0.05", "Jangan disalahartikan sebagai": "Kebenaran mutlak"},
        {"Istilah": "Effect size", "Makna sederhana": "Besar/kekuatan temuan secara praktis", "Jangan disalahartikan sebagai": "Pengganti desain riset yang baik"},
        {"Istilah": "Confidence interval", "Makna sederhana": "Rentang estimasi yang masuk akal", "Jangan disalahartikan sebagai": "Rentang 95% data mentah"},
        {"Istilah": "Normalitas", "Makna sederhana": "Pola distribusi data/residual mendekati lonceng", "Jangan disalahartikan sebagai": "Wajib sempurna untuk semua kondisi"},
        {"Istilah": "Reliabilitas", "Makna sederhana": "Konsistensi item dalam mengukur konstruk", "Jangan disalahartikan sebagai": "Validitas/kebenaran konstruk"},
        {"Istilah": "EFA", "Makna sederhana": "Mencari struktur faktor/dimensi dari banyak item", "Jangan disalahartikan sebagai": "Bukti final tanpa teori/validasi"},
    ])


def data_format_recipes_table():
    return pd.DataFrame([
        {"Kebutuhan analisis": "T-test independen", "Format terbaik": "1 kolom skor + 1 kolom grup", "Contoh": "nilai, kelompok"},
        {"Kebutuhan analisis": "Paired t-test", "Format terbaik": "1 baris = 1 responden; kolom pre dan post", "Contoh": "pretest, posttest"},
        {"Kebutuhan analisis": "ANOVA", "Format terbaik": "1 kolom skor + 1 kolom grup 3+ kategori", "Contoh": "kepuasan, kelas"},
        {"Kebutuhan analisis": "Korelasi/regresi", "Format terbaik": "Setiap variabel numerik dalam kolom terpisah", "Contoh": "motivasi, jam_belajar, nilai"},
        {"Kebutuhan analisis": "Reliabilitas/EFA", "Format terbaik": "Setiap item kuesioner menjadi kolom numerik", "Contoh": "item1, item2, item3"},
        {"Kebutuhan analisis": "Chi-square", "Format terbaik": "Dua kolom kategori, bukan angka total ringkasan", "Contoh": "gender, pilihan_produk"},
    ])


def render_quick_start(df, num_cols, cat_cols, all_cols):
    st.subheader("🚀 Mulai Cepat — Dipandu, Tidak Membingungkan")
    st.caption("Halaman ini merangkum kondisi data dan memberi urutan kerja. Detail statistik tetap tersedia, tetapi dibuka bertahap.")
    issues = analyze_data_compatibility(df, st.session_state.get("metadata"))
    score = compatibility_score(issues)
    label, klass = status_color_label(score)

    a, b, c, d = st.columns(4)
    a.metric("Skor kesiapan", f"{score}/100", label)
    b.metric("Baris data", f"{df.shape[0]:,}")
    c.metric("Kolom numerik", len(num_cols))
    d.metric("Kolom kategori", len(cat_cols))

    st.markdown(f"<div class='{klass}'><b>Status data: {label}</b><br><span class='tiny'>Gunakan rekomendasi langkah berikutnya di bawah agar analisis tidak asal klik.</span></div>", unsafe_allow_html=True)

    st.markdown("### Jalur kerja yang disarankan")
    next_actions = build_next_best_actions(df, issues, score)
    st.dataframe(next_actions, use_container_width=True, hide_index=True)

    q1, q2, q3 = st.columns(3)
    with q1:
        st.markdown(card_html("Saya belum tahu uji apa", "Buka Smart Assistant → Wizard Uji Otomatis. Pilih tujuan riset dan variabel; aplikasi akan menyarankan uji.", "🧙", "guide"), unsafe_allow_html=True)
    with q2:
        st.markdown(card_html("Data saya berantakan", "Buka Kompatibilitas Data atau Data Repair Assistant. Aplikasi memberi tahu apa yang perlu diubah, ditambah, atau diganti.", "🧰", "guide"), unsafe_allow_html=True)
    with q3:
        st.markdown(card_html("Saya butuh laporan", "Jalankan analisis, lalu buka Insight Riset dan Output & Ekspor untuk narasi dan file laporan.", "📝", "guide"), unsafe_allow_html=True)

    st.markdown("### Analisis yang mungkin cocok untuk dataset ini")
    st.dataframe(suggest_analysis_table(df), use_container_width=True, hide_index=True)

    with st.expander("Lihat peran tiap kolom — cocok untuk Y, X, grup, atau item?", expanded=detail_is_full()):
        st.dataframe(variable_role_suggestions(df), use_container_width=True, hide_index=True)

    with st.expander("Peta lengkap fitur aplikasi", expanded=False):
        st.dataframe(comprehensive_ui_feature_map(), use_container_width=True, hide_index=True)

    if st.button("💾 Simpan ringkasan mulai cepat ke Output Viewer", key="save_quick_start_summary"):
        summary = pd.concat([
            next_actions.assign(Tabel="Langkah berikutnya"),
            suggest_analysis_table(df).rename(columns={"Uji/Analisis": "Langkah"}).assign(Tabel="Saran analisis"),
        ], ignore_index=True, sort=False)
        add_report("Ringkasan Mulai Cepat", summary, f"Skor kesiapan data: {score}/100 ({label}).")
        st.success("Ringkasan disimpan ke Output Viewer.")


def render_reference_center(df):
    st.subheader("📚 Panduan Lengkap & Glosarium")
    st.caption("Bagian ini menyimpan detail agar UI utama tetap bersih. Cocok untuk belajar, mengecek asumsi, dan menyiapkan laporan.")
    guide_mode = st.radio(
        "Pilih panduan",
        ["Pohon Keputusan Uji", "Asumsi & Solusi", "Normalitas & Distribusi", "Effect Size", "Format Data", "Glosarium", "Checklist Laporan"],
        horizontal=True,
        key="reference_center_mode",
    )
    if guide_mode == "Pohon Keputusan Uji":
        st.dataframe(analysis_decision_matrix(df), use_container_width=True, hide_index=True)
    elif guide_mode == "Asumsi & Solusi":
        st.dataframe(assumption_playbook_table(), use_container_width=True, hide_index=True)
    elif guide_mode == "Normalitas & Distribusi":
        st.markdown("#### Skewness & Kurtosis")
        st.dataframe(skewness_kurtosis_reference_table(), use_container_width=True, hide_index=True)
        st.markdown("#### Jika uji normalitas tidak tercapai")
        st.dataframe(normality_solution_table(), use_container_width=True, hide_index=True)
        st.info("Gunakan keputusan gabungan: p-value normalitas, Q-Q plot/histogram, skewness-kurtosis, ukuran sampel, dan tujuan analisis. Jangan hanya memakai satu angka.")
    elif guide_mode == "Effect Size":
        st.dataframe(effect_size_reference_table(), use_container_width=True, hide_index=True)
        st.info("Baca effect size bersama konteks riset. Nilai kecil tetap bisa penting jika dampaknya luas; nilai besar perlu tetap dicek validitas datanya.")
    elif guide_mode == "Format Data":
        st.dataframe(data_format_recipes_table(), use_container_width=True, hide_index=True)
        st.markdown("<span class='step-pill'>1 baris = 1 responden/observasi</span> <span class='step-pill'>1 kolom = 1 variabel</span> <span class='step-pill'>Kode kategori konsisten</span> <span class='step-pill'>Missing value jelas</span>", unsafe_allow_html=True)
    elif guide_mode == "Glosarium":
        st.dataframe(glossary_table(), use_container_width=True, hide_index=True)
    else:
        checklist = pd.DataFrame([
            {"Bagian laporan": "Deskripsi data", "Harus ada": "N, mean/median, SD/IQR, frekuensi kategori"},
            {"Bagian laporan": "Asumsi", "Harus ada": "Normalitas/homogenitas/outlier/VIF sesuai analisis"},
            {"Bagian laporan": "Hasil utama", "Harus ada": "Statistik uji, df, p-value, CI, effect size"},
            {"Bagian laporan": "Makna riset", "Harus ada": "Arah temuan, besar dampak, kaitan dengan teori"},
            {"Bagian laporan": "Keterbatasan", "Harus ada": "Sampel, missing, asumsi, desain, generalisasi"},
            {"Bagian laporan": "Rekomendasi", "Harus ada": "Analisis lanjutan atau keputusan praktis"},
        ])
        st.dataframe(checklist, use_container_width=True, hide_index=True)

# Header
left, right = st.columns([0.72, 0.28])
with left:
    st.title(f"📊 {APP_NAME}")
    st.markdown(f"<div class='small-note'>{APP_SUBTITLE}</div>", unsafe_allow_html=True)
with right:
    st.metric("Output tersimpan", len(st.session_state.report_items))



# -----------------------------------------------------------------------------
# v5.0 Advanced Research Analytics: SPSS-like modules with defensive UI
# -----------------------------------------------------------------------------
def _analysis_safe_note(name):
    return f"Bagian {name} mengalami kendala, tetapi aplikasi tetap berjalan. Cek pilihan variabel dan format data, lalu coba lagi."


def _quote_col(col):
    return str(col).replace('"', '\\"')


def _modeling_frame(df, cols):
    """Return clean modeling frame with simple variable aliases v0, v1, ..."""
    cols = [c for c in cols if c in df.columns]
    aliases = {c: f"v{i}" for i, c in enumerate(cols)}
    mdf = df[cols].rename(columns=aliases).copy()
    return mdf.dropna(), aliases


def _fmt_ci(lo, hi, digits=4):
    if pd.isna(lo) or pd.isna(hi):
        return "NA"
    return f"[{lo:.{digits}f}, {hi:.{digits}f}]"


def _cohen_d_independent(x, y):
    x = pd.to_numeric(pd.Series(x), errors="coerce").dropna().to_numpy(dtype=float)
    y = pd.to_numeric(pd.Series(y), errors="coerce").dropna().to_numpy(dtype=float)
    if len(x) < 2 or len(y) < 2:
        return np.nan, np.nan
    pooled = np.sqrt(((len(x)-1)*np.var(x, ddof=1) + (len(y)-1)*np.var(y, ddof=1)) / (len(x)+len(y)-2))
    if pooled == 0 or np.isnan(pooled):
        return np.nan, np.nan
    d = (np.mean(x) - np.mean(y)) / pooled
    correction = 1 - (3 / (4*(len(x)+len(y))-9)) if (len(x)+len(y)) > 2 else 1
    return d, d * correction


def _cohen_d_paired(x, y):
    pair = pd.DataFrame({"x": x, "y": y}).apply(pd.to_numeric, errors="coerce").dropna()
    if len(pair) < 2:
        return np.nan
    diff = pair["x"] - pair["y"]
    sd = diff.std(ddof=1)
    return diff.mean() / sd if sd and not np.isnan(sd) else np.nan


def _cramers_v_from_table(table):
    try:
        chi2, p, dof, expected = stats.chi2_contingency(table)
        n = np.asarray(table).sum()
        if n == 0:
            return np.nan
        r, k = table.shape
        denom = n * max(1, min(k-1, r-1))
        return np.sqrt(chi2 / denom) if denom > 0 else np.nan
    except Exception:
        return np.nan


def _bootstrap_array(data, statistic_func, n_boot=2000, seed=123):
    data = np.asarray(data)
    n = len(data)
    if n < 2:
        return np.nan, np.nan, np.nan
    rng = np.random.default_rng(seed)
    vals = []
    for _ in range(int(n_boot)):
        idx = rng.integers(0, n, n)
        try:
            val = statistic_func(data[idx])
            if np.isfinite(val):
                vals.append(float(val))
        except Exception:
            continue
    if not vals:
        return np.nan, np.nan, np.nan
    vals = np.asarray(vals)
    return float(np.mean(vals)), float(np.percentile(vals, 2.5)), float(np.percentile(vals, 97.5))


def _bootstrap_two_arrays(x, y, statistic_func, paired=False, n_boot=2000, seed=123):
    x = np.asarray(x, dtype=float)
    y = np.asarray(y, dtype=float)
    rng = np.random.default_rng(seed)
    vals = []
    if paired:
        n = min(len(x), len(y))
        if n < 2:
            return np.nan, np.nan, np.nan
        x, y = x[:n], y[:n]
        for _ in range(int(n_boot)):
            idx = rng.integers(0, n, n)
            try:
                val = statistic_func(x[idx], y[idx])
                if np.isfinite(val):
                    vals.append(float(val))
            except Exception:
                continue
    else:
        if len(x) < 2 or len(y) < 2:
            return np.nan, np.nan, np.nan
        for _ in range(int(n_boot)):
            ix = rng.integers(0, len(x), len(x))
            iy = rng.integers(0, len(y), len(y))
            try:
                val = statistic_func(x[ix], y[iy])
                if np.isfinite(val):
                    vals.append(float(val))
            except Exception:
                continue
    if not vals:
        return np.nan, np.nan, np.nan
    vals = np.asarray(vals)
    return float(np.mean(vals)), float(np.percentile(vals, 2.5)), float(np.percentile(vals, 97.5))


def _small_table_message(message):
    return pd.DataFrame({"Informasi": [message]})


def render_bootstrap_effect_size(df, num_cols, cat_cols):
    st.markdown("### 🧪 Bootstrapping & Effect Size")
    st.caption("Gunakan saat ingin confidence interval yang lebih robust dan makna praktis, bukan hanya p-value.")
    mode = st.radio(
        "Jenis perhitungan",
        ["Mean satu variabel", "Selisih dua rata-rata", "Korelasi", "Regresi linear koefisien", "Effect size cepat"],
        horizontal=True,
        key="v5_boot_mode",
    )
    n_boot = int(st.number_input("Jumlah bootstrap resampling", min_value=200, max_value=10000, value=2000, step=200, key="v5_boot_n"))
    seed = int(st.number_input("Seed", min_value=1, max_value=999999, value=123, step=1, key="v5_boot_seed"))

    if mode == "Mean satu variabel":
        if not num_cols:
            st.warning("Butuh minimal 1 variabel numerik.")
            return
        col = st.selectbox("Variabel numerik", num_cols, key="v5_boot_mean_col")
        values = pd.to_numeric(df[col], errors="coerce").dropna().to_numpy(dtype=float)
        if st.button("Hitung bootstrap mean", key="v5_boot_mean_btn"):
            boot_mean, lo, hi = _bootstrap_array(values, np.mean, n_boot, seed)
            res = pd.DataFrame([{
                "Variabel": col,
                "N valid": len(values),
                "Mean sampel": np.mean(values) if len(values) else np.nan,
                "Bootstrap mean": boot_mean,
                "CI 95% bawah": lo,
                "CI 95% atas": hi,
                "Makna": "Rata-rata populasi diperkirakan berada dalam rentang CI 95%. Jika rentang sempit, estimasi lebih presisi.",
            }])
            show_table("Bootstrap Mean", res, "Bootstrap membantu saat asumsi normalitas meragukan atau sampel relatif kecil.")

    elif mode == "Selisih dua rata-rata":
        if len(num_cols) < 1:
            st.warning("Butuh variabel numerik sebagai skor.")
            return
        compare_style = st.radio("Format data", ["Satu skor + satu grup", "Dua kolom skor"], horizontal=True, key="v5_boot_diff_style")
        if compare_style == "Satu skor + satu grup":
            score = st.selectbox("Variabel skor", num_cols, key="v5_boot_score")
            if not cat_cols:
                st.warning("Butuh variabel grup/kategori.")
                return
            group = st.selectbox("Variabel grup", cat_cols, key="v5_boot_group")
            levels = df[group].dropna().astype(str).unique().tolist()
            if len(levels) < 2:
                st.warning("Variabel grup harus punya minimal 2 kategori.")
                return
            chosen = st.multiselect("Pilih 2 kelompok", levels, default=levels[:2], key="v5_boot_levels")
            if len(chosen) == 2 and st.button("Hitung bootstrap selisih mean", key="v5_boot_diff_btn"):
                x = pd.to_numeric(df.loc[df[group].astype(str) == chosen[0], score], errors="coerce").dropna().to_numpy(float)
                y = pd.to_numeric(df.loc[df[group].astype(str) == chosen[1], score], errors="coerce").dropna().to_numpy(float)
                stat, lo, hi = _bootstrap_two_arrays(x, y, lambda a, b: np.mean(a) - np.mean(b), False, n_boot, seed)
                d, g = _cohen_d_independent(x, y)
                res = pd.DataFrame([{
                    "Skor": score,
                    "Kelompok 1": chosen[0], "N1": len(x), "Mean1": np.mean(x) if len(x) else np.nan,
                    "Kelompok 2": chosen[1], "N2": len(y), "Mean2": np.mean(y) if len(y) else np.nan,
                    "Selisih Mean": np.mean(x)-np.mean(y) if len(x) and len(y) else np.nan,
                    "Bootstrap CI 95%": _fmt_ci(lo, hi),
                    "Cohen's d": d, "Hedges' g": g,
                    "Makna": "Jika CI tidak melewati 0, perbedaan rata-rata relatif kuat secara statistik. Cohen's d/g menunjukkan besar perbedaan praktis.",
                }])
                show_table("Bootstrap Selisih Mean + Effect Size", res)
        else:
            c1 = st.selectbox("Kolom skor 1", num_cols, key="v5_boot_col1")
            c2 = st.selectbox("Kolom skor 2", num_cols, index=1 if len(num_cols)>1 else 0, key="v5_boot_col2")
            paired = st.checkbox("Data berpasangan/pre-post", value=True, key="v5_boot_paired")
            if c1 == c2:
                st.info("Pilih dua kolom berbeda.")
            elif st.button("Hitung bootstrap selisih kolom", key="v5_boot_diffcols_btn"):
                pair = df[[c1, c2]].apply(pd.to_numeric, errors="coerce").dropna() if paired else None
                x = pair[c1].to_numpy(float) if paired else pd.to_numeric(df[c1], errors="coerce").dropna().to_numpy(float)
                y = pair[c2].to_numpy(float) if paired else pd.to_numeric(df[c2], errors="coerce").dropna().to_numpy(float)
                stat, lo, hi = _bootstrap_two_arrays(x, y, lambda a, b: np.mean(a)-np.mean(b), paired, n_boot, seed)
                effect = _cohen_d_paired(x, y) if paired else _cohen_d_independent(x, y)[0]
                res = pd.DataFrame([{"Kolom 1": c1, "Kolom 2": c2, "N": len(x) if paired else f"{len(x)} / {len(y)}", "Selisih Mean": stat, "CI 95%": _fmt_ci(lo, hi), "Effect Size": effect}])
                show_table("Bootstrap Selisih Dua Kolom", res)

    elif mode == "Korelasi":
        if len(num_cols) < 2:
            st.warning("Butuh minimal 2 variabel numerik.")
            return
        xcol = st.selectbox("Variabel X", num_cols, key="v5_boot_corr_x")
        ycol = st.selectbox("Variabel Y", [c for c in num_cols if c != xcol] or num_cols, key="v5_boot_corr_y")
        method = st.radio("Metode", ["Pearson", "Spearman"], horizontal=True, key="v5_boot_corr_method")
        if st.button("Hitung bootstrap korelasi", key="v5_boot_corr_btn"):
            pair = df[[xcol, ycol]].apply(pd.to_numeric, errors="coerce").dropna()
            arr = pair.to_numpy(float)
            def corr_stat(sample):
                if method == "Spearman":
                    return stats.spearmanr(sample[:, 0], sample[:, 1]).statistic
                return stats.pearsonr(sample[:, 0], sample[:, 1]).statistic
            boot, lo, hi = _bootstrap_array(arr, corr_stat, n_boot, seed)
            p = stats.spearmanr(pair[xcol], pair[ycol]).pvalue if method == "Spearman" else stats.pearsonr(pair[xcol], pair[ycol]).pvalue
            res = pd.DataFrame([{"X": xcol, "Y": ycol, "N": len(pair), "r sampel": corr_stat(arr), "p-value": p, "Bootstrap r": boot, "CI 95%": _fmt_ci(lo, hi), "Makna": "Arah ditunjukkan oleh tanda r; kekuatan praktis dibaca dari besar |r| dan CI."}])
            show_table("Bootstrap Korelasi", res)

    elif mode == "Regresi linear koefisien":
        if sm is None or len(num_cols) < 2:
            st.warning("Butuh statsmodels dan minimal 2 variabel numerik.")
            return
        y = st.selectbox("Variabel dependen/Y", num_cols, key="v5_boot_reg_y")
        xs = st.multiselect("Prediktor numerik/X", [c for c in num_cols if c != y], default=[c for c in num_cols if c != y][:1], key="v5_boot_reg_xs")
        if xs and st.button("Hitung bootstrap koefisien regresi", key="v5_boot_reg_btn"):
            data = df[[y] + xs].apply(pd.to_numeric, errors="coerce").dropna()
            if len(data) < len(xs) + 3:
                st.warning("Kasus lengkap terlalu sedikit untuk regresi.")
                return
            X = sm.add_constant(data[xs])
            model = sm.OLS(data[y], X).fit()
            rng = np.random.default_rng(seed)
            boot_coefs = []
            for _ in range(n_boot):
                idx = rng.integers(0, len(data), len(data))
                sample = data.iloc[idx]
                try:
                    bm = sm.OLS(sample[y], sm.add_constant(sample[xs])).fit()
                    boot_coefs.append(bm.params.reindex(model.params.index).to_numpy(float))
                except Exception:
                    continue
            if not boot_coefs:
                st.warning("Bootstrap regresi gagal dihitung. Cek multikolinearitas atau data konstan.")
                return
            boot_arr = np.vstack(boot_coefs)
            rows = []
            for i, term in enumerate(model.params.index):
                rows.append({"Term": term, "Koefisien OLS": model.params[term], "p-value": model.pvalues[term], "CI bootstrap bawah": np.percentile(boot_arr[:, i], 2.5), "CI bootstrap atas": np.percentile(boot_arr[:, i], 97.5)})
            res = pd.DataFrame(rows)
            show_table("Bootstrap Koefisien Regresi", res, "Jika CI bootstrap prediktor tidak melewati 0, kontribusi prediktor relatif kuat.")

    else:
        effect_mode = st.radio("Effect size", ["Cohen's d dua kelompok", "Cohen's dz berpasangan", "Cramer's V crosstab"], horizontal=True, key="v5_effect_mode")
        if effect_mode == "Cohen's d dua kelompok":
            if not num_cols or not cat_cols:
                st.warning("Butuh skor numerik dan grup kategori.")
                return
            score = st.selectbox("Skor", num_cols, key="v5_eff_score")
            group = st.selectbox("Grup", cat_cols, key="v5_eff_group")
            levels = df[group].dropna().astype(str).unique().tolist()
            chosen = st.multiselect("Pilih 2 grup", levels, default=levels[:2], key="v5_eff_levels")
            if len(chosen) == 2 and st.button("Hitung Cohen's d", key="v5_eff_d_btn"):
                x = pd.to_numeric(df.loc[df[group].astype(str)==chosen[0], score], errors="coerce").dropna()
                y = pd.to_numeric(df.loc[df[group].astype(str)==chosen[1], score], errors="coerce").dropna()
                d, g = _cohen_d_independent(x, y)
                show_table("Effect Size Cohen's d", pd.DataFrame([{"Skor": score, "Grup 1": chosen[0], "Grup 2": chosen[1], "Cohen's d": d, "Hedges' g": g, "Interpretasi": _effect_size_label(d, "d") if not pd.isna(d) else "NA"}]))
        elif effect_mode == "Cohen's dz berpasangan":
            if len(num_cols) < 2:
                st.warning("Butuh dua kolom numerik.")
                return
            before = st.selectbox("Kolom 1", num_cols, key="v5_eff_before")
            after = st.selectbox("Kolom 2", [c for c in num_cols if c != before] or num_cols, key="v5_eff_after")
            if st.button("Hitung dz", key="v5_eff_dz_btn"):
                dz = _cohen_d_paired(df[before], df[after])
                show_table("Effect Size Paired", pd.DataFrame([{"Kolom 1": before, "Kolom 2": after, "Cohen's dz": dz, "Interpretasi": _effect_size_label(dz, "d") if not pd.isna(dz) else "NA"}]))
        else:
            if len(cat_cols) < 2:
                st.warning("Butuh dua variabel kategori.")
                return
            a = st.selectbox("Kategori 1", cat_cols, key="v5_eff_cat1")
            b = st.selectbox("Kategori 2", [c for c in cat_cols if c != a] or cat_cols, key="v5_eff_cat2")
            if st.button("Hitung Cramer's V", key="v5_eff_v_btn"):
                tab = pd.crosstab(df[a], df[b])
                v = _cramers_v_from_table(tab)
                show_table("Effect Size Cramer's V", pd.DataFrame([{"Variabel 1": a, "Variabel 2": b, "Cramer's V": v, "Interpretasi": _effect_size_label(v, "r") if not pd.isna(v) else "NA"}]))
                st.dataframe(tab, use_container_width=True)


def render_glm_family(df, num_cols, cat_cols, all_cols):
    st.markdown("### 📐 ANCOVA, MANOVA, dan Repeated Measures")
    st.caption("Modul lanjutan ini dibuat bertahap dan defensif. Gunakan setelah data sudah bersih dan variabel jelas.")
    analysis = st.radio("Pilih analisis", ["ANCOVA", "MANOVA", "Repeated Measures ANOVA"], horizontal=True, key="v5_glm_mode")
    alpha = st.number_input("Alpha", min_value=0.001, max_value=0.20, value=float(st.session_state.get("active_alpha", 0.05)), step=0.01, format="%.3f", key="v5_glm_alpha")

    if analysis == "ANCOVA":
        if sm is None or not num_cols or not cat_cols:
            st.warning("ANCOVA butuh statsmodels, 1 variabel dependen numerik, 1 faktor kategori, dan kovariat numerik.")
            return
        y = st.selectbox("Variabel dependen/Y numerik", num_cols, key="v5_ancova_y")
        factor = st.selectbox("Faktor/kategori", cat_cols, key="v5_ancova_factor")
        covs = st.multiselect("Kovariat numerik", [c for c in num_cols if c != y], default=[c for c in num_cols if c != y][:1], key="v5_ancova_covs")
        typ = st.radio("Type SS", [2, 3], horizontal=True, key="v5_ancova_typ")
        if covs and st.button("Jalankan ANCOVA", key="v5_ancova_btn"):
            cols = [y, factor] + covs
            data, aliases = _modeling_frame(df, cols)
            if len(data) < len(covs) + data[aliases[factor]].nunique() + 3:
                st.warning("Kasus lengkap terlalu sedikit untuk ANCOVA.")
                return
            formula = f"{aliases[y]} ~ C({aliases[factor]}) + " + " + ".join(aliases[c] for c in covs)
            if int(typ) == 3:
                formula = f"{aliases[y]} ~ C({aliases[factor]}) + " + " + ".join(aliases[c] for c in covs)
            model = smf.ols(formula, data=data).fit()
            table = sm.stats.anova_lm(model, typ=int(typ)).reset_index().rename(columns={"index": "Efek"})
            if "sum_sq" in table.columns:
                resid_ss = float(table.loc[table["Efek"].astype(str).str.lower().eq("residual"), "sum_sq"].iloc[0]) if any(table["Efek"].astype(str).str.lower().eq("residual")) else np.nan
                table["partial_eta_sq"] = table["sum_sq"].apply(lambda ss: ss/(ss+resid_ss) if pd.notna(resid_ss) and (ss+resid_ss)>0 else np.nan)
            table["Keputusan"] = table.get("PR(>F)", pd.Series([np.nan]*len(table))).apply(lambda p: decision_text(p, alpha))
            note = "ANCOVA menguji perbedaan rata-rata Y antar kelompok setelah mengontrol kovariat. Baca efek faktor utama setelah kontrol kovariat."
            show_table("ANCOVA", table, note)
            st.caption("Model: " + formula)
            add_report("Ringkasan Model ANCOVA", pd.DataFrame({"Parameter": model.params.index, "Koefisien": model.params.values, "p-value": model.pvalues.values}))

    elif analysis == "MANOVA":
        if sm is None or len(num_cols) < 2 or not all_cols:
            st.warning("MANOVA butuh minimal 2 variabel dependen numerik dan minimal 1 prediktor.")
            return
        dvs = st.multiselect("Variabel dependen numerik/lebih dari satu", num_cols, default=num_cols[:min(2, len(num_cols))], key="v5_manova_dvs")
        predictors = st.multiselect("Prediktor/faktor/kovariat", [c for c in all_cols if c not in dvs], default=cat_cols[:1] if cat_cols else [c for c in num_cols if c not in dvs][:1], key="v5_manova_preds")
        if len(dvs) >= 2 and predictors and st.button("Jalankan MANOVA", key="v5_manova_btn"):
            try:
                from statsmodels.multivariate.manova import MANOVA
                cols = dvs + predictors
                data, aliases = _modeling_frame(df, cols)
                terms = []
                for p in predictors:
                    if p in cat_cols or not pd.api.types.is_numeric_dtype(df[p]):
                        terms.append(f"C({aliases[p]})")
                    else:
                        terms.append(aliases[p])
                formula = " + ".join(aliases[d] for d in dvs) + " ~ " + " + ".join(terms)
                fit = MANOVA.from_formula(formula, data=data)
                text_summary = str(fit.mv_test())
                st.text(text_summary)
                report_df = pd.DataFrame({"MANOVA Summary": text_summary.splitlines()[:500]})
                add_report("MANOVA", report_df, "MANOVA menguji apakah kombinasi beberapa variabel dependen berbeda/terpengaruh oleh prediktor.")
            except Exception as exc:
                st.error("MANOVA gagal dihitung. Pastikan data tidak singular, kategori tidak terlalu banyak, dan kasus lengkap cukup.")
                st.exception(exc)

    else:
        st.info("Repeated Measures ANOVA pada modul ini memakai format long: satu baris = satu pengukuran pada satu subjek.")
        if sm is None:
            st.warning("Butuh statsmodels.")
            return
        subject = st.selectbox("ID subjek/responden", all_cols, key="v5_rm_subject")
        within_candidates = [c for c in all_cols if c != subject]
        within = st.selectbox("Faktor waktu/kondisi (within-subject)", within_candidates, key="v5_rm_within")
        dv = st.selectbox("Variabel skor/Y numerik", [c for c in num_cols if c not in [subject, within]] or num_cols, key="v5_rm_dv")
        if st.button("Jalankan Repeated Measures ANOVA", key="v5_rm_btn"):
            try:
                from statsmodels.stats.anova import AnovaRM
                data = df[[subject, within, dv]].dropna().copy()
                data[dv] = pd.to_numeric(data[dv], errors="coerce")
                data = data.dropna()
                if data[subject].nunique() < 2 or data[within].nunique() < 2:
                    st.warning("Butuh minimal 2 subjek dan 2 waktu/kondisi.")
                    return
                fit = AnovaRM(data, depvar=dv, subject=subject, within=[within]).fit()
                table = fit.anova_table.reset_index().rename(columns={"index": "Efek"})
                table["Keputusan"] = table.get("Pr > F", pd.Series([np.nan]*len(table))).apply(lambda p: decision_text(p, alpha))
                show_table("Repeated Measures ANOVA", table, "Digunakan untuk membandingkan skor subjek yang sama pada beberapa waktu/kondisi.")
            except Exception as exc:
                st.error("Repeated Measures ANOVA gagal. Pastikan data long dan seimbang; setiap subjek sebaiknya punya semua waktu/kondisi.")
                st.exception(exc)


def render_mediation_moderation(df, num_cols, cat_cols):
    st.markdown("### 🔗 Mediasi & Moderasi")
    st.caption("Cocok untuk riset sosial, pendidikan, psikologi, manajemen: menguji mekanisme dan kondisi pengaruh.")
    if sm is None:
        st.warning("Butuh statsmodels.")
        return
    mode = st.radio("Analisis", ["Mediasi sederhana", "Moderasi sederhana"], horizontal=True, key="v5_medmod_mode")
    alpha = st.number_input("Alpha", min_value=0.001, max_value=0.20, value=float(st.session_state.get("active_alpha", 0.05)), step=0.01, format="%.3f", key="v5_medmod_alpha")

    if len(num_cols) < 3:
        st.warning("Modul ini membutuhkan minimal 3 variabel numerik.")
        return

    if mode == "Mediasi sederhana":
        x = st.selectbox("X / variabel independen", num_cols, key="v5_med_x")
        m = st.selectbox("M / mediator", [c for c in num_cols if c != x] or num_cols, key="v5_med_m")
        y = st.selectbox("Y / variabel dependen", [c for c in num_cols if c not in [x, m]] or num_cols, key="v5_med_y")
        covs = st.multiselect("Kovariat opsional", [c for c in num_cols if c not in [x, m, y]], key="v5_med_covs")
        n_boot = int(st.number_input("Bootstrap indirect effect", min_value=200, max_value=10000, value=2000, step=200, key="v5_med_boot"))
        if st.button("Jalankan Mediasi", key="v5_med_btn"):
            data = df[[x, m, y] + covs].apply(pd.to_numeric, errors="coerce").dropna()
            if len(data) < len(covs) + 10:
                st.warning("Kasus lengkap terlalu sedikit untuk mediasi.")
                return
            X_a = sm.add_constant(data[[x] + covs])
            a_model = sm.OLS(data[m], X_a).fit()
            X_b = sm.add_constant(data[[x, m] + covs])
            b_model = sm.OLS(data[y], X_b).fit()
            X_c = sm.add_constant(data[[x] + covs])
            c_model = sm.OLS(data[y], X_c).fit()
            a = a_model.params.get(x, np.nan)
            b = b_model.params.get(m, np.nan)
            cp = b_model.params.get(x, np.nan)
            c_total = c_model.params.get(x, np.nan)
            indirect = a * b
            rng = np.random.default_rng(123)
            boot_vals = []
            for _ in range(n_boot):
                idx = rng.integers(0, len(data), len(data))
                sample = data.iloc[idx]
                try:
                    am = sm.OLS(sample[m], sm.add_constant(sample[[x] + covs])).fit().params.get(x, np.nan)
                    bm = sm.OLS(sample[y], sm.add_constant(sample[[x, m] + covs])).fit().params.get(m, np.nan)
                    val = am * bm
                    if np.isfinite(val):
                        boot_vals.append(val)
                except Exception:
                    continue
            lo, hi = (np.percentile(boot_vals, 2.5), np.percentile(boot_vals, 97.5)) if boot_vals else (np.nan, np.nan)
            res = pd.DataFrame([
                {"Jalur": "a: X → M", "Koefisien": a, "p-value": a_model.pvalues.get(x, np.nan), "Makna": "Apakah X berhubungan dengan mediator."},
                {"Jalur": "b: M → Y | X", "Koefisien": b, "p-value": b_model.pvalues.get(m, np.nan), "Makna": "Apakah mediator menjelaskan Y setelah mengontrol X."},
                {"Jalur": "c': X → Y | M", "Koefisien": cp, "p-value": b_model.pvalues.get(x, np.nan), "Makna": "Efek langsung X ke Y setelah mediator masuk."},
                {"Jalur": "c total: X → Y", "Koefisien": c_total, "p-value": c_model.pvalues.get(x, np.nan), "Makna": "Efek total X ke Y."},
                {"Jalur": "Indirect a*b", "Koefisien": indirect, "p-value": np.nan, "CI Bootstrap 95%": _fmt_ci(lo, hi), "Makna": "Mediasi didukung jika CI indirect effect tidak melewati 0."},
            ])
            show_table("Mediasi Sederhana", res, "Interpretasi utama mediasi dibaca dari bootstrap CI pada indirect effect.")

    else:
        x = st.selectbox("X / prediktor numerik", num_cols, key="v5_mod_x")
        w = st.selectbox("W / moderator numerik", [c for c in num_cols if c != x] or num_cols, key="v5_mod_w")
        y = st.selectbox("Y / dependen numerik", [c for c in num_cols if c not in [x, w]] or num_cols, key="v5_mod_y")
        covs = st.multiselect("Kovariat opsional", [c for c in num_cols if c not in [x, w, y]], key="v5_mod_covs")
        if st.button("Jalankan Moderasi", key="v5_mod_btn"):
            data = df[[x, w, y] + covs].apply(pd.to_numeric, errors="coerce").dropna().copy()
            if len(data) < len(covs) + 10:
                st.warning("Kasus lengkap terlalu sedikit untuk moderasi.")
                return
            data["X_centered"] = data[x] - data[x].mean()
            data["W_centered"] = data[w] - data[w].mean()
            data["XxW"] = data["X_centered"] * data["W_centered"]
            Xcols = ["X_centered", "W_centered", "XxW"] + covs
            model = sm.OLS(data[y], sm.add_constant(data[Xcols])).fit()
            table = pd.DataFrame({"Term": model.params.index, "Koefisien": model.params.values, "p-value": model.pvalues.values, "CI bawah": model.conf_int()[0].values, "CI atas": model.conf_int()[1].values})
            table["Makna"] = table["Term"].apply(lambda t: "Interaksi signifikan berarti pengaruh X terhadap Y berubah tergantung tingkat moderator W." if t == "XxW" else "Parameter model regresi.")
            show_table("Moderasi Sederhana", table, "Fokus utama moderasi adalah term interaksi XxW.")
            st.caption(f"R² = {model.rsquared:.4f}; Adjusted R² = {model.rsquared_adj:.4f}")


def _simple_exponential_smoothing(values, alpha=0.3):
    values = np.asarray(values, dtype=float)
    if len(values) == 0:
        return np.array([])
    smoothed = [values[0]]
    for val in values[1:]:
        smoothed.append(alpha * val + (1 - alpha) * smoothed[-1])
    return np.asarray(smoothed)


def render_forecasting(df, num_cols, all_cols):
    st.markdown("### 📅 Forecasting Sederhana")
    st.caption("Untuk data runtun waktu dasar: trend, moving average, dan exponential smoothing. Cocok untuk eksplorasi awal, bukan model ARIMA penuh.")
    if not num_cols:
        st.warning("Butuh minimal satu variabel numerik.")
        return
    value_col = st.selectbox("Variabel nilai", num_cols, key="v5_forecast_value")
    date_col = st.selectbox("Kolom waktu/tanggal opsional", ["(urutan baris)"] + all_cols, key="v5_forecast_date")
    method = st.radio("Metode", ["Moving Average", "Exponential Smoothing", "Trend Linear"], horizontal=True, key="v5_forecast_method")
    horizon = int(st.number_input("Jumlah periode forecast", min_value=1, max_value=60, value=5, step=1, key="v5_forecast_horizon"))
    data = df[[value_col] + ([] if date_col == "(urutan baris)" else [date_col])].copy()
    data[value_col] = pd.to_numeric(data[value_col], errors="coerce")
    if date_col != "(urutan baris)":
        data[date_col] = pd.to_datetime(data[date_col], errors="coerce")
        data = data.dropna().sort_values(date_col)
    else:
        data = data.dropna().reset_index(drop=True)
        data["Periode"] = np.arange(1, len(data)+1)
        date_col = "Periode"
    if len(data) < 4:
        st.warning("Data runtun waktu minimal 4 observasi agar forecast sederhana masuk akal.")
        return
    window = 3
    alpha_smooth = 0.3
    if method == "Moving Average":
        max_win = min(24, len(data))
        if max_win <= 2:
            window = 2
            st.caption("Window moving average otomatis = 2 karena data sangat pendek.")
        else:
            window = int(st.number_input("Window moving average", min_value=2, max_value=max_win, value=min(3, max_win), step=1, key="v5_forecast_ma_window"))
    elif method == "Exponential Smoothing":
        alpha_smooth = float(st.number_input("Alpha smoothing", min_value=0.01, max_value=0.99, value=0.30, step=0.05, format="%.2f", key="v5_forecast_alpha"))
    if st.button("Buat forecast", key="v5_forecast_btn"):
        y = data[value_col].to_numpy(float)
        if method == "Moving Average":
            ma = pd.Series(y).rolling(window=window, min_periods=1).mean().to_numpy()
            future = np.repeat(ma[-1], horizon)
            fitted = ma
        elif method == "Exponential Smoothing":
            fitted = _simple_exponential_smoothing(y, alpha_smooth)
            future = np.repeat(fitted[-1], horizon)
        else:
            x = np.arange(len(y))
            coef = np.polyfit(x, y, 1)
            fitted = coef[0]*x + coef[1]
            future_x = np.arange(len(y), len(y)+horizon)
            future = coef[0]*future_x + coef[1]
        forecast_df = pd.DataFrame({"Periode_ke": np.arange(len(y)+1, len(y)+horizon+1), "Forecast": future})
        summary = pd.DataFrame([{"Metode": method, "N historis": len(y), "Forecast periode berikutnya": future[0], "Catatan": "Gunakan sebagai estimasi awal; validasi dengan data aktual jika tersedia."}])
        show_table("Forecasting Sederhana", summary)
        st.dataframe(forecast_df, use_container_width=True)
        add_report("Forecast Detail", forecast_df, "Forecast sederhana berdasarkan metode yang dipilih.")
        if px is not None:
            plot_df = pd.DataFrame({"Periode": np.arange(1, len(y)+1), "Aktual": y, "Fitted": fitted})
            st.plotly_chart(px.line(plot_df, x="Periode", y=["Aktual", "Fitted"], title="Aktual vs fitted"), use_container_width=True)


def render_missing_custom_tables(df, num_cols, cat_cols, all_cols):
    st.markdown("### 🧩 Missing Value Analysis & Custom Tables")
    task = st.radio("Pilih modul", ["Missing Value Analysis", "Custom Crosstab", "Tabel Ringkasan by Group"], horizontal=True, key="v5_missing_custom_task")
    if task == "Missing Value Analysis":
        miss = pd.DataFrame({
            "Variabel": df.columns,
            "Missing n": df.isna().sum().values,
            "Missing %": (df.isna().mean().values * 100),
            "Tipe": [str(df[c].dtype) for c in df.columns],
            "Saran": ["Prioritas cek/imputasi" if df[c].isna().mean() >= 0.10 else "Masih relatif aman" if df[c].isna().mean() > 0 else "Tidak ada missing" for c in df.columns],
        }).sort_values("Missing %", ascending=False)
        show_table("Missing Value Summary", miss, "Missing value tinggi dapat mengurangi sampel efektif dan memengaruhi estimasi.", save=False)
        if st.button("Simpan Missing Value Summary", key="v5_save_missing"):
            add_report("Missing Value Summary", miss, "Ringkasan missing value per variabel.")
            st.success("Tersimpan ke Output Viewer.")
        with st.expander("Pola missing value", expanded=False):
            pattern = df.isna().astype(int)
            pattern_cols = pattern.columns.tolist()
            if len(pattern_cols) > 20:
                st.caption("Menampilkan 20 kolom pertama agar tabel tetap ringan.")
                pattern = pattern.iloc[:, :20]
            pattern.columns = [f"kolom_{i+1}" for i in range(pattern.shape[1])]
            patt = pattern.groupby(pattern.columns.tolist(), dropna=False).size().reset_index(name="Jumlah kasus") if not pattern.empty else _small_table_message("Tidak ada data")
            st.dataframe(patt.sort_values("Jumlah kasus", ascending=False).head(100), use_container_width=True)
            st.caption("Pada pola missing: 1 = missing, 0 = terisi. Nama kolom disingkat agar aman jika ada nama variabel duplikat.")
        st.markdown("**Apa yang sebaiknya dilakukan jika banyak missing?**")
        st.markdown("- Cek apakah missing berasal dari kesalahan input/kode seperti 99, -, atau NA.\n- Jika missing kecil dan acak, listwise deletion bisa dipertimbangkan.\n- Jika missing sistematis, jangan langsung hapus; laporkan pola dan pertimbangkan imputasi.\n- Untuk skala kuesioner, hindari imputasi jika responden melewati banyak item inti.")

    elif task == "Custom Crosstab":
        if len(cat_cols) < 2:
            st.warning("Butuh minimal 2 variabel kategori.")
            return
        row = st.selectbox("Baris", cat_cols, key="v5_xtab_row")
        col = st.selectbox("Kolom", [c for c in cat_cols if c != row] or cat_cols, key="v5_xtab_col")
        percent = st.radio("Persentase", ["Jumlah", "% baris", "% kolom", "% total"], horizontal=True, key="v5_xtab_percent")
        if st.button("Buat Custom Crosstab", key="v5_xtab_btn"):
            tab = pd.crosstab(df[row], df[col])
            if percent == "% baris":
                shown = pd.crosstab(df[row], df[col], normalize="index") * 100
            elif percent == "% kolom":
                shown = pd.crosstab(df[row], df[col], normalize="columns") * 100
            elif percent == "% total":
                shown = pd.crosstab(df[row], df[col], normalize="all") * 100
            else:
                shown = tab
            st.dataframe(shown, use_container_width=True)
            chi2, p, dof, exp = stats.chi2_contingency(tab) if tab.size else (np.nan, np.nan, np.nan, None)
            v = _cramers_v_from_table(tab)
            summary = pd.DataFrame([{"Baris": row, "Kolom": col, "Chi-square": chi2, "df": dof, "p-value": p, "Cramer's V": v, "Makna": "Crosstab menunjukkan distribusi kategori; Cramer's V menunjukkan kekuatan asosiasi."}])
            show_table("Custom Crosstab Summary", summary)
            add_report("Custom Crosstab Table", shown.reset_index(), f"Crosstab {row} × {col}; mode: {percent}")

    else:
        if not num_cols:
            st.warning("Butuh minimal 1 variabel numerik.")
            return
        if not cat_cols:
            st.warning("Butuh minimal 1 variabel grup/kategori.")
            return
        group = st.selectbox("Grup", cat_cols, key="v5_sum_group")
        values = st.multiselect("Variabel numerik", num_cols, default=num_cols[:min(4, len(num_cols))], key="v5_sum_values")
        stats_choice = st.multiselect("Statistik", ["N", "Mean", "SD", "Median", "Min", "Max"], default=["N", "Mean", "SD"], key="v5_sum_stats")
        if values and st.button("Buat Tabel Ringkasan", key="v5_sum_btn"):
            grouped = df.groupby(group, dropna=False)[values]
            frames = []
            if "N" in stats_choice: frames.append(grouped.count().add_suffix("_N"))
            if "Mean" in stats_choice: frames.append(grouped.mean(numeric_only=True).add_suffix("_Mean"))
            if "SD" in stats_choice: frames.append(grouped.std(numeric_only=True).add_suffix("_SD"))
            if "Median" in stats_choice: frames.append(grouped.median(numeric_only=True).add_suffix("_Median"))
            if "Min" in stats_choice: frames.append(grouped.min(numeric_only=True).add_suffix("_Min"))
            if "Max" in stats_choice: frames.append(grouped.max(numeric_only=True).add_suffix("_Max"))
            table = pd.concat(frames, axis=1).reset_index() if frames else _small_table_message("Tidak ada statistik dipilih")
            show_table("Tabel Ringkasan by Group", table, "Tabel ini berguna untuk laporan deskriptif per kelompok.")


def render_validation_benchmark(df, num_cols, cat_cols):
    st.markdown("### 🛡️ Validasi, Reproducibility & Benchmark")
    st.caption("Bagian ini membantu meminimalkan bug analisis dan membuat hasil lebih dapat dipercaya.")
    checks = []
    checks.append({"Aspek": "Data aktif", "Status": "OK" if df is not None and not df.empty else "Bermasalah", "Detail": f"{df.shape[0]} baris × {df.shape[1]} kolom" if df is not None else "Tidak ada data"})
    checks.append({"Aspek": "Variabel numerik", "Status": "OK" if len(num_cols) else "Perlu ditambah", "Detail": f"{len(num_cols)} kolom numerik"})
    checks.append({"Aspek": "Variabel kategori", "Status": "OK" if len(cat_cols) else "Opsional", "Detail": f"{len(cat_cols)} kolom kategori"})
    checks.append({"Aspek": "Missing value", "Status": "Perlu cek" if df.isna().sum().sum() > 0 else "OK", "Detail": f"{int(df.isna().sum().sum())} sel kosong"})
    checks.append({"Aspek": "Output viewer", "Status": "OK" if st.session_state.report_items else "Kosong", "Detail": f"{len(st.session_state.report_items)} output tersimpan"})
    checks.append({"Aspek": "Syntax log", "Status": "OK" if st.session_state.syntax_log else "Kosong", "Detail": f"{len(st.session_state.syntax_log)} langkah tercatat"})
    checks.append({"Aspek": "Package statistik", "Status": "OK" if sm is not None else "Terbatas", "Detail": "statsmodels tersedia" if sm is not None else "statsmodels tidak tersedia"})
    table = pd.DataFrame(checks)
    st.dataframe(table, use_container_width=True, hide_index=True)
    if st.button("Simpan checklist validasi", key="v5_validation_save"):
        add_report("Validation Checklist", table, "Checklist reproduksibilitas dan kesiapan analisis.")
        st.success("Checklist tersimpan.")

    st.markdown("#### Benchmark manual yang disarankan")
    st.markdown("""
- Untuk riset formal, uji beberapa kasus kecil dengan **SPSS/R/JASP** dan bandingkan angka utama: mean, SD, t/F/r, p-value, dan effect size.
- Simpan versi data, filter, transformasi, dan output di Output Viewer.
- Jangan hanya mengandalkan p-value; laporkan ukuran efek, CI, asumsi, dan keterbatasan.
- Jika ada hasil berbeda antar-software, cek: missing value handling, coding kategori, jenis sum of squares, dan opsi equal variance.
""")


def render_advanced_research_analytics(df, num_cols, cat_cols, all_cols):
    st.subheader("🔬 Analisis Lanjutan v5.0")
    st.markdown("""
    Modul ini menambahkan fitur yang sering dibutuhkan agar software lebih mendekati alat statistik komprehensif, tetapi tetap dibuat bertahap agar tidak membingungkan.
    Gunakan **Smart Assistant** dulu jika belum yakin uji yang tepat.
    """)
    st.info("Desain anti-bug: tidak memakai tab bertumpuk, semua input memakai key eksplisit, slider berisiko diganti number input/radio, dan tiap modul diberi fallback.")
    sub = st.radio(
        "Pilih fitur lanjutan",
        ["🧪 Bootstrap & Effect Size", "📐 ANCOVA/MANOVA/RM-ANOVA", "🔗 Mediasi/Moderasi", "📅 Forecasting", "🧩 Missing & Custom Tables", "🛡️ Validasi"],
        horizontal=True,
        key="v5_advanced_submenu",
    )
    try:
        if sub == "🧪 Bootstrap & Effect Size":
            render_bootstrap_effect_size(df, num_cols, cat_cols)
        elif sub == "📐 ANCOVA/MANOVA/RM-ANOVA":
            render_glm_family(df, num_cols, cat_cols, all_cols)
        elif sub == "🔗 Mediasi/Moderasi":
            render_mediation_moderation(df, num_cols, cat_cols)
        elif sub == "📅 Forecasting":
            render_forecasting(df, num_cols, all_cols)
        elif sub == "🧩 Missing & Custom Tables":
            render_missing_custom_tables(df, num_cols, cat_cols, all_cols)
        else:
            render_validation_benchmark(df, num_cols, cat_cols)
    except BaseException as exc:
        if _is_streamlit_control_exception(exc):
            raise
        st.error(_analysis_safe_note(sub))
        st.exception(exc)

# Sidebar input
st.sidebar.header("📥 Input Data")
source = st.sidebar.radio("Sumber data", ["Upload File", "Input Manual", "Data Contoh"], index=0)

st.sidebar.divider()
st.sidebar.subheader("🧭 Tampilan")
st.session_state.ui_mode = st.sidebar.radio(
    "Mode pengguna",
    ["Pemula", "Ahli"],
    index=0 if st.session_state.get("ui_mode", "Pemula") == "Pemula" else 1,
    help="Pemula menampilkan alur ringkas; Ahli membuka semua menu teknis.",
    key="ui_mode_selector",
)
st.session_state.detail_level = st.sidebar.radio(
    "Level detail",
    ["Ringkas", "Lengkap"],
    index=0 if st.session_state.get("detail_level", "Ringkas") == "Ringkas" else 1,
    help="Ringkas menyembunyikan tabel/penjelasan lanjutan di expander.",
    key="detail_level_selector",
)
with st.sidebar.expander("Cara pakai cepat", expanded=False):
    st.markdown("""
    1. Upload data atau pakai data contoh.  
    2. Buka **🚀 Mulai Cepat**.  
    3. Ikuti rekomendasi perbaikan data.  
    4. Pilih uji melalui **Smart Assistant**.  
    5. Maknai hasil di **Insight Riset** dan ekspor laporan.
    """)

if st.sidebar.button("🗑️ Reset semua data & output"):
    for key in ["df", "raw_df", "file_name", "report_items", "metadata", "syntax_log", "split_by", "last_action"]:
        st.session_state[key] = [] if key in ["report_items", "syntax_log"] else None
    st.rerun()

if source == "Upload File":
    uploaded = st.sidebar.file_uploader("Upload CSV, Excel, atau SPSS .sav", type=["csv", "xlsx", "xls", "sav"])
    csv_sep = st.sidebar.selectbox("Pemisah CSV", [",", ";", "\t", "|"], index=0)
    decimal = st.sidebar.selectbox("Tanda desimal", [".", ","], index=0)
    sheet_name = None
    if uploaded is not None and uploaded.name.split(".")[-1].lower() in ["xlsx", "xls"]:
        try:
            sheets = pd.ExcelFile(io.BytesIO(uploaded.getvalue())).sheet_names
            sheet_name = st.sidebar.selectbox("Sheet Excel", sheets)
        except Exception:
            sheet_name = None
    if uploaded is not None:
        try:
            df_loaded = load_data(uploaded, csv_sep=csv_sep, decimal=decimal, sheet_name=sheet_name)
            st.session_state.df = df_loaded
            st.session_state.raw_df = df_loaded.copy()
            st.session_state.file_name = uploaded.name
            st.session_state.metadata = build_metadata(df_loaded)
            log_syntax(f"GET DATA /TYPE={uploaded.name.split('.')[-1].upper()} /FILE='{uploaded.name}'.")
            st.sidebar.success(f"✅ {uploaded.name} berhasil dimuat")
        except Exception as exc:
            st.sidebar.error(f"❌ Gagal membaca file: {exc}")

elif source == "Input Manual":
    st.sidebar.info("Baris pertama boleh berisi nama kolom. Pisahkan kolom dengan koma.")
    manual = st.sidebar.text_area(
        "Tempel data",
        "kelompok,nilai,motivasi\nA,72,69\nA,75,73\nB,81,78\nB,84,80\nC,88,85",
        height=180,
    )
    has_header = st.sidebar.checkbox("Baris pertama adalah nama kolom", value=True)
    if st.sidebar.button("🔹 Buat DataFrame"):
        try:
            st.session_state.df = pd.read_csv(io.StringIO(manual), header=0 if has_header else None)
            if not has_header:
                st.session_state.df.columns = [f"V{i+1}" for i in range(st.session_state.df.shape[1])]
            st.session_state.raw_df = st.session_state.df.copy()
            st.session_state.file_name = "manual_input.csv"
            st.session_state.metadata = build_metadata(st.session_state.df)
            log_syntax("DATASET NAME manual_input.")
            st.sidebar.success("✅ Data manual berhasil dibuat")
        except Exception as exc:
            st.sidebar.error(f"❌ Format input belum valid: {exc}")

else:
    if st.sidebar.button("✨ Muat data contoh"):
        st.session_state.df = sample_dataframe()
        st.session_state.raw_df = st.session_state.df.copy()
        st.session_state.file_name = "sample_statistik_pro.csv"
        st.session_state.metadata = build_metadata(st.session_state.df)
        log_syntax("GET DATA /TYPE=SAMPLE.")
        st.sidebar.success("✅ Data contoh dimuat")


df = st.session_state.df
if df is not None:
    sync_metadata(df)

if df is None:
    st.info("📂 Silakan unggah data, tempel data manual, atau gunakan data contoh dari sidebar.")
    with st.expander("Apa yang baru di Statistik Pro+?"):
        st.markdown(
            """
            - Import CSV, Excel, dan SPSS `.sav`.
            - Data View + Variable View, value labels, user-missing values, dan transformasi data.
            - Output Viewer, ekspor Excel/HTML/Markdown/Word, serta interpretasi otomatis.
            - Statistik deskriptif, frekuensi, normalitas, korelasi, crosstab, chi-square.
            - T-test, ANOVA + Tukey HSD, uji nonparametrik, regresi linear/logistik.
            - Reliabilitas Cronbach's alpha, PCA, EFA/PAF, visualisasi interaktif, dan ekspor output ke Excel/Word/HTML.
            - Mode Pemula/Ahli, halaman Mulai Cepat, pohon keputusan uji, effect size guide, dan glosarium statistik.
            """
        )
    st.stop()

# Tabs
num_cols = numeric_cols(df)
cat_cols = categorical_cols(df)
all_cols = df.columns.tolist()

# Navigasi utama stabil: hanya menu aktif yang dirender
if ui_mode_is_beginner():
    section_labels = ['🚀 Mulai Cepat', '🧙 Smart Assistant', '🧰 Kompatibilitas Data', '📋 Deskriptif', '🧪 Uji Statistik', '📈 Regresi', '🔬 Analisis Lanjutan', '🧭 Reliabilitas & Faktor', '🎨 Visualisasi', '🧠 Insight Riset', '📤 Output & Ekspor', '📚 Panduan']
else:
    section_labels = ['🚀 Mulai Cepat', '🗂️ Data', '🧰 Kompatibilitas Data', '🧙 Smart Assistant', '🔁 Transform', '📋 Deskriptif', '🧪 Uji Statistik', '📈 Regresi', '🔬 Analisis Lanjutan', '🧭 Reliabilitas & Faktor', '🎨 Visualisasi', '🧠 Insight Riset', '📤 Output & Ekspor', '📚 Panduan']
nav_method = getattr(st, "segmented_control", None)
nav_key_suffix = "beginner" if ui_mode_is_beginner() else "expert"
if nav_method is not None:
    try:
        active_section = nav_method("Navigasi utama", section_labels, selection_mode="single", default=section_labels[0], label_visibility="collapsed", key=f"main_section_nav_{nav_key_suffix}")
    except TypeError:
        active_section = st.radio("Navigasi utama", section_labels, horizontal=True, label_visibility="collapsed", key=f"main_section_nav_radio_{nav_key_suffix}")
else:
    active_section = st.radio("Navigasi utama", section_labels, horizontal=True, label_visibility="collapsed", key=f"main_section_nav_radio_{nav_key_suffix}")
if active_section is None or active_section not in section_labels:
    active_section = section_labels[0]

if active_section == '🚀 Mulai Cepat':
    try:
        render_quick_start(df, num_cols, cat_cols, all_cols)
    except BaseException as exc:
        if _is_streamlit_control_exception(exc):
            raise
        st.error("Bagian 🚀 Mulai Cepat mengalami kendala, tetapi aplikasi tetap berjalan.")
        st.exception(exc)

elif active_section == '📚 Panduan':
    try:
        render_reference_center(df)
    except BaseException as exc:
        if _is_streamlit_control_exception(exc):
            raise
        st.error("Bagian 📚 Panduan mengalami kendala, tetapi aplikasi tetap berjalan.")
        st.exception(exc)

elif active_section == '🗂️ Data':
    try:
        st.subheader("🗂️ Data Management")
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Baris", f"{df.shape[0]:,}")
        c2.metric("Kolom", f"{df.shape[1]:,}")
        c3.metric("Numerik", len(num_cols))
        c4.metric("Kategorik/Teks", len(cat_cols))

        st.caption(f"File aktif: {st.session_state.file_name or 'tanpa nama'}")

        with st.expander("✏️ Edit data langsung", expanded=False):
            edited = st.data_editor(df, use_container_width=True, num_rows="dynamic", key="data_editor")
            if st.button("💾 Simpan hasil edit ke sesi"):
                st.session_state.df = edited.copy()
                st.success("Perubahan data disimpan di sesi aplikasi.")
                st.rerun()

        st.markdown("### Data View")
        use_labels = st.checkbox("Tampilkan value labels jika tersedia", value=False)
        preview_df = display_with_value_labels(df, st.session_state.metadata) if use_labels else df
        st.dataframe(preview_df.head(100), use_container_width=True)

        with st.expander("🧾 Variable View ala SPSS", expanded=False):
            st.caption("Atur label variabel, measurement level, value labels, dan missing values khusus. Format value labels: `1=Laki-laki; 2=Perempuan`. Format missing: `99, 999, NA`.")
            metadata = sync_metadata(df)
            edited_meta = st.data_editor(
                metadata,
                use_container_width=True,
                num_rows="fixed",
                column_config={
                    "Measure": st.column_config.SelectboxColumn("Measure", options=["Nominal", "Ordinal", "Scale"]),
                    "Role": st.column_config.SelectboxColumn("Role", options=["Input", "Target", "Both", "None", "Partition", "Split"]),
                    "Decimals": st.column_config.NumberColumn("Decimals", min_value=0, max_value=8, step=1),
                },
                key="metadata_editor",
            )
            c_meta1, c_meta2 = st.columns(2)
            with c_meta1:
                if st.button("💾 Simpan Variable View"):
                    st.session_state.metadata = edited_meta.copy()
                    st.success("Variable View disimpan.")
            with c_meta2:
                if st.button("Terapkan user-missing values ke data"):
                    st.session_state.metadata = edited_meta.copy()
                    st.session_state.df = apply_metadata_to_data(df, edited_meta)
                    log_syntax("MISSING VALUES applied from Variable View.")
                    st.success("Missing values khusus sudah diterapkan sebagai NaN.")
                    st.rerun()

        col_a, col_b = st.columns(2)
        with col_a:
            show_table(
                "Kamus Variabel",
                pd.DataFrame(
                    {
                        "Variabel": all_cols,
                        "Tipe Data": [str(df[c].dtype) for c in all_cols],
                        "Non-Null": [int(df[c].notna().sum()) for c in all_cols],
                        "Missing": [int(df[c].isna().sum()) for c in all_cols],
                        "Unique": [int(df[c].nunique(dropna=True)) for c in all_cols],
                    }
                ),
                save=False,
            )
        with col_b:
            missing = pd.DataFrame(
                {
                    "Variabel": all_cols,
                    "Missing": [df[c].isna().sum() for c in all_cols],
                    "Missing %": [(df[c].isna().mean() * 100) for c in all_cols],
                }
            ).sort_values("Missing %", ascending=False)
            show_table("Ringkasan Missing Values", missing.round(2), save=False)

        with st.expander("🧹 Pembersihan cepat"):
            cleaning_action = st.selectbox(
                "Aksi",
                ["Tidak ada", "Hapus baris dengan missing", "Isi missing numerik dengan mean", "Isi missing numerik dengan median", "Konversi kolom ke numerik"],
            )
            target_cols = st.multiselect("Pilih kolom", all_cols, default=num_cols[: min(3, len(num_cols))])
            if st.button("Terapkan pembersihan"):
                new_df = df.copy()
                if cleaning_action == "Hapus baris dengan missing":
                    new_df = new_df.dropna(subset=target_cols if target_cols else None)
                elif cleaning_action == "Isi missing numerik dengan mean":
                    for c in target_cols:
                        new_df[c] = safe_numeric(new_df[c])
                        new_df[c] = new_df[c].fillna(new_df[c].mean())
                elif cleaning_action == "Isi missing numerik dengan median":
                    for c in target_cols:
                        new_df[c] = safe_numeric(new_df[c])
                        new_df[c] = new_df[c].fillna(new_df[c].median())
                elif cleaning_action == "Konversi kolom ke numerik":
                    for c in target_cols:
                        new_df[c] = safe_numeric(new_df[c])
                st.session_state.df = new_df
                st.success("✅ Pembersihan diterapkan.")
                st.rerun()

    except BaseException as exc:
        if _is_streamlit_control_exception(exc):
            raise
        st.error("Bagian 🗂️ Data mengalami kendala, tetapi aplikasi tetap berjalan.")
        st.exception(exc)


elif active_section == '🧰 Kompatibilitas Data':
    try:
        st.subheader("🧰 Cek Kompatibilitas Data & Panduan Perbaikan")
        st.caption("Bagian ini dibuat untuk pengguna awam: aplikasi mendiagnosis data, menjelaskan masalahnya, lalu memberi langkah perbaikan yang praktis sebelum analisis statistik dijalankan.")

        issues = analyze_data_compatibility(df, st.session_state.get("metadata"))
        profile = _column_profile(df, st.session_state.get("metadata"))
        suggestions = suggest_compatible_analyses(df)
        score = compatibility_score(issues)

        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Skor kesiapan", f"{score}/100")
        c2.metric("Masalah kritis/tinggi", int(issues["Prioritas"].isin(["Kritis", "Tinggi"]).sum()))
        c3.metric("Kolom numeric", len(num_cols))
        c4.metric("Kolom kategori/teks", len(cat_cols))

        if score >= 80:
            st.success("✅ Data cukup siap untuk eksplorasi dan sebagian besar analisis dasar. Tetap cek asumsi sesuai uji yang dipilih.")
        elif score >= 60:
            st.warning("⚠️ Data bisa dianalisis, tetapi ada beberapa hal yang sebaiknya dirapikan agar hasil lebih mudah dimaknai.")
        else:
            st.error("🚧 Data belum ideal. Ikuti rekomendasi perbaikan di bawah agar analisis tidak salah baca.")

        st.markdown("### 1) Diagnosis masalah data")
        priority_filter = st.multiselect(
            "Tampilkan prioritas",
            ["Kritis", "Tinggi", "Sedang", "Ringan", "Info", "OK"],
            default=["Kritis", "Tinggi", "Sedang", "Ringan", "Info"],
            key="compat_priority_filter",
        )
        view_issues = issues[issues["Prioritas"].isin(priority_filter)] if priority_filter else issues
        st.dataframe(view_issues, use_container_width=True, hide_index=True)

        st.markdown("### 2) Apa yang sebaiknya dilakukan dulu?")
        recipe = beginner_data_recipe(df, issues)
        st.dataframe(recipe, use_container_width=True, hide_index=True)

        st.markdown("### 3) Profil setiap kolom")
        st.caption("Gunakan tabel ini untuk mengetahui apakah kolom terbaca sebagai angka, kategori, tanggal, ID, kosong, atau perlu diganti tipe datanya.")
        st.dataframe(profile, use_container_width=True, hide_index=True)

        st.markdown("### 4) Uji statistik apa yang cocok dengan data ini?")
        st.dataframe(suggestions, use_container_width=True, hide_index=True)

        st.markdown("### 5) Checklist khusus sebelum menjalankan analisis")
        chosen_analysis = st.selectbox(
            "Pilih analisis yang ingin dicek",
            ["Independent T-Test", "ANOVA", "Korelasi", "Regresi Linear", "Reliabilitas / Cronbach Alpha", "EFA"],
            key="compat_analysis_select",
        )
        st.dataframe(analysis_specific_guidance(df, chosen_analysis), use_container_width=True, hide_index=True)

        st.markdown("### 6) Format data yang paling ramah untuk aplikasi")
        st.info(
            """
            **Aturan sederhana:** 1 baris = 1 responden/observasi, 1 kolom = 1 variabel.  
            Gunakan nama kolom pendek tanpa spasi, misalnya `jenis_kelamin`, `usia`, `skor_pre`, `skor_post`, `motivasi_1`.  
            Untuk angka, jangan campur dengan teks seperti `Rp 10.000`, `80%`, atau `tinggi=5`. Pisahkan label kategori lewat **Variable View**.
            """
        )

        export_df = pd.concat(
            {
                "diagnosis": issues.reset_index(drop=True),
                "profil_kolom": profile.reset_index(drop=True),
                "saran_analisis": suggestions.reset_index(drop=True),
            },
            names=["Bagian", "No"],
        ).reset_index(level=0)
        st.download_button(
            "⬇️ Download Laporan Kompatibilitas CSV",
            data=export_df.to_csv(index=False).encode("utf-8"),
            file_name="laporan_kompatibilitas_data.csv",
            mime="text/csv",
            key="download_compat_csv",
        )

        if st.button("💾 Simpan Diagnosis ke Output Viewer", key="save_compat_report"):
            add_report("Diagnosis Kompatibilitas Data", issues, f"Skor kesiapan data: {score}/100. Gunakan rekomendasi ini sebelum analisis inferensial.")
            add_report("Saran Analisis yang Cocok", suggestions, "Daftar analisis yang cocok/tidak cocok berdasarkan struktur data aktif.")
            st.success("Diagnosis kompatibilitas disimpan ke Output Viewer.")

    except BaseException as exc:
        if _is_streamlit_control_exception(exc):
            raise
        st.error("Bagian 🧰 Kompatibilitas Data mengalami kendala, tetapi aplikasi tetap berjalan.")
        st.exception(exc)


elif active_section == '🧙 Smart Assistant':
    try:
        st.subheader("🧙 Smart Statistical Assistant")
        st.caption("Mode ini memandu user awam: cek kualitas data, pilih tujuan riset, rekomendasi uji, perbaiki data, hitung sampel, dan buat narasi laporan.")

        issues_now = analyze_data_compatibility(df, st.session_state.metadata)
        score_now = compatibility_score(issues_now)
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Skor kesiapan data", f"{score_now}/100")
        c2.metric("Baris", f"{df.shape[0]:,}")
        c3.metric("Variabel numerik", len(num_cols))
        c4.metric("Variabel kategori", len(cat_cols))

        mode = st.radio(
            "Pilih alat bantu",
            ["Wizard Uji Otomatis", "Research Design Planner", "Data Repair Assistant", "Sample Size & Power", "Kalkulator Statistik", "Effect Size & Asumsi", "Template Narasi Laporan"],
            horizontal=True,
            key="smart_assistant_mode",
        )

        if mode == "Wizard Uji Otomatis":
            st.markdown("### 🧭 Wizard Pemilihan Uji")
            st.info("Pilih tujuan riset dan variabel. Aplikasi akan menyarankan uji yang paling masuk akal, alternatif jika asumsi lemah, dan langkah berikutnya.")
            objective = st.selectbox(
                "Tujuan riset/analisis",
                [
                    "Membandingkan rata-rata antar kelompok",
                    "Melihat hubungan antar variabel",
                    "Memprediksi variabel hasil",
                    "Menguji kuesioner/skala",
                    "Menganalisis data kategori",
                    "Eksplorasi awal/deskriptif",
                ],
                key="wizard_objective",
            )
            col_a, col_b, col_c = st.columns(3)
            with col_a:
                y_sel = st.selectbox("Variabel hasil / Y / skor", [None] + all_cols, format_func=lambda x: "— pilih —" if x is None else str(x), key="wizard_y")
            with col_b:
                x_sel = st.selectbox("Variabel X / prediktor / pembanding", [None] + all_cols, format_func=lambda x: "— pilih —" if x is None else str(x), key="wizard_x")
            with col_c:
                group_sel = st.selectbox("Variabel grup/kategori", [None] + all_cols, format_func=lambda x: "— pilih —" if x is None else str(x), key="wizard_group")

            type_rows = pd.DataFrame([
                {"Pilihan": "Y/skor", "Kolom": y_sel or "-", "Tipe terbaca": classify_variable(df, y_sel) if y_sel else "-"},
                {"Pilihan": "X/prediktor", "Kolom": x_sel or "-", "Tipe terbaca": classify_variable(df, x_sel) if x_sel else "-"},
                {"Pilihan": "Grup", "Kolom": group_sel or "-", "Tipe terbaca": classify_variable(df, group_sel) if group_sel else "-"},
            ])
            st.dataframe(type_rows, use_container_width=True, hide_index=True)

            rec = smart_test_recommendation(df, objective, y=y_sel, x=x_sel, group=group_sel)
            st.markdown("#### Rekomendasi uji")
            st.dataframe(rec, use_container_width=True, hide_index=True)

            effect_table = compute_quick_effect_sizes(df, y=y_sel, x=x_sel, group=group_sel)
            if not effect_table.empty:
                st.markdown("#### Effect size cepat dari variabel terpilih")
                st.dataframe(effect_table, use_container_width=True, hide_index=True)
                st.caption("Effect size membantu menjawab: hasil ini sekadar signifikan, atau benar-benar bermakna secara praktis?")

            st.markdown("#### Checklist sebelum menjalankan uji")
            if objective == "Membandingkan rata-rata antar kelompok":
                checklist_name = "ANOVA" if group_sel and group_sel in df.columns and df[group_sel].dropna().nunique() >= 3 else "Independent T-Test"
            elif objective == "Melihat hubungan antar variabel":
                checklist_name = "Korelasi"
            elif objective == "Memprediksi variabel hasil":
                checklist_name = "Regresi Linear"
            elif objective == "Menguji kuesioner/skala":
                checklist_name = "Reliabilitas / Cronbach Alpha"
            else:
                checklist_name = "Deskriptif"
            guide = analysis_specific_guidance(df, checklist_name)
            st.dataframe(guide, use_container_width=True, hide_index=True)

            if st.button("💾 Simpan rekomendasi wizard ke Output Viewer", key="save_wizard_recommendation"):
                merged = pd.concat([rec, guide.rename(columns={"Checklist": "Uji/Analisis Disarankan", "Status": "Kondisi Data", "Apa yang perlu dilakukan": "Langkah Berikutnya"})], ignore_index=True, sort=False)
                add_report("Rekomendasi Smart Wizard", merged, f"Tujuan: {objective}. Skor kesiapan data: {score_now}/100.")
                st.success("Rekomendasi disimpan ke Output Viewer.")

        elif mode == "Research Design Planner":
            st.markdown("### 🧭 Research Design Planner")
            st.caption("Bantu merapikan desain riset sebelum analisis: pertanyaan riset, desain data, variabel, sampel, asumsi, dan laporan.")
            rd1, rd2 = st.columns(2)
            with rd1:
                design = st.selectbox(
                    "Desain penelitian/data",
                    ["Survei/cross-sectional", "Eksperimen dua kelompok", "Pretest-posttest", "Korelasi/regresi", "Kuesioner/skala", "Data kategori", "Longitudinal/time series sederhana"],
                    key="research_design_type",
                )
            with rd2:
                objective_plan = st.selectbox(
                    "Tujuan utama",
                    ["Deskriptif", "Membandingkan kelompok", "Melihat hubungan", "Memprediksi hasil", "Menguji instrumen", "Mencari faktor/dimensi"],
                    key="research_design_objective",
                )
            context_note = st.text_area("Konteks sampel/populasi", placeholder="Contoh: 120 mahasiswa semester 3, data dikumpulkan melalui kuesioner Likert.", key="research_design_context")
            plan_df = research_design_planner_table(design, objective_plan, context_note)
            st.dataframe(plan_df, use_container_width=True, hide_index=True)
            st.markdown("#### Variabel yang tampaknya cocok dari dataset")
            st.dataframe(variable_role_suggestions(df), use_container_width=True, hide_index=True)
            if st.button("💾 Simpan rencana desain riset", key="save_research_design_plan"):
                add_report("Research Design Planner", plan_df, f"Desain: {design}. Tujuan: {objective_plan}.")
                st.success("Rencana desain riset disimpan ke Output Viewer.")

        elif mode == "Data Repair Assistant":
            st.markdown("### 🛠️ Data Repair Assistant")
            st.caption("Aplikasi mendeteksi masalah umum lalu menawarkan perbaikan satu klik. Data asli tetap ada di sesi sampai tombol diterapkan.")
            actions = detect_repair_actions(df)
            if actions.empty:
                st.success("Tidak ditemukan tindakan perbaikan otomatis yang mendesak. Lanjutkan ke wizard atau analisis.")
            else:
                st.dataframe(actions.drop(columns=["Kode"]), use_container_width=True, hide_index=True)
                st.warning("Saran aman: simpan file asli dulu. Perbaikan otomatis bagus untuk masalah format umum, tetapi tetap cek konteks riset.")
                for _, row in actions.iterrows():
                    action_code = row["Kode"]
                    if st.button(f"Terapkan: {row['Tindakan']}", key=f"repair_{action_code}"):
                        repaired, note = apply_repair_action(df, action_code)
                        st.session_state.df = repaired
                        st.session_state.metadata = build_metadata(repaired)
                        log_syntax(f"SMART REPAIR /ACTION={action_code}.")
                        add_report("Data Repair Assistant", pd.DataFrame([{"Aksi": action_code, "Catatan": note}]), note)
                        st.success(note)
                        st.rerun()

            with st.expander("Preview profil kolom setelah/ sebelum perbaikan", expanded=False):
                st.dataframe(_column_profile(df, st.session_state.metadata), use_container_width=True, hide_index=True)

        elif mode == "Sample Size & Power":
            st.markdown("### 📏 Sample Size & Power Calculator")
            st.caption("Gunakan untuk proposal riset atau mengecek apakah sampel kira-kira memadai. Untuk desain kompleks, tetap konsultasikan metode statistik.")
            test_type = st.selectbox(
                "Jenis perhitungan",
                ["Independent t-test", "One-sample t-test", "Paired t-test", "One-way ANOVA", "Correlation", "Multiple regression", "Survey proportion / margin of error"],
                key="power_test_type",
            )
            p1, p2, p3, p4 = st.columns(4)
            with p1:
                alpha_power = st.number_input("Alpha", min_value=0.001, max_value=0.20, value=0.05, step=0.01, format="%.3f", key="power_alpha")
            with p2:
                target_power = st.number_input("Power", min_value=0.50, max_value=0.99, value=0.80, step=0.01, format="%.2f", key="power_value")
            with p3:
                es_default = 0.5 if test_type != "One-way ANOVA" else 0.25
                es = st.number_input("Effect size", min_value=0.001, max_value=5.0, value=es_default, step=0.05, format="%.3f", key="power_effect")
            with p4:
                groups_or_pred = st.number_input("Jumlah grup/prediktor", min_value=1, max_value=50, value=3 if test_type == "One-way ANOVA" else 2, step=1, key="power_groups_pred")
            r_val = None
            moe_val = None
            prop_val = 0.5
            if test_type == "Correlation":
                r_val = st.number_input("Target korelasi r", min_value=0.01, max_value=0.99, value=0.30, step=0.01, format="%.2f", key="power_r")
            if test_type == "Survey proportion / margin of error":
                moe_val = st.number_input("Margin of error", min_value=0.001, max_value=0.50, value=0.05, step=0.01, format="%.3f", key="power_moe")
                prop_val = st.number_input("Proporsi awal p", min_value=0.01, max_value=0.99, value=0.50, step=0.01, format="%.2f", key="power_prop")
            n_needed, unit, method = normal_power_sample_size(test_type, alpha_power, target_power, es, groups=groups_or_pred, predictors=groups_or_pred, r=r_val, moe=moe_val, proportion=prop_val)
            if n_needed:
                st.metric("Sampel minimal perkiraan", f"{n_needed:,} {unit}")
                st.caption(method)
                ss_table = pd.DataFrame([{"Jenis": test_type, "Alpha": alpha_power, "Power": target_power, "Effect size": es, "N minimal": n_needed, "Satuan": unit, "Metode": method}])
                st.dataframe(ss_table, use_container_width=True, hide_index=True)
                if st.button("💾 Simpan perhitungan sample size", key="save_power_calc"):
                    add_report("Sample Size & Power Calculator", ss_table, method)
                    st.success("Perhitungan disimpan ke Output Viewer.")
            else:
                st.error(method)
            st.info("Panduan effect size cepat: Cohen's d 0.2 kecil, 0.5 sedang, 0.8 besar; r 0.1 kecil, 0.3 sedang, 0.5 besar; Cohen's f ANOVA 0.10 kecil, 0.25 sedang, 0.40 besar.")

        elif mode == "Kalkulator Statistik":
            st.markdown("### 🧮 Kalkulator Statistik Manual")
            calc_mode = st.radio("Jenis kalkulator", ["Deskriptif dari daftar angka", "Distribusi & nilai kritis", "Z-score/T-score"], horizontal=True, key="calc_mode")
            if calc_mode == "Deskriptif dari daftar angka":
                numbers_text = st.text_area("Masukkan angka, pisahkan dengan koma/spasi/baris", "12, 15, 14, 18, 20, 17, 16", height=120, key="manual_numbers")
                vals = parse_number_list(numbers_text)
                desc_calc = descriptive_calculator_table(vals)
                st.dataframe(desc_calc, use_container_width=True, hide_index=True)
                if st.button("💾 Simpan hasil kalkulator deskriptif", key="save_desc_calc"):
                    add_report("Kalkulator Deskriptif Manual", desc_calc, "Statistik deskriptif dari daftar angka manual.")
                    st.success("Hasil disimpan ke Output Viewer.")
            elif calc_mode == "Distribusi & nilai kritis":
                d1, d2, d3, d4 = st.columns(4)
                with d1:
                    dist_name = st.selectbox("Distribusi", ["Normal/Z", "t", "Chi-square", "F", "Binomial", "Poisson"], key="dist_name")
                with d2:
                    dist_mode = st.selectbox("Hitung", ["P(X ≤ nilai)", "P(X ≥ nilai)", "Nilai kritis kanan/dua sisi"], key="dist_mode")
                with d3:
                    value_x = st.number_input("Nilai X", value=1.96, step=0.1, format="%.4f", key="dist_x")
                with d4:
                    alpha_dist = st.number_input("Alpha", min_value=0.001, max_value=0.50, value=0.05, step=0.01, format="%.3f", key="dist_alpha")
                e1, e2, e3 = st.columns(3)
                with e1:
                    df1_val = st.number_input("df / df1", min_value=1, max_value=100000, value=10, step=1, key="dist_df1")
                with e2:
                    df2_val = st.number_input("df2", min_value=1, max_value=100000, value=20, step=1, key="dist_df2")
                with e3:
                    n_bin = st.number_input("n binomial", min_value=1, max_value=100000, value=20, step=1, key="dist_n")
                f1, f2 = st.columns(2)
                with f1:
                    p_bin = st.number_input("p binomial", min_value=0.001, max_value=0.999, value=0.50, step=0.01, format="%.3f", key="dist_p")
                with f2:
                    lam_pois = st.number_input("lambda poisson", min_value=0.001, max_value=100000.0, value=3.0, step=0.5, format="%.3f", key="dist_lambda")
                result = distribution_calculator(dist_name, dist_mode, value_x, df1=df1_val, df2=df2_val, n=n_bin, p=p_bin, lam=lam_pois, alpha=alpha_dist)
                st.metric("Hasil", f"{result:.6f}" if pd.notna(result) else "NA")
            else:
                zc1, zc2, zc3 = st.columns(3)
                with zc1:
                    raw_x = st.number_input("Nilai X", value=75.0, step=1.0, key="z_x")
                with zc2:
                    mean_x = st.number_input("Mean", value=70.0, step=1.0, key="z_mean")
                with zc3:
                    sd_x = st.number_input("Std. deviasi", min_value=0.0001, value=10.0, step=1.0, key="z_sd")
                z = (raw_x - mean_x) / sd_x
                t_score = 50 + 10*z
                pct = stats.norm.cdf(z) * 100
                z_table = pd.DataFrame([{"X": raw_x, "Mean": mean_x, "SD": sd_x, "Z-score": z, "T-score": t_score, "Persentil approx": pct}])
                st.dataframe(z_table, use_container_width=True, hide_index=True)

        elif mode == "Effect Size & Asumsi":
            st.markdown("### 📐 Effect Size & Asumsi")
            st.caption("Panduan ringkas untuk memastikan hasil tidak hanya signifikan, tetapi juga bermakna dan layak dilaporkan.")
            sub_mode = st.radio("Pilih panduan", ["Effect Size", "Asumsi & Solusi", "Normalitas & Distribusi", "Pohon Keputusan Uji"], horizontal=True, key="effect_assumption_submode")
            if sub_mode == "Effect Size":
                st.dataframe(effect_size_reference_table(), use_container_width=True, hide_index=True)
                st.info("Gunakan effect size untuk menjawab besar dampak. p-value menjawab ada/tidaknya bukti statistik; effect size menjawab seberapa berarti temuannya.")
            elif sub_mode == "Asumsi & Solusi":
                st.dataframe(assumption_playbook_table(), use_container_width=True, hide_index=True)
            elif sub_mode == "Normalitas & Distribusi":
                st.markdown("#### Makna skewness dan kurtosis")
                st.dataframe(skewness_kurtosis_reference_table(), use_container_width=True, hide_index=True)
                st.markdown("#### Tindakan jika normalitas tidak terpenuhi")
                st.dataframe(normality_solution_table(), use_container_width=True, hide_index=True)
            else:
                st.dataframe(analysis_decision_matrix(df), use_container_width=True, hide_index=True)
            if st.button("💾 Simpan panduan effect/asumsi", key="save_effect_assumption_guide"):
                table = effect_size_reference_table() if sub_mode == "Effect Size" else assumption_playbook_table() if sub_mode == "Asumsi & Solusi" else normality_solution_table() if sub_mode == "Normalitas & Distribusi" else analysis_decision_matrix(df)
                add_report(f"Panduan {sub_mode}", table, "Panduan interpretasi dan pemilihan analisis statistik.")
                st.success("Panduan disimpan ke Output Viewer.")

        else:
            st.markdown("### 📝 Template Narasi Laporan")
            report_kind = st.selectbox("Jenis template", ["BAB 4 Skripsi/Tesis", "APA Style", "Ringkasan Manajerial"], key="report_template_kind")
            research_title = st.text_input("Judul riset", value="", placeholder="Contoh: Pengaruh Motivasi Belajar terhadap Prestasi Siswa", key="research_title_input")
            hypothesis = st.text_area("Pertanyaan/hipotesis riset", value="", placeholder="Contoh: Motivasi belajar berhubungan positif dengan prestasi siswa.", key="hypothesis_input")
            output_titles = [item.get("title", f"Output {i+1}") for i, item in enumerate(st.session_state.report_items)]
            selected_outputs = st.multiselect("Output yang ingin dijadikan dasar narasi", output_titles, default=output_titles[-3:] if output_titles else [], key="template_outputs")
            template_md = build_report_template(report_kind, research_title, hypothesis, selected_outputs)
            st.markdown(template_md)
            st.download_button("⬇️ Download template narasi Markdown", template_md.encode("utf-8"), file_name="template_narasi_riset.md", mime="text/markdown", key="download_template_md")
            if st.button("💾 Simpan template narasi ke Output Viewer", key="save_template_report"):
                add_report("Template Narasi Laporan", pd.DataFrame([{"Jenis": report_kind, "Narasi": template_md}]), "Template narasi otomatis untuk membantu penyusunan laporan penelitian.")
                st.success("Template narasi disimpan ke Output Viewer.")

    except Exception as exc:
        if _is_streamlit_control_exception(exc):
            raise
        st.error("Bagian 🧙 Smart Assistant mengalami kendala, tetapi aplikasi tetap berjalan.")
        st.exception(exc)

elif active_section == '🔁 Transform':
    try:
        st.subheader("🔁 Transformasi Data")
        st.caption("Fitur ini meniru menu Transform/Data di SPSS: compute, recode, reverse coding, standardize, filter/select cases, dan split file.")

        transform_type = st.selectbox(
            "Pilih transformasi",
            [
                "Compute Variable",
                "Recode into Different Variable",
                "Reverse Coding Likert",
                "Standardize / Z-score",
                "Filter / Select Cases",
                "Split File untuk Output",
                "Rename / Drop Variables",
            ],
        )

        if transform_type == "Compute Variable":
            new_name = st.text_input("Nama variabel baru", "total_skor")
            formula = st.text_input("Formula pandas", "item_1 + item_2 + item_3 + item_4")
            st.caption("Gunakan nama kolom persis seperti di data. Contoh: `motivasi - kecemasan`, `(item_1 + item_2)/2`, atau `np.log(nilai_akhir)`.")
            if st.button("Buat variabel", type="primary"):
                try:
                    new_df = df.copy()
                    safe_name = safe_col_name(new_name, new_df.columns)
                    local_dict = {c: new_df[c] for c in new_df.columns}
                    local_dict.update({"np": np, "pd": pd})
                    new_df[safe_name] = eval(formula, {"__builtins__": {}}, local_dict)
                    st.session_state.df = new_df
                    st.session_state.metadata = build_metadata(new_df) if st.session_state.metadata is None else sync_metadata(new_df)
                    log_syntax(f"COMPUTE {safe_name} = {formula}.")
                    st.success(f"Variabel `{safe_name}` berhasil dibuat.")
                    st.rerun()
                except Exception as exc:
                    st.error(f"Formula belum valid: {exc}")

        elif transform_type == "Recode into Different Variable":
            col = st.selectbox("Variabel sumber", all_cols)
            new_name = st.text_input("Nama variabel hasil recode", f"{col}_recode")
            rules = st.text_area("Aturan recode", "1=Rendah\n2=Sedang\n3=Tinggi\n4:5=Sangat Tinggi", height=150)
            default = st.radio("Nilai lain", ["copy", "missing"], horizontal=True, format_func=lambda x: "Salin nilai asli" if x == "copy" else "Jadikan missing")
            if st.button("Jalankan Recode", type="primary"):
                try:
                    new_df = df.copy()
                    safe_name = safe_col_name(new_name, new_df.columns)
                    new_df[safe_name] = recode_series(new_df[col], rules, default=default)
                    st.session_state.df = new_df
                    st.session_state.metadata = sync_metadata(new_df)
                    log_syntax(f"RECODE {col} INTO {safe_name} /RULES='{rules.replace(chr(10), '; ')}'.")
                    st.success(f"Recode selesai: `{safe_name}` ditambahkan.")
                    st.rerun()
                except Exception as exc:
                    st.error(f"Recode gagal: {exc}")

        elif transform_type == "Reverse Coding Likert":
            cols = st.multiselect("Pilih item yang perlu dibalik", num_cols, default=[c for c in num_cols if "rev" in c.lower()][:5])
            c1, c2 = st.columns(2)
            with c1:
                min_score = st.number_input("Skor minimum", value=1.0)
            with c2:
                max_score = st.number_input("Skor maksimum", value=5.0)
            prefix = st.text_input("Prefix variabel baru", "rev_")
            if st.button("Buat reverse-coded item", type="primary"):
                if not cols:
                    st.error("Pilih minimal 1 item.")
                else:
                    new_df = df.copy()
                    for c in cols:
                        new_col = safe_col_name(f"{prefix}{c}", new_df.columns)
                        new_df[new_col] = (min_score + max_score) - safe_numeric(new_df[c])
                    st.session_state.df = new_df
                    st.session_state.metadata = sync_metadata(new_df)
                    log_syntax(f"RECODE reverse {cols} range {min_score}-{max_score}.")
                    st.success("Reverse coding selesai.")
                    st.rerun()

        elif transform_type == "Standardize / Z-score":
            cols = st.multiselect("Variabel numerik", num_cols, default=num_cols[: min(4, len(num_cols))])
            if st.button("Buat Z-score", type="primary"):
                if not cols:
                    st.error("Pilih minimal 1 variabel numerik.")
                else:
                    new_df = df.copy()
                    for c in cols:
                        s_num = safe_numeric(new_df[c])
                        sd = s_num.std(ddof=1)
                        new_df[safe_col_name(f"z_{c}", new_df.columns)] = (s_num - s_num.mean()) / sd if sd != 0 else np.nan
                    st.session_state.df = new_df
                    st.session_state.metadata = sync_metadata(new_df)
                    log_syntax(f"DESCRIPTIVES VARIABLES={','.join(cols)} /SAVE ZSCORES.")
                    st.success("Z-score berhasil dibuat.")
                    st.rerun()

        elif transform_type == "Filter / Select Cases":
            query = st.text_input("Kondisi filter pandas", "kelompok == 'A' or nilai_akhir >= 75")
            st.caption("Contoh: `gender == 'Wanita'`, `usia >= 18`, `kelompok in ['A','B']`. Gunakan nama kolom yang valid.")
            overwrite = st.checkbox("Terapkan sebagai data aktif", value=False)
            if st.button("Preview / Terapkan Filter", type="primary"):
                try:
                    filtered = df.query(query)
                    st.success(f"Hasil filter: {len(filtered)} dari {len(df)} baris.")
                    st.dataframe(filtered.head(100), use_container_width=True)
                    if overwrite:
                        st.session_state.df = filtered.copy()
                        st.session_state.metadata = sync_metadata(filtered)
                        log_syntax(f"SELECT IF ({query}).")
                        st.rerun()
                except Exception as exc:
                    st.error(f"Kondisi filter belum valid: {exc}")

        elif transform_type == "Split File untuk Output":
            split = st.selectbox("Kelompokkan output berdasarkan", ["(tidak ada)"] + all_cols, index=0)
            if st.button("Set Split File", type="primary"):
                st.session_state.split_by = split
                log_syntax("SPLIT FILE OFF." if split == "(tidak ada)" else f"SPLIT FILE BY {split}.")
                st.success(f"Split file diset ke: {split}")
            if st.session_state.get("split_by", "(tidak ada)") != "(tidak ada)":
                st.info(f"Split aktif: output deskriptif tertentu dapat diringkas per `{st.session_state.split_by}`.")

        else:
            action = st.radio("Aksi", ["Rename", "Drop"], horizontal=True)
            if action == "Rename":
                col = st.selectbox("Variabel lama", all_cols)
                new_name = st.text_input("Nama baru", col)
                if st.button("Rename Variable", type="primary"):
                    safe_name = safe_col_name(new_name, [c for c in df.columns if c != col])
                    new_df = df.rename(columns={col: safe_name})
                    st.session_state.df = new_df
                    st.session_state.metadata = build_metadata(new_df)
                    log_syntax(f"RENAME VARIABLES ({col}={safe_name}).")
                    st.success("Variabel berhasil diganti nama.")
                    st.rerun()
            else:
                cols = st.multiselect("Variabel yang akan dihapus", all_cols)
                if st.button("Drop Variables", type="primary"):
                    if not cols:
                        st.error("Pilih minimal 1 variabel.")
                    else:
                        new_df = df.drop(columns=cols)
                        st.session_state.df = new_df
                        st.session_state.metadata = sync_metadata(new_df)
                        log_syntax(f"DROP VARIABLES {','.join(cols)}.")
                        st.success("Variabel berhasil dihapus.")
                        st.rerun()

    except BaseException as exc:
        if _is_streamlit_control_exception(exc):
            raise
        st.error("Bagian 🔁 Transform mengalami kendala, tetapi aplikasi tetap berjalan.")
        st.exception(exc)

elif active_section == '📋 Deskriptif':
    try:
        st.subheader("📋 Statistik Deskriptif")
        if not num_cols and not cat_cols:
            st.warning("Data belum memiliki variabel yang bisa dianalisis.")
        else:
            desc_cols = st.multiselect("Variabel numerik", num_cols, default=num_cols[: min(6, len(num_cols))])
            if desc_cols and st.button("Hitung Deskriptif", type="primary"):
                show_table("Statistik Deskriptif + Makna Skewness/Kurtosis", descriptive_table(df, desc_cols))
                show_table("Uji Normalitas + Rekomendasi Tindakan", normality_table(df, desc_cols), "p-value < α mengindikasikan penyimpangan dari normalitas. Baca kolom tindakan untuk memilih transformasi, uji alternatif, atau cek outlier.")
                with st.expander("📌 Cara memaknai skewness, kurtosis, dan normalitas", expanded=detail_is_full()):
                    st.dataframe(skewness_kurtosis_reference_table(), use_container_width=True, hide_index=True)
                    st.dataframe(normality_solution_table(), use_container_width=True, hide_index=True)

            st.markdown("### Frekuensi Kategori")
            freq_col = st.selectbox("Pilih variabel", all_cols)
            if st.button("Buat Tabel Frekuensi"):
                show_table(f"Tabel Frekuensi: {freq_col}", frequency_table(df, freq_col))

            st.markdown("### Explore by Group")
            if num_cols and all_cols:
                dv = st.selectbox("Variabel numerik", num_cols, key="desc_group_dv")
                grp = st.selectbox("Kelompok", all_cols, key="desc_group_grp")
                if st.button("Ringkas per Kelompok"):
                    grouped = df.assign(**{dv: safe_numeric(df[dv])}).groupby(grp, dropna=False)[dv].agg(
                        N="count", Mean="mean", Median="median", SD="std", Min="min", Max="max"
                    ).reset_index()
                    show_table(f"Deskriptif {dv} berdasarkan {grp}", grouped.round(4))

    except BaseException as exc:
        if _is_streamlit_control_exception(exc):
            raise
        st.error("Bagian 📋 Deskriptif mengalami kendala, tetapi aplikasi tetap berjalan.")
        st.exception(exc)

elif active_section == '🧪 Uji Statistik':
    try:
        st.subheader("🧪 Uji Statistik")
        alpha = st.slider("Tingkat signifikansi (α)", 0.01, 0.10, 0.05, 0.01, key="alpha_tests")
        st.session_state.active_alpha = alpha
        analysis_options = ["Crosstab & Chi-Square"]
        if num_cols:
            analysis_options = [
                "One-Sample T-Test",
                "Independent T-Test",
                "Paired T-Test",
                "One-Way ANOVA",
                "Two-Way ANOVA",
                "Korelasi",
                "Crosstab & Chi-Square",
                "Uji Nonparametrik",
                "Uji Asumsi",
            ]
        test_type = st.selectbox("Pilih analisis", analysis_options)

        if test_type == "One-Sample T-Test":
            col = st.selectbox("Variabel sampel", num_cols)
            mu = st.number_input("Nilai mean hipotesis (μ₀)", value=0.0)
            if st.button("Jalankan One-Sample T-Test", type="primary"):
                show_table(f"One-Sample T-Test: {col}", one_sample_ttest(safe_numeric(df[col]), mu, alpha))

        elif test_type == "Independent T-Test":
            mode = st.radio("Format data", ["Wide: dua kolom sampel", "Long: DV + kolom grup"], horizontal=True)
            equal_var = st.checkbox("Asumsi varian sama (Student). Matikan untuk Welch.", value=False)
            if mode.startswith("Wide"):
                c1, c2 = st.columns(2)
                with c1:
                    col1 = st.selectbox("Sampel/Grup 1", num_cols, index=0 if num_cols else None)
                with c2:
                    col2 = st.selectbox("Sampel/Grup 2", num_cols, index=1 if len(num_cols) > 1 else 0)
                if st.button("Jalankan Independent T-Test", type="primary"):
                    if col1 == col2:
                        st.error("Pilih dua kolom berbeda.")
                    else:
                        show_table(
                            f"Independent T-Test: {col1} vs {col2}",
                            independent_ttest(safe_numeric(df[col1]), safe_numeric(df[col2]), col1, col2, alpha, equal_var),
                        )
            else:
                dv = st.selectbox("Dependent variable / nilai", num_cols, key="tt_dv")
                grp = st.selectbox("Kolom grup", all_cols, key="tt_grp")
                levels = df[grp].dropna().astype(str).unique().tolist()
                selected_levels = st.multiselect("Pilih tepat 2 kategori", levels, default=levels[:2])
                if st.button("Jalankan Independent T-Test Long", type="primary"):
                    if len(selected_levels) != 2:
                        st.error("Pilih tepat 2 kategori.")
                    else:
                        s1 = safe_numeric(df.loc[df[grp].astype(str) == selected_levels[0], dv])
                        s2 = safe_numeric(df.loc[df[grp].astype(str) == selected_levels[1], dv])
                        show_table(
                            f"Independent T-Test: {dv} berdasarkan {grp}",
                            independent_ttest(s1, s2, selected_levels[0], selected_levels[1], alpha, equal_var),
                        )

        elif test_type == "Paired T-Test":
            c1, c2 = st.columns(2)
            with c1:
                before = st.selectbox("Before/Pre", num_cols)
            with c2:
                after = st.selectbox("After/Post", num_cols, index=1 if len(num_cols) > 1 else 0)
            if st.button("Jalankan Paired T-Test", type="primary"):
                if before == after:
                    st.error("Pilih dua kolom berbeda.")
                else:
                    show_table(f"Paired T-Test: {before} vs {after}", paired_ttest(safe_numeric(df[before]), safe_numeric(df[after]), before, after, alpha))

        elif test_type == "One-Way ANOVA":
            mode = st.radio("Format data", ["Wide: setiap kolom = grup", "Long: DV + kolom grup"], horizontal=True)
            long_for_posthoc = None
            if mode.startswith("Wide"):
                groups = st.multiselect("Pilih kolom grup", num_cols, default=num_cols[: min(3, len(num_cols))])
                if st.button("Jalankan ANOVA", type="primary"):
                    if len(groups) < 2:
                        st.error("Minimal 2 grup.")
                    else:
                        table, effects, long_for_posthoc = anova_wide(df, groups, alpha)
                        show_table("ANOVA Table", table)
                        show_table("Effect Size & Catatan Unequal N", effects)
                        with st.expander("📏 Ukuran grup dan opsi jika jumlah data tidak sama", expanded=effects.iloc[0].get("Status N") != "Seimbang"):
                            group_map = {g: safe_numeric(df[g]).dropna() for g in groups}
                            show_table("Ringkasan N per Grup", group_size_summary(group_map), "ANOVA biasa dapat menghitung jumlah grup yang tidak sama, tetapi interpretasi perlu melihat homogenitas varians.")
                            show_table("Welch ANOVA — alternatif aman untuk unequal N/varians", welch_anova_from_groups(group_map, alpha))
                            if len(groups) > 2:
                                show_table("Post-Hoc Games-Howell — alternatif Tukey saat varians/N tidak sama", games_howell_table(long_for_posthoc, alpha))
                        if effects.iloc[0]["Keputusan"].startswith("Signifikan") and len(groups) > 2:
                            show_table("Post-Hoc Tukey HSD", tukey_table(long_for_posthoc, alpha))
            else:
                dv = st.selectbox("Dependent variable / nilai", num_cols, key="anova_dv")
                grp = st.selectbox("Kolom grup", all_cols, key="anova_grp")
                if st.button("Jalankan ANOVA Long", type="primary"):
                    if df[grp].nunique(dropna=True) < 2:
                        st.error("Kolom grup perlu minimal 2 kategori.")
                    else:
                        table, effects, long_for_posthoc = anova_long(df, dv, grp, alpha)
                        show_table(f"ANOVA Table: {dv} berdasarkan {grp}", table)
                        show_table("Effect Size & Catatan Unequal N", effects)
                        with st.expander("📏 Ukuran grup dan opsi jika jumlah data tidak sama", expanded=effects.iloc[0].get("Status N") != "Seimbang"):
                            group_map = {str(level): safe_numeric(df.loc[df[grp].astype(str) == str(level), dv]).dropna() for level in df[grp].dropna().astype(str).unique().tolist()}
                            show_table("Ringkasan N per Grup", group_size_summary(group_map), "ANOVA biasa dapat menghitung jumlah grup yang tidak sama, tetapi interpretasi perlu melihat homogenitas varians.")
                            show_table("Welch ANOVA — alternatif aman untuk unequal N/varians", welch_anova_from_groups(group_map, alpha))
                            if df[grp].nunique(dropna=True) > 2:
                                show_table("Post-Hoc Games-Howell — alternatif Tukey saat varians/N tidak sama", games_howell_table(long_for_posthoc, alpha))
                        if effects.iloc[0]["Keputusan"].startswith("Signifikan") and df[grp].nunique(dropna=True) > 2:
                            show_table("Post-Hoc Tukey HSD", tukey_table(long_for_posthoc, alpha))

        elif test_type == "Two-Way ANOVA":
            if sm is None or smf is None:
                st.error("statsmodels belum tersedia. Pastikan requirements.txt sudah di-install.")
            else:
                dv = st.selectbox("Dependent variable / nilai", num_cols, key="twoway_dv")
                f1 = st.selectbox("Faktor 1", all_cols, key="twoway_f1")
                f2 = st.selectbox("Faktor 2", [c for c in all_cols if c != f1], key="twoway_f2")
                typ = st.selectbox("Type Sum of Squares", [2, 3], index=0)
                if st.button("Jalankan Two-Way ANOVA", type="primary"):
                    work = df[[dv, f1, f2]].copy()
                    work[dv] = safe_numeric(work[dv])
                    work = work.dropna()
                    if work[f1].nunique() < 2 or work[f2].nunique() < 2:
                        st.error("Masing-masing faktor perlu minimal 2 kategori.")
                    else:
                        formula = f'Q("{dv}") ~ C(Q("{f1}")) + C(Q("{f2}")) + C(Q("{f1}")):C(Q("{f2}"))'
                        model = smf.ols(formula, data=work).fit()
                        anova_tbl = sm.stats.anova_lm(model, typ=typ).reset_index().rename(columns={"index": "Source", "PR(>F)": "p-value"})
                        show_table(f"Two-Way ANOVA: {dv} ~ {f1} * {f2}", anova_tbl.round(5))
                        log_syntax(f"UNIANOVA {dv} BY {f1} {f2} /METHOD=SSTYPE({typ}) /DESIGN={f1} {f2} {f1}*{f2}.")

        elif test_type == "Korelasi":
            corr_cols = st.multiselect("Variabel numerik", num_cols, default=num_cols[: min(5, len(num_cols))])
            method = st.selectbox("Metode", ["pearson", "spearman", "kendall"])
            if st.button("Hitung Korelasi", type="primary"):
                if len(corr_cols) < 2:
                    st.error("Pilih minimal 2 variabel.")
                else:
                    corr = df[corr_cols].apply(pd.to_numeric, errors="coerce").corr(method=method)
                    show_table(f"Matriks Korelasi ({method})", corr.round(4))
                    rows = []
                    for i, a in enumerate(corr_cols):
                        for b in corr_cols[i + 1 :]:
                            pair = df[[a, b]].apply(pd.to_numeric, errors="coerce").dropna()
                            if len(pair) > 2:
                                if method == "pearson":
                                    r, p = stats.pearsonr(pair[a], pair[b])
                                elif method == "spearman":
                                    r, p = stats.spearmanr(pair[a], pair[b])
                                else:
                                    r, p = stats.kendalltau(pair[a], pair[b])
                                rows.append({"Variabel 1": a, "Variabel 2": b, "N": len(pair), "r/τ": r, "p-value": p, "Keputusan": decision_text(p, alpha)})
                    show_table("Signifikansi Korelasi", pd.DataFrame(rows).round(5))

        elif test_type == "Crosstab & Chi-Square":
            c1, c2 = st.columns(2)
            with c1:
                row_col = st.selectbox("Baris", all_cols)
            with c2:
                col_col = st.selectbox("Kolom", all_cols, index=1 if len(all_cols) > 1 else 0)
            if st.button("Buat Crosstab & Chi-Square", type="primary"):
                if row_col == col_col:
                    st.error("Pilih dua variabel berbeda.")
                else:
                    ct = pd.crosstab(df[row_col], df[col_col], margins=True)
                    show_table(f"Crosstab: {row_col} x {col_col}", ct.reset_index())
                    observed = pd.crosstab(df[row_col], df[col_col])
                    chi2, p, dof, expected = stats.chi2_contingency(observed)
                    n = observed.to_numpy().sum()
                    r, k = observed.shape
                    cramers_v = np.sqrt((chi2 / n) / min(k - 1, r - 1)) if min(k - 1, r - 1) > 0 else np.nan
                    result = pd.DataFrame([{"Chi-square": chi2, "df": dof, "p-value": p, "Cramer's V": cramers_v, "Keputusan": decision_text(p, alpha)}])
                    show_table("Chi-Square Test", result.round(5))
                    show_table("Expected Count", pd.DataFrame(expected, index=observed.index, columns=observed.columns).round(3))

        elif test_type == "Uji Nonparametrik":
            np_test = st.selectbox("Jenis uji", ["Mann-Whitney U", "Wilcoxon Signed-Rank", "Kruskal-Wallis", "Friedman Test"])
            if np_test == "Mann-Whitney U":
                c1, c2 = st.columns(2)
                with c1:
                    a = st.selectbox("Sampel 1", num_cols)
                with c2:
                    b = st.selectbox("Sampel 2", num_cols, index=1 if len(num_cols) > 1 else 0)
                if st.button("Jalankan Mann-Whitney U", type="primary"):
                    u, p = stats.mannwhitneyu(safe_numeric(df[a]).dropna(), safe_numeric(df[b]).dropna(), alternative="two-sided")
                    show_table("Mann-Whitney U", pd.DataFrame([{"U": u, "p-value": p, "Keputusan": decision_text(p, alpha)}]).round(5))
            elif np_test == "Wilcoxon Signed-Rank":
                c1, c2 = st.columns(2)
                with c1:
                    a = st.selectbox("Before/Pre", num_cols)
                with c2:
                    b = st.selectbox("After/Post", num_cols, index=1 if len(num_cols) > 1 else 0)
                if st.button("Jalankan Wilcoxon", type="primary"):
                    pair = df[[a, b]].apply(pd.to_numeric, errors="coerce").dropna()
                    w, p = stats.wilcoxon(pair[a], pair[b])
                    show_table("Wilcoxon Signed-Rank", pd.DataFrame([{"W": w, "p-value": p, "Keputusan": decision_text(p, alpha)}]).round(5))
            elif np_test == "Kruskal-Wallis":
                mode = st.radio("Format data", ["Wide: setiap kolom = grup", "Long: DV + kolom grup"], horizontal=True, key="kw_mode")
                if mode.startswith("Wide"):
                    groups = st.multiselect("Pilih kolom grup numerik", num_cols, default=num_cols[: min(3, len(num_cols))])
                    if st.button("Jalankan Kruskal-Wallis", type="primary"):
                        if len(groups) < 2:
                            st.error("Pilih minimal 2 grup.")
                        else:
                            h, p = stats.kruskal(*[safe_numeric(df[g]).dropna() for g in groups])
                            show_table("Kruskal-Wallis", pd.DataFrame([{"H": h, "df": len(groups) - 1, "p-value": p, "Keputusan": decision_text(p, alpha)}]).round(5))
                            if p < alpha and sp_posthoc is not None and len(groups) > 2:
                                wide_data = df[groups].apply(pd.to_numeric, errors="coerce")
                                show_table("Dunn Post-Hoc (p adjusted Bonferroni)", sp_posthoc.posthoc_dunn(wide_data.melt(value_name="nilai", var_name="grup").dropna(), val_col="nilai", group_col="grup", p_adjust="bonferroni").round(5))
                else:
                    dv = st.selectbox("Dependent variable / nilai", num_cols, key="kw_dv")
                    grp = st.selectbox("Kolom grup", all_cols, key="kw_grp")
                    if st.button("Jalankan Kruskal-Wallis Long", type="primary"):
                        work = df[[dv, grp]].copy()
                        work[dv] = safe_numeric(work[dv])
                        work = work.dropna()
                        groups_data = [g[dv] for _, g in work.groupby(grp)]
                        if len(groups_data) < 2:
                            st.error("Perlu minimal 2 grup.")
                        else:
                            h, p = stats.kruskal(*groups_data)
                            show_table("Kruskal-Wallis", pd.DataFrame([{"H": h, "df": len(groups_data) - 1, "p-value": p, "Keputusan": decision_text(p, alpha)}]).round(5))
                            if p < alpha and sp_posthoc is not None and len(groups_data) > 2:
                                show_table("Dunn Post-Hoc (p adjusted Bonferroni)", sp_posthoc.posthoc_dunn(work, val_col=dv, group_col=grp, p_adjust="bonferroni").round(5))
            else:
                groups = st.multiselect("Pilih kolom repeated measures", num_cols, default=num_cols[: min(3, len(num_cols))])
                if st.button("Jalankan Friedman", type="primary"):
                    if len(groups) < 3:
                        st.error("Friedman membutuhkan minimal 3 kondisi/kolom berpasangan.")
                    else:
                        pair = df[groups].apply(pd.to_numeric, errors="coerce").dropna()
                        stat, p = stats.friedmanchisquare(*[pair[g] for g in groups])
                        show_table("Friedman Test", pd.DataFrame([{"Chi-square": stat, "df": len(groups)-1, "p-value": p, "Keputusan": decision_text(p, alpha)}]).round(5))

        else:
            st.markdown("#### Normalitas, Homogenitas, dan Outlier")
            check_cols = st.multiselect("Variabel numerik", num_cols, default=num_cols[: min(4, len(num_cols))], key="assumption_cols")
            group_for_levene = st.selectbox("Kolom grup untuk Levene", ["(tidak ada)"] + all_cols)
            if st.button("Jalankan Uji Asumsi", type="primary"):
                if check_cols:
                    show_table("Normality Tests + Rekomendasi", normality_table(df, check_cols), "Shapiro dipakai untuk N ≤ 5000; D'Agostino untuk N besar. Jika p-value < α, jangan hanya berhenti; baca rekomendasi tindakan.")
                    with st.expander("Apa yang dilakukan jika normalitas tidak tercapai?", expanded=True):
                        st.dataframe(normality_solution_table(), use_container_width=True, hide_index=True)
                    out_rows = []
                    for col in check_cols:
                        s = safe_numeric(df[col])
                        z = np.abs(stats.zscore(s.dropna())) if s.dropna().std(ddof=1) != 0 else np.array([])
                        out_rows.append({"Variabel": col, "Outlier |z| > 3": int((z > 3).sum())})
                    show_table("Deteksi Outlier Sederhana", pd.DataFrame(out_rows))
                if group_for_levene != "(tidak ada)" and check_cols:
                    rows = []
                    for col in check_cols:
                        work = df[[col, group_for_levene]].copy()
                        work[col] = safe_numeric(work[col])
                        work = work.dropna()
                        groups = [g[col] for _, g in work.groupby(group_for_levene)]
                        if len(groups) >= 2 and all(len(g) >= 2 for g in groups):
                            stat, p = stats.levene(*groups)
                            rows.append({"Variabel": col, "Levene Statistic": stat, "p-value": p, "Interpretasi": "Homogen" if p >= alpha else "Tidak homogen"})
                    if rows:
                        show_table("Levene's Test", pd.DataFrame(rows).round(5))

    except BaseException as exc:
        if _is_streamlit_control_exception(exc):
            raise
        st.error("Bagian 🧪 Uji Statistik mengalami kendala, tetapi aplikasi tetap berjalan.")
        st.exception(exc)

elif active_section == '📈 Regresi':
    try:
        st.subheader("📈 Regresi")
        if sm is None:
            st.error("statsmodels belum tersedia. Pastikan requirements.txt sudah di-install.")
        else:
            model_options = ["Regresi Logistik Biner"] if not num_cols else ["Regresi Linear", "Regresi Logistik Biner"]
            model_type = st.selectbox("Model", model_options)
            if model_type == "Regresi Linear":
                y_col = st.selectbox("Dependent variable (Y)", num_cols)
                predictors = st.multiselect("Independent variables (X)", [c for c in all_cols if c != y_col], default=[c for c in num_cols if c != y_col][:2])
                if st.button("Jalankan Regresi Linear", type="primary"):
                    if not predictors:
                        st.error("Pilih minimal 1 prediktor.")
                    else:
                        work = df[[y_col] + predictors].dropna().copy()
                        y = safe_numeric(work[y_col])
                        X = make_design_matrix(work, predictors)
                        valid = y.notna()
                        y, X = y.loc[valid], X.loc[valid]
                        Xc = sm.add_constant(X, has_constant="add")
                        model = sm.OLS(y, Xc).fit()
                        fit_table = pd.DataFrame(
                            [
                                {
                                    "N": int(model.nobs),
                                    "R²": model.rsquared,
                                    "Adj. R²": model.rsquared_adj,
                                    "F": model.fvalue,
                                    "Prob(F)": model.f_pvalue,
                                    "AIC": model.aic,
                                    "BIC": model.bic,
                                }
                            ]
                        ).round(5)
                        coef_table = pd.DataFrame(
                            {
                                "Coef": model.params,
                                "Std.Err": model.bse,
                                "t": model.tvalues,
                                "p-value": model.pvalues,
                                "CI Lower": model.conf_int()[0],
                                "CI Upper": model.conf_int()[1],
                            }
                        ).reset_index(names="Term").round(5)
                        show_table("Model Summary - Regresi Linear", fit_table)
                        show_table("Coefficients - Regresi Linear", coef_table)
                        diag_rows = []
                        try:
                            jb_stat, jb_p, skew, kurt = sm.stats.jarque_bera(model.resid)
                            dw = sm.stats.durbin_watson(model.resid)
                            bp_stat, bp_p, _, _ = het_breuschpagan(model.resid, model.model.exog) if het_breuschpagan is not None else (np.nan, np.nan, np.nan, np.nan)
                            diag_rows.append({"Diagnostic": "Jarque-Bera normality", "Statistic": jb_stat, "p-value": jb_p, "Interpretasi": "Residual normal" if jb_p >= 0.05 else "Residual tidak normal"})
                            diag_rows.append({"Diagnostic": "Durbin-Watson", "Statistic": dw, "p-value": np.nan, "Interpretasi": "Mendekati 2 berarti autokorelasi rendah"})
                            diag_rows.append({"Diagnostic": "Breusch-Pagan heteroskedasticity", "Statistic": bp_stat, "p-value": bp_p, "Interpretasi": "Homoskedastik" if bp_p >= 0.05 else "Ada indikasi heteroskedastisitas"})
                            show_table("Regression Diagnostics", pd.DataFrame(diag_rows).round(5))
                        except Exception:
                            pass
                        if X.shape[1] >= 2:
                            show_table("Collinearity Diagnostics (VIF)", vif_table(X))
            else:
                y_col = st.selectbox("Dependent binary variable (0/1 atau 2 kategori)", all_cols)
                predictors = st.multiselect("Independent variables (X)", [c for c in all_cols if c != y_col], default=[c for c in num_cols if c != y_col][:2])
                if st.button("Jalankan Regresi Logistik", type="primary"):
                    work = df[[y_col] + predictors].dropna().copy()
                    y_raw = work[y_col]
                    uniques = y_raw.dropna().unique().tolist()
                    if len(uniques) != 2:
                        st.error("Regresi logistik biner memerlukan dependent variable dengan tepat 2 kategori.")
                    elif not predictors:
                        st.error("Pilih minimal 1 prediktor.")
                    else:
                        mapping = {uniques[0]: 0, uniques[1]: 1}
                        y = y_raw.map(mapping).astype(float)
                        X = make_design_matrix(work, predictors)
                        Xc = sm.add_constant(X, has_constant="add")
                        try:
                            model = sm.Logit(y, Xc).fit(disp=False)
                            coef = pd.DataFrame(
                                {
                                    "Coef(logit)": model.params,
                                    "Std.Err": model.bse,
                                    "z": model.tvalues,
                                    "p-value": model.pvalues,
                                    "Odds Ratio": np.exp(model.params),
                                    "OR CI Lower": np.exp(model.conf_int()[0]),
                                    "OR CI Upper": np.exp(model.conf_int()[1]),
                                }
                            ).reset_index(names="Term").round(5)
                            summary = pd.DataFrame(
                                [
                                    {
                                        "N": int(model.nobs),
                                        "Pseudo R²": model.prsquared,
                                        "LLR p-value": model.llr_pvalue,
                                        "AIC": model.aic,
                                        "BIC": model.bic,
                                        "Encoding Y": f"{uniques[0]}=0, {uniques[1]}=1",
                                    }
                                ]
                            ).round(5)
                            show_table("Model Summary - Regresi Logistik", summary)
                            show_table("Coefficients - Regresi Logistik", coef)
                        except Exception as exc:
                            st.error(f"Model tidak dapat diestimasi: {exc}")

    except BaseException as exc:
        if _is_streamlit_control_exception(exc):
            raise
        st.error("Bagian 📈 Regresi mengalami kendala, tetapi aplikasi tetap berjalan.")
        st.exception(exc)

elif active_section == '🔬 Analisis Lanjutan':
    try:
        render_advanced_research_analytics(df, num_cols, cat_cols, all_cols)
    except BaseException as exc:
        if _is_streamlit_control_exception(exc):
            raise
        st.error("Bagian 🔬 Analisis Lanjutan mengalami kendala, tetapi aplikasi tetap berjalan.")
        st.exception(exc)

elif active_section == '🧭 Reliabilitas & Faktor':
    try:
        st.subheader("🧭 Reliabilitas, PCA & EFA dengan Interpretasi Riset")
        st.markdown("### Cronbach's Alpha")
        item_cols = st.multiselect("Pilih item skala", num_cols, default=[c for c in num_cols if c.lower().startswith("item")][:6])
        if st.button("Hitung Cronbach's Alpha", type="primary"):
            if len(item_cols) < 2:
                st.error("Pilih minimal 2 item.")
            else:
                complete_cases = len(df[item_cols].apply(pd.to_numeric, errors="coerce").dropna())
                alpha_value, item_table = cronbach_alpha(df[item_cols])
                reliability_summary = pd.DataFrame([{"Cronbach's Alpha": alpha_value, "N Items": len(item_cols), "Complete Cases": complete_cases}]).round(5)
                show_table("Reliability Statistics", reliability_summary)
                show_table("Item-Total Statistics", item_table)
                reliability_interpretation = interpret_cronbach_reliability(alpha_value, item_table, n_items=len(item_cols), n_cases=complete_cases)
                show_table("Makna Riset - Reliabilitas", reliability_interpretation, note="Interpretasi ini membantu membaca kualitas instrumen, item yang perlu ditinjau, dan langkah tindak lanjut sebelum skor total dipakai.")

                with st.expander("📝 Contoh narasi laporan reliabilitas"):
                    level_text = reliability_interpretation.iloc[0]["Makna untuk riset"] if not reliability_interpretation.empty else "Reliabilitas belum dapat dimaknai."
                    st.markdown(f"""
Berdasarkan hasil uji reliabilitas, diperoleh nilai Cronbach's Alpha sebesar **{_format_number(alpha_value)}** dengan jumlah item **{len(item_cols)}**. {level_text}

Dengan demikian, keputusan penggunaan skor total/rata-rata skala perlu mempertimbangkan nilai alpha, korelasi item-total, serta kesesuaian item dengan konstruk teoritis penelitian.
""")

        st.markdown("### Principal Component Analysis (PCA)")
        pca_cols = st.multiselect("Variabel PCA", num_cols, default=num_cols[: min(5, len(num_cols))], key="pca_cols")

        # Streamlit slider membutuhkan min_value < max_value. Jika variabel PCA kurang dari 2,
        # jangan tampilkan slider agar aplikasi tidak crash pada dataset dengan 1 kolom numerik.
        if len(pca_cols) >= 2:
            max_components = min(len(pca_cols), 10)
            n_comp = st.slider("Jumlah komponen", 1, max_components, min(2, max_components))
        else:
            n_comp = 1
            st.info("Pilih minimal 2 variabel numerik untuk mengaktifkan PCA.")

        if st.button("Jalankan PCA"):
            if StandardScaler is None or PCA is None:
                st.error("scikit-learn belum tersedia. Pastikan requirements.txt sudah di-install.")
            elif len(pca_cols) < 2:
                st.error("Pilih minimal 2 variabel untuk PCA.")
            else:
                data = df[pca_cols].apply(pd.to_numeric, errors="coerce").dropna()
                if len(data) < 2:
                    st.error("Data lengkap untuk PCA minimal 2 baris setelah missing value dihapus.")
                else:
                    max_allowed_components = min(len(pca_cols), len(data), 10)
                    n_comp_safe = min(n_comp, max_allowed_components)
                    scaler = StandardScaler()
                    X_scaled = scaler.fit_transform(data)
                    pca = PCA(n_components=n_comp_safe)
                    pca.fit(X_scaled)
                    explained = pd.DataFrame(
                        {
                            "Component": [f"PC{i+1}" for i in range(n_comp_safe)],
                            "Eigenvalue": pca.explained_variance_,
                            "Explained Variance %": pca.explained_variance_ratio_ * 100,
                            "Cumulative %": np.cumsum(pca.explained_variance_ratio_) * 100,
                        }
                    ).round(5)
                    loadings = pd.DataFrame(pca.components_.T, index=pca_cols, columns=[f"PC{i+1}" for i in range(n_comp_safe)]).reset_index(names="Variable").round(5)
                    show_table("PCA Explained Variance", explained)
                    show_table("PCA Component Loadings", loadings)
                    pca_interpretation = interpret_pca_results(explained, loadings, selected_cols=pca_cols, n_rows=len(data))
                    show_table("Makna Riset - PCA", pca_interpretation, note="Interpretasi ini membantu menentukan apakah komponen sudah cukup menjelaskan data, variabel mana yang dominan, dan apa langkah berikutnya.")

                    with st.expander("📝 Contoh narasi laporan PCA"):
                        cum_val = None
                        cum_col = _first_existing_column(explained, ["Cumulative %", "Cumulative Variance %"])
                        if cum_col is not None:
                            vals = pd.to_numeric(explained[cum_col], errors="coerce").dropna()
                            if len(vals):
                                cum_val = float(vals.iloc[-1])
                        eig_col = _first_existing_column(explained, ["Eigenvalue"])
                        eig_count = None
                        if eig_col is not None:
                            eig_count = int((pd.to_numeric(explained[eig_col], errors="coerce") > 1).sum())
                        st.markdown(f"""
Analisis PCA dilakukan terhadap **{len(pca_cols)} variabel** dengan **{len(data)} kasus lengkap**. Komponen yang diekstraksi menjelaskan sekitar **{_format_number(cum_val, 2) if cum_val is not None else 'NA'}%** variasi data. Berdasarkan aturan eigenvalue > 1, terdapat sekitar **{eig_count if eig_count is not None else 'NA'}** komponen yang dapat dipertimbangkan.

Komponen dapat diberi nama berdasarkan variabel dengan loading tertinggi. Hasil PCA sebaiknya digunakan sebagai reduksi data/indikator ringkas, dan tetap dikaitkan dengan teori penelitian.
""")

        st.markdown("### Exploratory Factor Analysis (EFA)")
        efa_cols = st.multiselect("Variabel EFA", num_cols, default=[c for c in num_cols if c.lower().startswith("item")][:6], key="efa_cols")
        # Streamlit slider membutuhkan min_value < max_value.
        # Untuk 2 variabel, jumlah faktor maksimum yang valid hanya 1,
        # jadi jangan memakai slider karena min=1 dan max=1 akan membuat aplikasi crash.
        if len(efa_cols) >= 2:
            max_factors = max(1, min(len(efa_cols) - 1, 8))
            if max_factors >= 2:
                n_factors = st.slider("Jumlah faktor", 1, max_factors, min(2, max_factors), key="n_factors")
            else:
                n_factors = 1
                st.info("Dengan 2 variabel, EFA hanya dapat mengekstraksi 1 faktor. Tambahkan minimal 3 variabel jika ingin memilih lebih dari 1 faktor.")
        else:
            n_factors = 1
            st.info("Pilih minimal 2 variabel numerik untuk EFA.")
        rotation = st.selectbox("Rotasi", ["varimax", "promax", "oblimin", None], index=0)
        efa_engine = st.selectbox(
            "Engine EFA",
            ["Fallback stabil (Principal Axis Factoring)", "Otomatis: factor-analyzer lalu fallback"],
            index=0,
            help="Fallback stabil direkomendasikan agar EFA tetap jalan tanpa konflik factor-analyzer/scikit-learn.",
        )
        if st.button("Jalankan EFA"):
            if len(efa_cols) < 2:
                st.error("Pilih minimal 2 variabel.")
            else:
                data = df[efa_cols].apply(pd.to_numeric, errors="coerce").dropna()
                if len(data) < 5:
                    st.error("EFA sebaiknya memiliki minimal 5 baris lengkap; idealnya jauh lebih besar.")
                else:
                    try:
                        prefer_fallback = efa_engine.startswith("Fallback")
                        safe_n_factors = max(1, min(int(n_factors), max(1, min(len(efa_cols) - 1, len(data) - 1, 8))))
                        kmo_table, loadings, variance, communalities, eigen_table, efa_note = run_efa_analysis(
                            df, efa_cols, safe_n_factors, rotation, prefer_fallback=prefer_fallback
                        )
                        show_table("KMO & Bartlett's Test", kmo_table)
                        show_table("EFA Factor Loadings", loadings)
                        show_table("EFA Variance Explained", variance)
                        show_table("EFA Communalities", communalities)
                        show_table("EFA Eigenvalues", eigen_table)
                        efa_interpretation = interpret_efa_results(kmo_table, loadings, variance, communalities, eigen_table)
                        show_table("Makna Riset - EFA", efa_interpretation, note="Interpretasi ini membantu membaca kelayakan faktor, item lemah, cross-loading, dan makna dimensi yang terbentuk.")
                        if "fallback" in efa_note.lower():
                            st.info(efa_note)
                        else:
                            st.success(efa_note)
                    except Exception as exc:
                        st.error(f"EFA gagal dihitung: {exc}")
                        st.info("Coba pilih Engine EFA: `Fallback stabil (Principal Axis Factoring)`, lalu jalankan ulang.")

    except BaseException as exc:
        if _is_streamlit_control_exception(exc):
            raise
        st.error("Bagian 🧭 Reliabilitas & Faktor mengalami kendala, tetapi aplikasi tetap berjalan.")
        st.exception(exc)

elif active_section == '🎨 Visualisasi':
    try:
        st.subheader("🎨 Visualisasi")
        if px is None:
            st.error("plotly belum tersedia. Pastikan requirements.txt sudah di-install.")
        elif not all_cols:
            st.warning("Tidak ada kolom untuk divisualisasikan.")
        elif not num_cols:
            chart_type = st.selectbox("Jenis grafik", ["Bar Chart"])
            x = st.selectbox("Kategori", all_cols)
            plot_df = df[x].value_counts(dropna=False).reset_index()
            plot_df.columns = [x, "Frekuensi"]
            fig = px.bar(plot_df, x=x, y="Frekuensi")
            st.plotly_chart(fig, use_container_width=True)
        else:
            chart_type = st.selectbox("Jenis grafik", ["Histogram", "Box Plot", "Scatter Plot", "Bar Chart", "Correlation Heatmap", "Q-Q Plot"])
            if chart_type == "Histogram":
                x = st.selectbox("Variabel", num_cols)
                color = st.selectbox("Warna berdasarkan", [None] + all_cols)
                fig = px.histogram(df, x=x, color=color, marginal="box", nbins=30)
                st.plotly_chart(fig, use_container_width=True)
            elif chart_type == "Box Plot":
                y = st.selectbox("Variabel numerik", num_cols)
                x = st.selectbox("Kelompok", [None] + all_cols)
                fig = px.box(df, x=x, y=y, points="all")
                st.plotly_chart(fig, use_container_width=True)
            elif chart_type == "Scatter Plot":
                x = st.selectbox("X", num_cols)
                y = st.selectbox("Y", num_cols, index=1 if len(num_cols) > 1 else 0)
                color = st.selectbox("Color", [None] + all_cols)
                trendline = st.checkbox("Tambah trendline OLS", value=True)
                fig = px.scatter(df, x=x, y=y, color=color, trendline="ols" if trendline else None)
                st.plotly_chart(fig, use_container_width=True)
            elif chart_type == "Bar Chart":
                x = st.selectbox("Kategori", all_cols)
                y = st.selectbox("Nilai numerik (opsional)", [None] + num_cols)
                if y is None:
                    plot_df = df[x].value_counts().reset_index()
                    plot_df.columns = [x, "Frekuensi"]
                    fig = px.bar(plot_df, x=x, y="Frekuensi")
                else:
                    plot_df = df.groupby(x, dropna=False)[y].mean().reset_index()
                    fig = px.bar(plot_df, x=x, y=y)
                st.plotly_chart(fig, use_container_width=True)
            elif chart_type == "Correlation Heatmap":
                corr_cols = st.multiselect("Variabel", num_cols, default=num_cols[: min(6, len(num_cols))], key="heatmap_cols")
                if len(corr_cols) >= 2:
                    corr = df[corr_cols].corr(numeric_only=True)
                    fig = px.imshow(corr, text_auto=True, aspect="auto")
                    st.plotly_chart(fig, use_container_width=True)
                else:
                    st.info("Pilih minimal 2 variabel.")
            else:
                import plotly.graph_objects as go
                x = st.selectbox("Variabel", num_cols, key="qq_var")
                s_num = safe_numeric(df[x]).dropna().sort_values()
                if len(s_num) >= 3:
                    theoretical = stats.norm.ppf((np.arange(1, len(s_num) + 1) - 0.5) / len(s_num), loc=s_num.mean(), scale=s_num.std(ddof=1))
                    fig = go.Figure()
                    fig.add_scatter(x=theoretical, y=s_num, mode="markers", name="Data")
                    fig.add_scatter(x=[theoretical.min(), theoretical.max()], y=[theoretical.min(), theoretical.max()], mode="lines", name="Normal line")
                    fig.update_layout(xaxis_title="Theoretical Quantiles", yaxis_title="Sample Quantiles")
                    st.plotly_chart(fig, use_container_width=True)
                else:
                    st.info("Q-Q plot membutuhkan minimal 3 data non-missing.")


    except BaseException as exc:
        if _is_streamlit_control_exception(exc):
            raise
        st.error("Bagian 🎨 Visualisasi mengalami kendala, tetapi aplikasi tetap berjalan.")
        st.exception(exc)

elif active_section == '🧠 Insight Riset':
    try:
        st.subheader("🧠 Insight & Makna Riset")
        st.markdown(
            """
            Modul ini mengubah output statistik menjadi **narasi riset**: apa makna hasilnya,
            seberapa kuat buktinya, apa implikasinya, dan apa langkah analisis berikutnya.
            Jalankan analisis di tab lain terlebih dahulu, lalu gunakan bagian ini untuk menyusun pembahasan.
            """
        )

        if not st.session_state.report_items:
            st.info("Belum ada output tersimpan. Jalankan deskriptif/uji statistik/regresi/reliabilitas/EFA terlebih dahulu.")
        else:
            insight_alpha = st.slider("α untuk interpretasi insight", 0.01, 0.10, float(st.session_state.get("active_alpha", 0.05)), 0.01, key="insight_alpha")
            output_titles = [item.get("title", f"Output {i+1}") for i, item in enumerate(st.session_state.report_items)]
            selected_outputs = st.multiselect(
                "Pilih output yang ingin dimaknai",
                output_titles,
                default=output_titles,
                help="Pilih semua output untuk sintesis lengkap, atau pilih beberapa output untuk insight spesifik.",
            )

            c_syn, c_detail = st.columns([0.45, 0.55])
            with c_syn:
                st.markdown("### Sintesis Besar")
                synthesis = build_research_synthesis(
                    [item for item in st.session_state.report_items if item.get("title") in selected_outputs],
                    alpha=insight_alpha,
                )
                st.dataframe(synthesis, use_container_width=True)
                st.caption("Sintesis ini membantu menyusun bagian pembahasan, bukan menggantikan penilaian teoritis peneliti.")

            with c_detail:
                st.markdown("### Insight per Output")
                insight_table = build_insight_table(st.session_state.report_items, alpha=insight_alpha, selected_titles=selected_outputs)
                st.dataframe(insight_table, use_container_width=True)

            st.markdown("### Template Narasi Pembahasan")
            narasi = []
            for _, row in synthesis.iterrows():
                narasi.append(f"**{row['Bagian']}.** {row['Sintesis']}")
            narasi_text = "\n\n".join(narasi)
            st.markdown(narasi_text)

            c_save1, c_save2 = st.columns(2)
            with c_save1:
                if st.button("💾 Simpan Sintesis ke Output Viewer"):
                    add_report("Sintesis Insight Riset", synthesis, "Sintesis otomatis untuk membantu pemaknaan hasil penelitian.")
                    st.success("Sintesis insight disimpan ke Output Viewer.")
            with c_save2:
                if st.button("💾 Simpan Detail Insight ke Output Viewer"):
                    add_report("Detail Insight per Output", insight_table, "Insight otomatis per output statistik.")
                    st.success("Detail insight disimpan ke Output Viewer.")

            st.download_button(
                "⬇️ Download Narasi Insight Markdown",
                data=narasi_text.encode("utf-8"),
                file_name="statistik_pro_insight_riset.md",
                mime="text/markdown",
            )

            with st.expander("📌 Cara membaca insight ini"):
                st.markdown(
                    """
                    - **Signifikan** berarti ada bukti statistik pada α yang dipilih, bukan otomatis penting secara praktis.
                    - **Tidak signifikan** berarti bukti belum cukup, bukan bukti mutlak bahwa efek/hubungan tidak ada.
                    - Selalu baca hasil bersama **effect size**, ukuran sampel, kualitas instrumen, asumsi statistik, dan teori.
                    - Untuk skripsi/tesis/artikel, gunakan insight ini sebagai draf pembahasan lalu sesuaikan dengan konteks variabel dan literatur.
                    """
                )

    except BaseException as exc:
        if _is_streamlit_control_exception(exc):
            raise
        st.error("Bagian 🧠 Insight Riset mengalami kendala, tetapi aplikasi tetap berjalan.")
        st.exception(exc)

elif active_section == '📤 Output & Ekspor':
    try:
        st.subheader("📤 Ekspor Data & Output")
        st.markdown("Gunakan bagian ini untuk menyimpan data aktif dan output analisis seperti jendela **Output Viewer** di SPSS.")

        c1, c2 = st.columns(2)
        with c1:
            csv_bytes = df.to_csv(index=False).encode("utf-8")
            st.download_button("⬇️ Download Data CSV", data=csv_bytes, file_name="statistik_pro_data.csv", mime="text/csv")
        with c2:
            excel_bytes = get_excel_download(df, st.session_state.report_items)
            st.download_button(
                "⬇️ Download Data + Output Excel",
                data=excel_bytes,
                file_name="statistik_pro_output.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )

        st.markdown("### Ekspor Laporan")
        if st.session_state.report_items:
            md_bytes = report_as_markdown(st.session_state.report_items).encode("utf-8")
            html_bytes = report_as_html(st.session_state.report_items).encode("utf-8")
            c3, c4, c5 = st.columns(3)
            with c3:
                st.download_button("⬇️ Download Report Markdown", data=md_bytes, file_name="statistik_pro_report.md", mime="text/markdown")
            with c4:
                st.download_button("⬇️ Download Report HTML", data=html_bytes, file_name="statistik_pro_report.html", mime="text/html")
            with c5:
                docx_bytes = report_as_docx(st.session_state.report_items)
                if docx_bytes is not None:
                    st.download_button("⬇️ Download Report Word", data=docx_bytes, file_name="statistik_pro_report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                else:
                    st.caption("Install `python-docx` untuk ekspor Word.")

        st.markdown("### Syntax / Audit Trail")
        if st.session_state.syntax_log:
            syntax_df = pd.DataFrame(st.session_state.syntax_log)
            st.dataframe(syntax_df, use_container_width=True)
            st.download_button("⬇️ Download Syntax Log", data="\n".join(syntax_df["command"].astype(str)).encode("utf-8"), file_name="statistik_pro_syntax.sps", mime="text/plain")
        else:
            st.info("Syntax log masih kosong. Transformasi dan beberapa analisis akan tercatat otomatis.")

        st.markdown("### Output Tersimpan")
        if st.session_state.report_items:
            for i, item in enumerate(st.session_state.report_items, start=1):
                with st.expander(f"{i}. {item['title']} — {item['created_at']}"):
                    if item.get("note"):
                        st.caption(item["note"])
                    if "table" in item:
                        st.dataframe(item["table"], use_container_width=True)
            if st.button("🧽 Bersihkan Output Viewer"):
                st.session_state.report_items = []
                st.rerun()
        else:
            st.info("Belum ada output tersimpan. Jalankan analisis di tab lain terlebih dahulu.")

    except BaseException as exc:
        if _is_streamlit_control_exception(exc):
            raise
        st.error("Bagian 📤 Output & Ekspor mengalami kendala, tetapi aplikasi tetap berjalan.")
        st.exception(exc)


with st.expander("📖 Catatan Metodologis"):
    st.markdown(
        """
        - Aplikasi ini membantu analisis statistik umum, tetapi bukan pengganti validasi metodologis dari peneliti/statistikawan.
        - Cek skala data, independensi observasi, normalitas, homogenitas varian, dan ukuran sampel sebelum menarik kesimpulan.
        - Untuk publikasi akademik, laporkan statistik uji, derajat kebebasan, p-value, confidence interval, effect size, dan asumsi yang diuji.
        - Hasil p-value sebaiknya dibaca bersama effect size dan konteks riset, bukan sebagai satu-satunya dasar keputusan.
        """
    )

st.markdown('<div class="statpro-footer-spacer"></div>', unsafe_allow_html=True)
# Footer fixed sudah dirender di awal; sidebar footer sengaja dihapus mulai v5.2.
